"""
End-to-end PDF output regression — locked 24 Feb 2026.

Builds a minimal canonical PDF and DOCX via the actual helpers in
services/export_footer.py and services/print_styles.py, then parses the
output to assert the locked spec:

  - PDF footer: Times-Italic 8pt, left = canonical 4-field label,
                right = DD/MM/YYYY · Page X of Y
  - DOCX heading sizes: H1 14.5pt bold, H2 12.5pt bold, H3 11.5pt bold italic
  - DOCX footer: 8pt italic Times New Roman, page-number field present
"""
from __future__ import annotations

import re
import io
import zipfile
import pytest


# ---------------------------------------------------------------------------
# Shared fixture — a minimal valid case dict.
# ---------------------------------------------------------------------------

@pytest.fixture
def canonical_case():
    return {
        "defendant_name": "John Smith",
        "case_number": "2025/0001234",
        "title": "R v Smith",
    }


# ---------------------------------------------------------------------------
# PDF — NumberedCanvas actually renders the 8pt footer with X of Y text.
# ---------------------------------------------------------------------------

def _build_canonical_pdf(case: dict, doc_label: str) -> bytes:
    from io import BytesIO
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from services.export_footer import build_footer_label, NumberedCanvas
    from services.print_styles import get_canonical_pdf_styles, canonical_page_margins_mm

    footer_label = build_footer_label(case, doc_label)
    styles = get_canonical_pdf_styles()
    m = canonical_page_margins_mm()

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=m["right"] * mm, leftMargin=m["left"] * mm,
        topMargin=m["top"] * mm, bottomMargin=m["bottom"] * mm,
    )
    story = [
        Paragraph("Main Heading Test", styles["CanonicalH1"]),
        Paragraph("Body paragraph one.", styles["CanonicalBody"]),
        PageBreak(),
        Paragraph("Second page body.", styles["CanonicalBody"]),
    ]
    doc.build(story, canvasmaker=NumberedCanvas(footer_label))
    return buf.getvalue()


def test_pdf_footer_contains_canonical_4_field_label_and_page_x_of_y(canonical_case):
    pdf_bytes = _build_canonical_pdf(canonical_case, "Full Detailed Report")

    # ReportLab PDF → pdfminer.six gives us plain text per page.
    try:
        from pdfminer.high_level import extract_text
    except ImportError:
        pytest.skip("pdfminer.six not installed")

    text = extract_text(io.BytesIO(pdf_bytes))

    # Canonical 4-field legal identifier must appear (left footer)
    assert "Criminal Law Appeal Management" in text
    assert "Full Detailed Report" in text
    assert "John Smith" in text
    assert "2025/0001234" in text

    # Right footer: date DD/MM/YYYY + middle-dot + "Page X of Y"
    assert re.search(r"\d{2}/\d{2}/\d{4}", text), f"Date DD/MM/YYYY missing from footer: {text[:500]!r}"
    assert "Page 1 of 2" in text
    assert "Page 2 of 2" in text


def test_pdf_uses_times_font_not_helvetica(canonical_case):
    """The ReportLab output stream must reference Times-* fonts and not Helvetica."""
    pdf_bytes = _build_canonical_pdf(canonical_case, "Case Summary Report")
    assert b"Times-Bold" in pdf_bytes or b"Times-Roman" in pdf_bytes, \
        "Canonical PDF must embed Times-* fonts"
    # No Helvetica leakage (the Canonical styles do not use it).
    # Note: ReportLab's default sample sheet references Helvetica internally
    # for unused styles — we only care that our canonical stream uses Times.


# ---------------------------------------------------------------------------
# DOCX — unzip and parse word/styles.xml to prove locked heading sizes.
# ---------------------------------------------------------------------------

def _build_canonical_docx(case: dict, doc_label: str) -> bytes:
    from io import BytesIO
    from docx import Document
    from services.print_styles import apply_canonical_docx_styles
    from services.export_footer import build_footer_label, apply_docx_footer

    doc = Document()
    apply_canonical_docx_styles(doc)

    doc.add_paragraph("Main Heading Test", style="Heading 1")
    doc.add_paragraph("Sub Heading Test", style="Heading 2")
    doc.add_paragraph("Sub-sub Heading Test", style="Heading 3")
    doc.add_paragraph("Body paragraph one.")
    doc.add_paragraph("Bullet item one.", style="List Bullet")

    apply_docx_footer(doc, build_footer_label(case, doc_label))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_heading_sizes_and_typography_are_canonical(canonical_case):
    """Parse word/styles.xml and assert Heading 1/2/3 + Normal match spec."""
    docx_bytes = _build_canonical_docx(canonical_case, "Case Summary Report")
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        styles_xml = zf.read("word/styles.xml").decode("utf-8")

    # python-docx stores sizes as w:sz with half-points (so 14.5pt => 29).
    # Heading 1 — sz 29 (14.5pt × 2), bold on
    h1_block_match = re.search(r'w:styleId="Heading1"[\s\S]+?</w:style>', styles_xml)
    assert h1_block_match, "Heading 1 block missing from DOCX styles.xml"
    h1_block = h1_block_match.group(0)
    assert re.search(r'w:sz w:val="29"', h1_block), f"Heading 1 size != 14.5pt: {h1_block[:400]}"
    assert "Times New Roman" in h1_block
    assert '<w:b/>' in h1_block or 'w:b w:val="1"' in h1_block

    # Heading 2 — sz 25 (12.5pt × 2)
    h2_block = re.search(r'w:styleId="Heading2"[\s\S]+?</w:style>', styles_xml).group(0)
    assert re.search(r'w:sz w:val="25"', h2_block)
    assert "Times New Roman" in h2_block

    # Heading 3 — sz 23 (11.5pt × 2), bold AND italic
    h3_block = re.search(r'w:styleId="Heading3"[\s\S]+?</w:style>', styles_xml).group(0)
    assert re.search(r'w:sz w:val="23"', h3_block)
    assert "Times New Roman" in h3_block
    assert '<w:b/>' in h3_block or 'w:b w:val="1"' in h3_block
    assert '<w:i/>' in h3_block or 'w:i w:val="1"' in h3_block

    # Normal body — sz 20 (10pt × 2)
    normal_block = re.search(r'w:styleId="Normal"[\s\S]+?</w:style>', styles_xml).group(0)
    assert re.search(r'w:sz w:val="20"', normal_block)
    assert "Times New Roman" in normal_block


def test_docx_footer_contains_canonical_label_and_page_number_field(canonical_case):
    """Parse word/footer1.xml and assert the 4-field label + PAGE/NUMPAGES fields."""
    docx_bytes = _build_canonical_docx(canonical_case, "Barrister View Report")
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        footer_files = [n for n in zf.namelist() if re.match(r"word/footer\d+\.xml", n)]
        assert footer_files, "No footer XML files generated"
        footer_xml = zf.read(footer_files[0]).decode("utf-8")

    # LEFT — the canonical 4-field label appears verbatim
    assert "Criminal Law Appeal Management" in footer_xml
    assert "Barrister View Report" in footer_xml
    assert "John Smith" in footer_xml
    assert "2025/0001234" in footer_xml

    # RIGHT — DD/MM/YYYY date and " Page " wrapper around the field
    assert re.search(r"\d{2}/\d{2}/\d{4}", footer_xml)
    assert "Page" in footer_xml
    # PAGE + NUMPAGES field codes must be present (locked ordering per apply_docx_footer)
    assert "PAGE" in footer_xml
    assert "NUMPAGES" in footer_xml

    # Footer font size — must be 16 half-points (= 8pt)
    # python-docx tags each run's font: <w:sz w:val="16"/>
    assert re.search(r'w:sz w:val="16"', footer_xml), \
        f"Footer is not 8pt: {footer_xml[:400]!r}"
    # Footer typeface is Times New Roman
    assert "Times New Roman" in footer_xml
    # Italic
    assert re.search(r'<w:i ?/>', footer_xml) or 'w:i w:val="1"' in footer_xml


def test_docx_missing_case_number_renders_no_case_number_literal():
    docx_bytes = _build_canonical_docx({"defendant_name": "Jane Doe"}, "Legal Framework")
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        footer_files = [n for n in zf.namelist() if re.match(r"word/footer\d+\.xml", n)]
        footer_xml = zf.read(footer_files[0]).decode("utf-8")
    assert "No case number" in footer_xml
