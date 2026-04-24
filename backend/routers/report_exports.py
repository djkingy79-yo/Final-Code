# ===========================================================================
# DO NOT UNDO — ENTIRE FILE PROTECTED
# Report PDF & DOCX Export Routes (extracted from server.py)
# All ReportLab styles, python-docx styles, font sizes, cover page layout,
# spacers, footers, and disclaimer text in this file are approved by Deb King
# and must be preserved. Font sizes are calibrated to match the frontend print
# CSS exactly (Body 11pt, H1 16pt, H2 15pt, H3 13pt, tables 9pt, footer 7pt).
# Do not refactor, rename, or reduce without explicit instruction.
# ===========================================================================

import re
from datetime import datetime, timezone

from fastapi import APIRouter, HTTPException, Request, Response

from config import db, logger
from auth_utils import get_current_user
from services.offence_helpers import get_export_legal_refs
from services.report_quality import _strip_report_placeholders, _clean_sentence_candidate, _is_valid_sentence_candidate
from services.barrister_generator import _coerce_utc_datetime

router = APIRouter(prefix="/api")


# ============ SHARED HELPERS ============
def _format_law_section(section: dict) -> str:
    """Render a law_sections[] entry as a single clean line.
    Handles: (a) Chapter/Part/Division/Schedule prefixes (no 's.' prefix),
    (b) duplicated jurisdiction (act already ends with '(NSW)') to avoid '(NSW) (NSW)'."""
    sec = (section.get("section") or "").strip()
    act = (section.get("act") or "").strip()
    juris = (section.get("jurisdiction") or "").strip().upper()

    # Decide section prefix
    sec_lower = sec.lower()
    structural = any(sec_lower.startswith(p) for p in ("chapter ", "part ", "division ", "schedule ", "subpart "))
    if structural or not sec:
        sec_display = sec
    else:
        sec_display = f"s {sec}" if not sec_lower.startswith(("s ", "s.", "ss ", "ss.")) else sec

    # Avoid duplicated jurisdiction
    act_tail = act[-len(f"({juris})"):].upper() if juris else ""
    has_juris = bool(juris) and act_tail == f"({juris})"
    juris_display = "" if has_juris else (f" ({juris})" if juris else "")

    return f"{sec_display} {act}{juris_display}".strip()


def _format_evidence_item(evidence) -> str:
    """Extract clean text from a supporting_evidence entry. Handles:
    (a) dict with quote/text/description, (b) Python-dict-looking string
    (legacy LLM output like \"{'document_id': 'optional', 'quote': 'real text'}\"),
    (c) plain string."""
    if isinstance(evidence, dict):
        for key in ("quote", "text", "description", "detail"):
            val = evidence.get(key)
            if val and isinstance(val, str) and val.strip() and val.strip().lower() != "optional":
                return val.strip()
        return ""
    if isinstance(evidence, str):
        if evidence.startswith("{") and "'optional'" in evidence:
            m = re.search(r"['\"]quote['\"]\s*:\s*['\"](.+?)['\"]\s*[,}]", evidence)
            return m.group(1) if m else ""
        return evidence.strip()
    return str(evidence).strip()


# ============ PDF EXPORT ============

def _format_export_display_date(value=None) -> str:
    parsed = _coerce_utc_datetime(value)
    if not parsed:
        parsed = datetime.now(timezone.utc)
    return parsed.astimezone(timezone.utc).strftime("%d/%m/%Y")


from services.export_footer import build_footer_label, NumberedCanvas, apply_docx_footer, REPORT_TYPE_LABELS  # noqa: E402


def _build_export_footer_label(case: dict, report_label: str, generated_at=None) -> str:
    return build_footer_label(case, report_label, generated_at)


def _build_export_footer_message() -> str:
    return ""


def _truncate_export_footer(text: str, max_chars: int = 118) -> str:
    value = (text or "").strip()
    if len(value) <= max_chars:
        return value
    return value[: max_chars - 1].rstrip() + "…"


def _extract_sentence_from_text(case: dict, analysis: str) -> str:
    if case.get('sentence') and str(case.get('sentence')).strip():
        return str(case.get('sentence')).strip()
    patterns = [
        r"(?:sentence\s+imposed\s+was|sentence\s+was|head\s+sentence\s+was|head\s+sentence:|sentenced?\s+to)\s+([^\.\n]{8,160})",
        r"(\d+\s+years?'?\s+with\s+a\s+non[- ]?parole\s+period\s+of\s+\d+\s+years?(?:\s+and\s+\d+\s+months?)?)",
        r"(\d+\s+years?'?(?:\s+and\s+\d+\s+months?)?\s*(?:imprisonment|gaol|jail|custody)?\s*(?:with\s+(?:a\s+)?non[- ]?parole\s+period\s+of\s+\d+\s+years?(?:\s+and\s+\d+\s+months?)?)?)",
        r"(?:was\s+)?sentenced?\s+to\s+(\d+\s*(?:years?|months?)\s+(?:and\s+\d+\s*(?:years?|months?)\s+)?(?:imprisonment|gaol|jail|custody)[^\n\.]{0,80})",
        r"(?:^|\n)\s*(?:Head\s+)?Sentence\s*:\s*(\d+[^\n]{5,100})",
        r"(?:sentenced?\s+to\s+)?(life\s+imprisonment|imprisonment\s+for\s+life|life\s+sentence)[^\n\.]{0,60}",
        r"(\d+[\s-]*(?:years?|months?)(?:'s?)?\s*(?:imprisonment|gaol|jail|custody|sentence|non[- ]?parole)[^\n\.]{0,80})",
        r"sentence\s+of\s+(\d+[^\n\.]{5,80})",
        r"((?:minimum|non[- ]?parole)\s+(?:period\s+)?of\s+\d+[^\n\.]{3,60})",
    ]
    for pattern in patterns:
        match = re.search(pattern, analysis or "", re.I | re.M)
        if match:
            candidate = _clean_sentence_candidate(match.group(1))
            if _is_valid_sentence_candidate(candidate):
                return candidate
    return "See report analysis"


async def _derive_export_sentence(case: dict, case_id: str, user_id: str, fallback_report: dict | None = None) -> str:
    standard_reports = await db.reports.find(
        {
            "case_id": case_id,
            "user_id": user_id,
            "report_type": {"$in": ["quick_summary", "full_detailed", "extensive_log", "barrister_view"]},
            "status": "completed",
        },
        {"_id": 0},
    ).sort("generated_at", -1).to_list(20)

    for report_type in ["quick_summary", "full_detailed", "extensive_log", "barrister_view"]:
        for item in [r for r in standard_reports if r.get("report_type") == report_type]:
            candidate = _extract_sentence_from_text(case, ((item.get("content") or {}).get("analysis") or ""))
            if candidate != "See report analysis":
                return candidate

    if fallback_report:
        candidate = _extract_sentence_from_text(case, ((fallback_report.get("content") or {}).get("analysis") or ""))
        if candidate != "See report analysis":
            return candidate

    return _extract_sentence_from_text(case, "")

@router.get("/cases/{case_id}/reports/{report_id}/export-pdf")
async def export_report_pdf(case_id: str, report_id: str, request: Request):
    """Export a report as PDF with Grounds of Merit and Legal References"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from io import BytesIO
    
    user = await get_current_user(request)
    
    # Get report
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get grounds of merit for this case
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Create PDF buffer
    # Canonical page margins (locked 24 Feb 2026).
    from services.print_styles import (
        canonical_page_margins_mm,
        CANONICAL_H1_PT, CANONICAL_H2_PT, CANONICAL_H3_PT,
        CANONICAL_BODY_PT,
    )
    _m = canonical_page_margins_mm()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=_m["right"]*mm,
        leftMargin=_m["left"]*mm,
        topMargin=_m["top"]*mm,
        bottomMargin=_m["bottom"]*mm
    )

    # Canonical typography (locked 24 Feb 2026): Times New Roman everywhere,
    # H1 14.5pt bold, H2 12.5pt bold, H3 11.5pt bold italic, body 10pt.
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ReportTitle',
        fontSize=CANONICAL_H1_PT,
        spaceAfter=6,
        alignment=TA_CENTER,
        fontName='Times-Bold',
        leading=CANONICAL_H1_PT * 1.25,
    ))
    styles.add(ParagraphStyle(
        name='ReportSubtitle',
        fontSize=CANONICAL_BODY_PT,
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName='Times-Roman',
        textColor=colors.HexColor('#475569')
    ))
    styles.add(ParagraphStyle(
        name='SectionHeader',
        fontSize=CANONICAL_H1_PT,
        spaceBefore=10,
        spaceAfter=6,
        fontName='Times-Bold',
        textColor=colors.HexColor('#0f172a'),
        keepWithNext=True
    ))
    styles.add(ParagraphStyle(
        name='SubHeader',
        fontSize=CANONICAL_H2_PT,
        spaceBefore=8,
        spaceAfter=3,
        fontName='Times-Bold',
        textColor=colors.HexColor('#1e293b'),
        keepWithNext=True
    ))
    styles.add(ParagraphStyle(
        name='ReportBodyText',
        fontSize=CANONICAL_BODY_PT,
        spaceAfter=6,
        alignment=TA_JUSTIFY,
        leading=CANONICAL_BODY_PT * 1.35,
        fontName='Times-Roman'
    ))
    styles.add(ParagraphStyle(
        name='BulletText',
        parent=styles['ReportBodyText'],
        leftIndent=20,
        bulletIndent=10,
        fontName='Times-Roman',
        fontSize=CANONICAL_BODY_PT,
        leading=CANONICAL_BODY_PT * 1.35
    ))
    styles.add(ParagraphStyle(
        name='LawSection',
        fontSize=CANONICAL_BODY_PT,
        spaceAfter=2,
        leftIndent=20,
        fontName='Times-Italic',
        textColor=colors.HexColor('#1e40af'),
        leading=CANONICAL_BODY_PT * 1.35
    ))
    styles.add(ParagraphStyle(
        name='GroundTitle',
        fontSize=CANONICAL_H3_PT,
        spaceBefore=6,
        spaceAfter=3,
        fontName='Times-BoldItalic',
        textColor=colors.HexColor('#1e3a8a'),
        keepWithNext=True
    ))
    styles.add(ParagraphStyle(
        name='NumberedSectionHeader',
        fontSize=CANONICAL_H1_PT,
        spaceBefore=10,
        spaceAfter=6,
        fontName='Times-Bold',
        textColor=colors.HexColor('#0f172a'),
        keepWithNext=True
    ))
    styles.add(ParagraphStyle(
        name='CoverMetaLabel',
        fontSize=8,
        fontName='Times-Bold',
        textColor=colors.HexColor('#64748b'),
        spaceAfter=1,
    ))
    styles.add(ParagraphStyle(
        name='CoverMetaValue',
        fontSize=10,
        fontName='Times-Roman',
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=3,
        leading=13,
    ))
    # DO_NOT_UNDO — Bold red disclaimer with white text in PDF
    styles.add(ParagraphStyle(
        name='CoverDisclaimer',
        fontSize=9,
        fontName='Times-Bold',
        textColor=colors.HexColor('#dc2626'),
        alignment=TA_CENTER,
        leading=12,
    ))

    def format_inline(text: str) -> str:
        clean = (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Convert markdown bold
        clean = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", clean)
        # Convert markdown links [text](url) to just text
        clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", clean)
        return clean

    def render_table(lines):
        rows = []
        for line in lines:
            parts = [p.strip() for p in line.strip().strip("|").split("|")]
            if not parts:
                continue
            if all(set(p) <= set("-:") for p in parts):
                continue
            # Escape special characters in table cells for reportlab
            safe_parts = []
            for p in parts:
                cell = re.sub(r"\*\*(.*?)\*\*", r"\1", p)
                cell = cell.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_parts.append(cell)
            rows.append(safe_parts)
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        rows = [r + [""] * (col_count - len(r)) for r in rows]
        try:
            col_width = doc.width / col_count
            para_rows = []
            cell_style = ParagraphStyle(name='CellText', fontSize=9, leading=12, fontName='Times-Roman', wordWrap='CJK')
            header_style = ParagraphStyle(name='HeaderCellText', fontSize=9, leading=12, fontName='Times-Bold', textColor=colors.white)
            for ri, row in enumerate(rows):
                style = header_style if ri == 0 else cell_style
                para_rows.append([Paragraph(c[:260], style) for c in row])
            table = Table(para_rows, colWidths=[col_width] * col_count, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1d4ed8')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            story.append(table)
            story.append(Spacer(1, 3*mm))
        except Exception as e:
            logger.warning(f"PDF table render failed: {e}")
            for row in rows:
                story.append(Paragraph(" | ".join(row), styles['ReportBodyText']))

    def render_markdown(text: str):
        lines = (text or "").splitlines()
        buffer = []
        table_lines = []

        def flush_paragraph():
            if buffer:
                paragraph = " ".join(buffer).strip()
                if paragraph:
                    try:
                        story.append(Paragraph(format_inline(paragraph), styles['ReportBodyText']))
                        story.append(Spacer(1, 1*mm))
                    except Exception as e:
                        logger.warning(f"PDF paragraph failed: {e}")
                        # Fallback: strip all XML-like content
                        safe = re.sub(r'<[^>]+>', '', paragraph)
                        story.append(Paragraph(safe, styles['ReportBodyText']))
            buffer.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                flush_paragraph()
                continue
            if stripped.startswith("|") and "|" in stripped:
                flush_paragraph()
                table_lines.append(stripped)
                continue
            if table_lines and not (stripped.startswith("|") and "|" in stripped):
                render_table(table_lines)
                table_lines = []
            if stripped.startswith("## "):
                flush_paragraph()
                # New major section starts on new page
                story.append(PageBreak())
                story.append(Paragraph(format_inline(stripped[3:].strip()), styles['SectionHeader']))
                story.append(Spacer(1, 2*mm))
                continue
            if re.match(r'^\d+\.\s+[A-Z][A-Z0-9\s/&()\-]{4,}$', stripped):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped), styles['NumberedSectionHeader']))
                story.append(Spacer(1, 1*mm))
                continue
            if stripped.startswith("### "):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped[4:].strip()), styles['SubHeader']))
                continue
            if stripped.startswith("#### "):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped[5:].strip()), styles['SubHeader']))
                continue
            if re.match(r'^Ground\s+\d+\s*:', stripped, re.I):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped), styles['GroundTitle']))
                continue
            if stripped.startswith("- ") or stripped.startswith("• "):
                flush_paragraph()
                bullet_text = stripped[2:].strip()
                story.append(Paragraph(format_inline(f"- {bullet_text}"), styles['BulletText']))
                continue
            # Handle numbered lists
            if re.match(r'^\d+\.\s', stripped):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped), styles['BulletText']))
                continue
            buffer.append(stripped)

        flush_paragraph()
        if table_lines:
            render_table(table_lines)

    story = []

    # Report Title
    report_type_labels = {
        'quick_summary': 'Case Summary Report (Free)',
        'full_detailed': 'Full Detailed Report ($150.00)',
        'extensive_log': 'Extensive Log Report ($200.00)',
        'barrister_view': 'Appellate Research Brief'
    }
    title = report_type_labels.get(report.get('report_type'), 'Legal Report')
    resolved_sentence = await _derive_export_sentence(case, case_id, user.user_id, report)
    footer_label = _build_export_footer_label(case, title, report.get('generated_at'))

    cover_meta = [
        ['Case Title', case.get('title', 'N/A')],
        ['Defendant', case.get('defendant_name', 'N/A')],
        ['Court / State', f"{case.get('court', 'Court')} — {(case.get('state') or 'Unspecified').upper()}"],
        ['Report', title],
        ['Offence', case.get('offence_type') or (case.get('offence_category') or 'Not recorded').replace('_', ' ').title()],
        ['Sentence', resolved_sentence],
    ]

    story.append(Spacer(1, 6*mm))
    story.append(Paragraph("Appeal Case Manager", styles['ReportSubtitle']))
    story.append(Paragraph(title, styles['ReportTitle']))
    story.append(Paragraph("Criminal Law Appeal Case Management", styles['ReportSubtitle']))
    story.append(Spacer(1, 4*mm))

    cover_table_rows = []
    for idx in range(0, len(cover_meta), 2):
        left = cover_meta[idx]
        right = cover_meta[idx + 1] if idx + 1 < len(cover_meta) else ["", ""]
        cover_table_rows.append([
            Paragraph(f"<b>{left[0]}</b><br/>{left[1]}", styles['CoverMetaValue']),
            Paragraph(f"<b>{right[0]}</b><br/>{right[1]}", styles['CoverMetaValue'])
        ])
    cover_table = Table(cover_table_rows, colWidths=[87*mm, 87*mm])
    cover_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
        ('BOX', (0, 0), (-1, -1), 0.6, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#e2e8f0')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(cover_table)
    story.append(Spacer(1, 6*mm))
    story.append(Paragraph(
        "IMPORTANT DISCLAIMER — NOT LEGAL ADVICE — This application is an educational research tool only and does NOT constitute legal advice. The creator is not a lawyer. All analysis and recommendations must be independently verified by a qualified Australian legal professional. Australian law only. No solicitor-client relationship is created.",
        styles['CoverDisclaimer']
    ))
    story.append(PageBreak())

    # Case Info Table — skip N/A fields
    # Get grounds for PDF header
    pdf_grounds = await db.grounds_of_merit.find({"case_id": case_id}, {"_id": 0}).to_list(100)
    numbered_canvas = NumberedCanvas(footer_label)
    case_data_rows = [
        ['Case Title:', case.get('title', 'N/A')],
        ['Defendant:', case.get('defendant_name', 'N/A')],
    ]
    if case.get('case_number') and case.get('case_number') != 'N/A':
        case_data_rows.append(['Case Number:', case['case_number']])
    if case.get('court') and case.get('court') != 'N/A':
        case_data_rows.append(['Court:', case['court']])
    if case.get('state'):
        case_data_rows.append(['Jurisdiction:', case['state'].upper()])
    if resolved_sentence and resolved_sentence != 'See report analysis':
        case_data_rows.append(['Sentence:', resolved_sentence])
    case_data_rows.append(['Grounds:', f"{len(pdf_grounds)} identified"])
    case_data_rows.append(['Generated:', report.get('generated_at', 'N/A')[:10] if report.get('generated_at') else 'N/A'])
    
    case_table = Table(case_data_rows, colWidths=[40*mm, 130*mm])
    case_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Times-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#475569')),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
    ]))
    story.append(case_table)
    story.append(Spacer(1, 6*mm))
    
    # Grounds of Merit Section
    if grounds:
        story.append(Paragraph("GROUNDS OF MERIT", styles['SectionHeader']))
        story.append(Spacer(1, 2*mm))
        
        ground_type_labels = {
            'procedural_error': 'Procedural Error',
            'fresh_evidence': 'Fresh Evidence',
            'miscarriage_of_justice': 'Miscarriage of Justice',
            'sentencing_error': 'Sentencing Error',
            'judicial_error': 'Judicial Error',
            'ineffective_counsel': 'Ineffective Counsel',
            'prosecution_misconduct': 'Prosecution Misconduct',
            'jury_irregularity': 'Jury Irregularity',
            'constitutional_violation': 'Constitutional Violation',
            'other': 'Other'
        }
        
        for idx, ground in enumerate(grounds, 1):
            ground_type = ground_type_labels.get(ground.get('ground_type'), 'Other')
            strength = ground.get('strength', 'moderate').capitalize()
            
            # Ground header
            story.append(Paragraph(
                f"{idx}. {ground.get('title', 'Unnamed Ground')} [{ground_type}] - Strength: {strength}",
                styles['GroundTitle']
            ))
            
            # Description
            if ground.get('description'):
                story.append(Paragraph(ground.get('description'), styles['ReportBodyText']))
            
            # Legal References (Law Sections)
            if ground.get('law_sections'):
                story.append(Paragraph("<b>Relevant Law Sections:</b>", styles['ReportBodyText']))
                for section in ground.get('law_sections', []):
                    section_text = f"- {_format_law_section(section)}"
                    story.append(Paragraph(section_text, styles['LawSection']))
            
            # Similar Cases
            if ground.get('similar_cases'):
                story.append(Paragraph("<b>Similar Cases:</b>", styles['ReportBodyText']))
                for case_ref in ground.get('similar_cases', []):
                    case_text = f"- {(case_ref.get('case_name') or '').strip()} {(case_ref.get('citation') or '').strip()}".strip()
                    story.append(Paragraph(case_text, styles['LawSection']))
            
            # Supporting Evidence
            if ground.get('supporting_evidence'):
                story.append(Paragraph("<b>Supporting Evidence:</b>", styles['ReportBodyText']))
                for evidence in ground.get('supporting_evidence', []):
                    clean = _format_evidence_item(evidence)
                    if clean:
                        story.append(Paragraph(f"- {clean}", styles['LawSection']))
            
            story.append(Spacer(1, 3*mm))
    
    # Legal Framework Reference — state-specific, NO NSW default
    story.append(Paragraph("LEGAL FRAMEWORK REFERENCE", styles['SectionHeader']))
    story.append(Spacer(1, 1*mm))
    export_state = (case.get('state') or '').lower()
    legal_refs = get_export_legal_refs(export_state)
    for ref in legal_refs:
        story.append(Paragraph(ref, styles['LawSection']))
    
    story.append(Spacer(1, 5*mm))
    
    # Main Analysis Content
    story.append(Paragraph("DETAILED ANALYSIS", styles['SectionHeader']))

    analysis_text = _strip_report_placeholders(report.get('content', {}).get('analysis', 'No analysis available.'))
    render_markdown(analysis_text)

    # DO_NOT_UNDO — Bold red disclaimer in PDF body
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph(
        "IMPORTANT DISCLAIMER — NOT LEGAL ADVICE",
        ParagraphStyle(name='DisclaimerTitle', fontSize=10, fontName='Times-Bold', alignment=TA_CENTER, textColor=colors.HexColor('#dc2626'), spaceAfter=3)
    ))
    story.append(Paragraph(
        "This application is an educational research tool only and does NOT constitute legal advice. It must NOT be relied upon as such. "
        "The creator of this application is not a lawyer. All analysis, findings, reports, and recommendations generated by this tool must be independently verified "
        "by a qualified Australian legal professional before any action is taken. This tool covers Australian law only. "
        "No solicitor-client relationship is created by using this service.",
        ParagraphStyle(name='DisclaimerBody', fontSize=8, fontName='Times-Bold', alignment=TA_CENTER, textColor=colors.HexColor('#dc2626'), leading=11)
    ))
    
    # Build PDF
    try:
        doc.build(story, canvasmaker=numbered_canvas)
    except Exception as e:
        logger.error(f"PDF build failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)[:200]}")
    buffer.seek(0)
    
    # Create filename
    safe_title = "".join(c for c in case.get('title', 'Report')[:30] if c.isalnum() or c in ' -_').strip()
    filename = f"{safe_title}_{report.get('report_type', 'report')}.pdf"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

@router.get("/cases/{case_id}/reports/{report_id}/export-docx")
async def export_report_docx(case_id: str, report_id: str, request: Request):
    """Export a report as DOCX (Microsoft Word) with Grounds of Merit and Legal References"""
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    
    user = await get_current_user(request)
    
    # Get report
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get grounds of merit
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Create DOCX document — canonical style engine (locked 24 Feb 2026).
    from services.print_styles import apply_canonical_docx_styles
    doc = DocxDocument()
    apply_canonical_docx_styles(doc)
    
    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("APPEAL CASE MANAGER")
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(13)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(15, 23, 42)

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_header.add_run("Criminal Law Appeal Case Management")
    sub_run.font.name = 'Times New Roman'
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    case_line = doc.add_paragraph()
    case_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    case_run = case_line.add_run(f"Case: {case.get('title', 'Unknown')}")
    case_run.font.name = 'Times New Roman'
    case_run.font.size = Pt(10)
    case_run.font.color.rgb = RGBColor(71, 85, 105)

    # Report Title
    report_title = REPORT_TYPE_LABELS.get(report.get('report_type'), 'Legal Report')
    resolved_sentence = await _derive_export_sentence(case, case_id, user.user_id, report)
    footer_label = _build_export_footer_label(case, report_title, report.get('generated_at'))

    apply_docx_footer(doc, footer_label)

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title_run = cover_title.add_run(report_title)
    cover_title_run.font.name = 'Times New Roman'
    cover_title_run.font.size = Pt(16)
    cover_title_run.font.bold = True
    cover_title_run.font.color.rgb = RGBColor(15, 23, 42)

    cover_subtitle = doc.add_paragraph()
    cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_subtitle_run = cover_subtitle.add_run("Appeal Case Manager")
    cover_subtitle_run.font.name = 'Times New Roman'
    cover_subtitle_run.font.size = Pt(10)
    cover_subtitle_run.font.bold = True
    cover_subtitle_run.font.color.rgb = RGBColor(29, 78, 216)

    cover_case = doc.add_paragraph()
    cover_case.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_case_run = cover_case.add_run(case.get('title', 'Unknown case'))
    cover_case_run.font.name = 'Times New Roman'
    cover_case_run.font.size = Pt(10)
    cover_case_run.font.color.rgb = RGBColor(71, 85, 105)

    cover_info = [
        ('Case Title', case.get('title', 'N/A')),
        ('Defendant', case.get('defendant_name', 'N/A')),
        ('Court / State', f"{case.get('court', 'Court')} — {(case.get('state') or 'Unspecified').upper()}"),
        ('Report', report_title),
        ('Offence', case.get('offence_type') or (case.get('offence_category') or 'Not recorded').replace('_', ' ').title()),
        ('Sentence', resolved_sentence),
    ]
    cover_table = doc.add_table(rows=len(cover_info), cols=2)
    cover_table.style = 'Table Grid'
    for row_idx, (label, value) in enumerate(cover_info):
        row = cover_table.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 116, 139)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(15, 23, 42)

    cover_disclaimer = doc.add_paragraph()
    cover_disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_disclaimer_run = cover_disclaimer.add_run(
        "IMPORTANT DISCLAIMER — NOT LEGAL ADVICE — This application is an educational research tool only and does NOT constitute legal advice. The creator is not a lawyer. All analysis and recommendations must be independently verified by a qualified Australian legal professional. Australian law only. No solicitor-client relationship is created."
    )
    # DO_NOT_UNDO — Cover page disclaimer in bold red
    cover_disclaimer_run.font.name = 'Times New Roman'
    cover_disclaimer_run.font.size = Pt(9)
    cover_disclaimer_run.font.bold = True
    cover_disclaimer_run.font.color.rgb = RGBColor(220, 38, 38)

    doc.add_page_break()

    title = doc.add_heading(report_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Case Information Table — skip N/A fields
    case_info = [
        ('Case Title:', case.get('title', 'N/A')),
        ('Defendant:', case.get('defendant_name', 'N/A')),
    ]
    if case.get('case_number') and case.get('case_number') != 'N/A':
        case_info.append(('Case Number:', case['case_number']))
    if case.get('court') and case.get('court') != 'N/A':
        case_info.append(('Court:', case['court']))
    if case.get('state'):
        case_info.append(('Jurisdiction:', case['state'].upper()))
    if resolved_sentence and resolved_sentence != 'See report analysis':
        case_info.append(('Sentence:', resolved_sentence))
    case_info.append(('Grounds:', f"{len(grounds)} identified"))
    case_info.append(('Generated:', report.get('generated_at', 'N/A')[:10] if report.get('generated_at') else 'N/A'))
    
    case_table = doc.add_table(rows=len(case_info), cols=2)
    case_table.style = 'Table Grid'
    
    for i, (label, value) in enumerate(case_info):
        row = case_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        row.cells[1].text = str(value)
        row.cells[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Grounds of Merit Section
    if grounds:
        doc.add_heading('GROUNDS OF MERIT', level=1)
        
        ground_type_labels = {
            'procedural_error': 'Procedural Error',
            'fresh_evidence': 'Fresh Evidence',
            'miscarriage_of_justice': 'Miscarriage of Justice',
            'sentencing_error': 'Sentencing Error',
            'judicial_error': 'Judicial Error',
            'ineffective_counsel': 'Ineffective Counsel',
            'prosecution_misconduct': 'Prosecution Misconduct',
            'jury_irregularity': 'Jury Irregularity',
            'constitutional_violation': 'Constitutional Violation',
            'other': 'Other'
        }
        
        for idx, ground in enumerate(grounds, 1):
            ground_type = ground_type_labels.get(ground.get('ground_type'), 'Other')
            strength = ground.get('strength', 'moderate').capitalize()
            
            # Ground heading
            doc.add_heading(
                f"{idx}. {ground.get('title', 'Unnamed Ground')} [{ground_type}] - Strength: {strength}",
                level=2
            )
            
            # Description
            if ground.get('description'):
                doc.add_paragraph(ground.get('description'))
            
            # Legal References
            if ground.get('law_sections'):
                law_para = doc.add_paragraph()
                law_run = law_para.add_run('Relevant Law Sections:')
                law_run.font.bold = True
                law_run.font.color.rgb = RGBColor(30, 64, 175)
                
                for section in ground.get('law_sections', []):
                    section_text = _format_law_section(section)
                    bullet = doc.add_paragraph(section_text, style='List Bullet')
                    for run in bullet.runs:
                        run.font.color.rgb = RGBColor(30, 64, 175)
            
            # Similar Cases
            if ground.get('similar_cases'):
                cases_para = doc.add_paragraph()
                cases_run = cases_para.add_run('Similar Cases:')
                cases_run.font.bold = True
                cases_run.font.color.rgb = RGBColor(30, 58, 138)
                
                for case_ref in ground.get('similar_cases', []):
                    case_text = f"{(case_ref.get('case_name') or '').strip()} {(case_ref.get('citation') or '').strip()}".strip()
                    if case_text:
                        doc.add_paragraph(case_text, style='List Bullet')
            
            # Supporting Evidence
            if ground.get('supporting_evidence'):
                evidence_para = doc.add_paragraph()
                evidence_run = evidence_para.add_run('Supporting Evidence:')
                evidence_run.font.bold = True
                evidence_run.font.color.rgb = RGBColor(5, 150, 105)
                
                for evidence in ground.get('supporting_evidence', []):
                    clean = _format_evidence_item(evidence)
                    if clean:
                        doc.add_paragraph(clean, style='List Bullet')
            
            doc.add_paragraph()
    
    # Legal Framework Reference — state-specific, NO NSW default
    doc.add_heading('LEGAL FRAMEWORK REFERENCE', level=1)
    
    export_state_docx = (case.get('state') or '').lower()
    legal_refs = get_export_legal_refs(export_state_docx)
    
    for ref in legal_refs:
        doc.add_paragraph(ref.lstrip('- '), style='List Bullet')
    
    doc.add_paragraph()
    
    # Detailed Analysis
    doc.add_heading('DETAILED ANALYSIS', level=1)

    def add_table_docx(lines):
        rows = []
        for line in lines:
            parts = [p.strip() for p in line.strip().strip("|").split("|")]
            if not parts:
                continue
            if all(set(p) <= set("-:") for p in parts):
                continue
            rows.append(parts)
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        table.autofit = False
        available_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        col_width = int(available_width / max(1, len(rows[0])))
        for r_idx, row in enumerate(rows):
            for c_idx, value in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.width = col_width
                cell.text = value.replace('**', '')
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                else:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
        doc.add_paragraph()

    def render_markdown_docx(text):
        lines = (text or "").splitlines()
        buffer = []
        table_lines = []

        def flush_paragraph():
            if buffer:
                paragraph = " ".join(buffer).replace('**', '').strip()
                if paragraph:
                    doc.add_paragraph(paragraph)
            buffer.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                flush_paragraph()
                continue
            if stripped.startswith("|") and "|" in stripped:
                flush_paragraph()
                table_lines.append(stripped)
                continue
            if table_lines and not (stripped.startswith("|") and "|" in stripped):
                add_table_docx(table_lines)
                table_lines = []
            if stripped.startswith("## "):
                flush_paragraph()
                doc.add_heading(stripped[3:].strip(), level=1)
                continue
            if stripped.startswith("### "):
                flush_paragraph()
                doc.add_heading(stripped[4:].strip(), level=2)
                continue
            if stripped.startswith("- ") or stripped.startswith("• "):
                flush_paragraph()
                doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
                continue
            buffer.append(stripped)

        flush_paragraph()
        if table_lines:
            add_table_docx(table_lines)

    analysis_text = _strip_report_placeholders(report.get('content', {}).get('analysis', 'No analysis available.'))
    render_markdown_docx(analysis_text)
    
    doc.add_paragraph()
    
    # DO_NOT_UNDO — Bold red disclaimer with white text
    disc_title = doc.add_paragraph()
    disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disc_title.add_run("NOT LEGAL ADVICE")
    disc_run.font.name = 'Times New Roman'
    disc_run.font.size = Pt(10)
    disc_run.font.bold = True
    disc_run.font.color.rgb = RGBColor(220, 38, 38)
    
    disc_body = doc.add_paragraph()
    disc_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_body_run = disc_body.add_run(
        "IMPORTANT DISCLAIMER — This application is an educational research tool only and does NOT constitute legal advice. It must NOT be relied upon "
        "as such. The creator of this application is not a lawyer. All analysis, findings, reports, and recommendations generated "
        "by this tool must be independently verified by a qualified Australian legal professional before any action is taken. "
        "This tool covers Australian law only. No solicitor-client relationship is created by using this service."
    )
    disc_body_run.font.size = Pt(8)
    disc_body_run.font.bold = True
    disc_body_run.font.color.rgb = RGBColor(220, 38, 38)
    disc_body_run.font.name = 'Times New Roman'
    
    # Footer already applied via apply_docx_footer above (line 695)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create filename
    safe_title = "".join(c for c in case.get('title', 'Report')[:30] if c.isalnum() or c in ' -_').strip()
    filename = f"{safe_title}_{report.get('report_type', 'report')}.docx"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

