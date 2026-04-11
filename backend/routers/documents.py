"""
DO NOT UNDO — ENTIRE FILE PROTECTED
Criminal Appeal AI - Documents Router (CRUD + OCR + Text Extraction)
All features in this file are approved. Do not remove, rename, or refactor.
Timeline auto-generation uses fuzzy dedup and year-only dates for unknown dates.
"""
from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Form
from starlette.background import BackgroundTasks
from typing import List
from datetime import datetime, timezone
import base64
import uuid
import re
import logging

from config import db
from auth_utils import get_current_user
from models import Document, DocumentSearchRequest
from services.document_helpers import extract_text_with_ocr
from services.llm_service import call_llm_with_fallback, call_llm_for_json

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["documents"])


# ==========================================================================
# DO NOT UNDO — AUTO-DETECT METADATA FUNCTION
# ==========================================================================
# This function auto-detects state, offence_category, offence_type,
# sentence, court, case_number from uploaded documents via LLM.
# It relies on CaseCreate having Optional[None] defaults (models/__init__.py).
# DO NOT remove this function. DO NOT skip calling it on upload.
# DO NOT hardcode "nsw" or "homicide" anywhere in models or routes.
# This has been broken and re-fixed 15+ times. LEAVE IT ALONE.
# ==========================================================================
async def _background_auto_detect_metadata(case_id: str, user_id: str, content_text: str, filename: str):
    """Background task: auto-detect case metadata from uploaded document via LLM"""
    try:
        case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
        if not case:
            return
        if case.get('offence_type') and case.get('sentence') and case.get('offence_category') and case.get('state'):
            return
        detect_prompt = f"""Analyse this criminal case document and extract metadata.
Return ONLY valid JSON (no markdown):
{{
  "offence_category": "<one of: homicide, assault, sexual_offences, robbery_theft, drug_offences, fraud_dishonesty, firearms_weapons, domestic_violence, public_order, terrorism, driving_offences, cybercrime, arson, perjury, environmental>",
  "offence_type": "<specific offence e.g. Murder, Assault Occasioning ABH>",
  "sentence": "<exact sentence if stated, else null>",
  "state": "<one of: nsw, vic, qld, sa, wa, tas, nt, act, federal>",
  "court": "<court name if found>",
  "case_number": "<case number if found>",
  "defendant_name": "<defendant full name if found>"
}}
Rules: Read the ACTUAL document. Do NOT default to homicide. If a field is unknown, set null. state must be lowercase.

DOCUMENT ({filename}):
{content_text[:20000]}"""
        system_msg = "Extract factual metadata from Australian criminal case documents. Do NOT invent facts. Do NOT default to NSW or homicide if the document does not explicitly state the jurisdiction or offence type. If the case is a Commonwealth/federal prosecution (e.g. Criminal Code Act 1995 (Cth), Crimes Act 1914 (Cth), or prosecuted by the Commonwealth DPP), set state to 'federal'. Use Australian English spelling. Return only valid JSON."
        raw = await call_llm_for_json(system_msg, detect_prompt, f"upload_detect_{case_id}", max_tokens=2000, timeout_seconds=60)
        meta = raw if isinstance(raw, dict) else {}
        if not meta:
            return
        valid_cats = ["homicide","assault","sexual_offences","robbery_theft","drug_offences","fraud_dishonesty","firearms_weapons","domestic_violence","public_order","terrorism","driving_offences","cybercrime","arson","perjury","environmental"]
        valid_sts = ["nsw","vic","qld","sa","wa","tas","nt","act","federal"]
        update_fields = {}
        if meta.get("offence_category") in valid_cats:
            update_fields["offence_category"] = meta["offence_category"]
        if meta.get("offence_type"):
            update_fields["offence_type"] = meta["offence_type"]
        if meta.get("sentence"):
            update_fields["sentence"] = meta["sentence"]
        if meta.get("state") and meta["state"].lower() in valid_sts:
            update_fields["state"] = meta["state"].lower()
        if meta.get("court"):
            update_fields["court"] = meta["court"]
        if meta.get("case_number") and not case.get("case_number"):
            update_fields["case_number"] = meta["case_number"]
        if meta.get("defendant_name") and not case.get("defendant_name"):
            update_fields["defendant_name"] = meta["defendant_name"]
        if update_fields:
            update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
            await db.cases.update_one({"case_id": case_id}, {"$set": update_fields})
            logger.info(f"Background auto-detected metadata for {case_id}: {list(update_fields.keys())}")
    except Exception as e:
        logger.warning(f"Background auto-detect failed for {case_id}: {e}")


async def _background_auto_generate(case_id: str, user_id: str):
    """Background task: auto-generate/UPDATE timeline events + case summary from ALL documents.
    Runs every time a new document is uploaded. Dedup prevents duplicate events."""
    try:
        case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
        if not case:
            return
        documents = await db.documents.find({"case_id": case_id}, {"_id": 0, "file_data": 0}).to_list(500)
        docs_with_text = [d for d in documents if d.get("content_text")]
        if not docs_with_text:
            return
        doc_context = ""
        for doc in docs_with_text:
            doc_context += f"\n--- {doc.get('filename','')} ---\n{doc.get('content_text','')[:6000]}\n"

        # Always regenerate summary to include ALL documents
        try:
            summary = await call_llm_with_fallback(
                "You are a legal analyst. Write a concise factual summary of this criminal case in Australian English. Use third person only. Do NOT invent facts. Do NOT default to NSW if the case is from another jurisdiction. Use forensic language — do not characterise guilt or innocence. 3-5 sentences maximum.",
                f"Summarise this case based on all available documents:\n{doc_context[:15000]}",
                f"summary_{case_id}_{len(docs_with_text)}",
                max_tokens=500,
                timeout_seconds=30,
            )
            if summary:
                await db.cases.update_one({"case_id": case_id}, {"$set": {"summary": summary.strip(), "updated_at": datetime.now(timezone.utc).isoformat()}})
                logger.info(f"Auto-generated/updated summary for {case_id} ({len(docs_with_text)} docs)")
        except Exception as e:
            logger.warning(f"Auto-summary failed for {case_id}: {e}")

        # Always extract timeline events — dedup prevents duplicates
        try:
            tl_prompt = f"""Extract a chronological timeline of events from these case documents.
Return ONLY a JSON array (no markdown). Each event:
{{"date":"YYYY-MM-DD or YYYY if exact date unknown","title":"Brief title","description":"Details","event_type":"incident|arrest|court_hearing|evidence|witness|legal_filing|verdict|appeal|other"}}

CRITICAL RULES:
- Only include events clearly identifiable in the documents
- Do NOT invent dates. If the exact date is unknown, use just the year (e.g. "2018")
- Do NOT use 1 January as a placeholder — use the year only instead
- Do NOT create duplicate events. Each distinct event should appear only once
- Use Australian English spelling (analyse, defence, offence, behaviour, characterisation)

CASE: {case.get('title','')} | DEFENDANT: {case.get('defendant_name','')}
DOCUMENTS:
{doc_context[:25000]}"""
            tl_result = await call_llm_for_json(
                "Extract timeline events from Australian criminal case documents. Do NOT invent dates, names, or events not in the documents. Do NOT default to NSW. Use Australian English spelling. Return only valid JSON array.",
                tl_prompt,
                f"tl_{case_id}_{len(docs_with_text)}",
                max_tokens=4000,
                timeout_seconds=90,
            )
            events = tl_result if isinstance(tl_result, list) else []
            if isinstance(events, list):
                created = 0
                # DO_NOT_UNDO — Timeline fuzzy dedup. Build fingerprints of existing events.
                # Without this, identical events multiply every time documents are processed.
                existing = await db.timeline_events.find({"case_id": case_id}, {"_id": 0, "title": 1}).to_list(200)
                existing_titles = [e.get("title", "").lower().strip() for e in existing]
                from fuzzywuzzy import fuzz
                for evt in events:
                    if not evt.get("title"):
                        continue
                    new_title = evt["title"].lower().strip()
                    # Skip if fuzzy duplicate of existing event
                    is_dup = any(fuzz.token_set_ratio(new_title, et) >= 65 for et in existing_titles)
                    if is_dup:
                        continue
                    event_date = evt.get("date", "")
                    # DO_NOT_UNDO — Fix Jan 1 placeholder dates. Use year only when exact date unknown.
                    if event_date and event_date.endswith("-01-01"):
                        event_date = event_date[:4]  # Keep year only
                    event_doc = {
                        "event_id": f"evt_{uuid.uuid4().hex[:12]}",
                        "case_id": case_id, "user_id": user_id,
                        "title": evt.get("title", ""), "description": evt.get("description", ""),
                        "event_date": event_date, "event_type": evt.get("event_type", "other"),
                        "source": "ai_generated", "created_at": datetime.now(timezone.utc).isoformat()
                    }
                    await db.timeline_events.insert_one(event_doc)
                    existing_titles.append(new_title)
                    created += 1
                logger.info(f"Auto-generated {created} new timeline events for {case_id} (dedup active)")
        except Exception as e:
            logger.warning(f"Auto-timeline failed for {case_id}: {e}")

        # Trigger pipeline extraction for any new documents not yet extracted
        try:
            from services.pipeline.extract import extract_document_artifacts
            # Check which documents don't have extracts yet
            new_extractions = 0
            for doc in docs_with_text:
                existing_extract = await db.document_extracts.find_one(
                    {"case_id": case_id, "document_id": doc["document_id"]}, {"_id": 1}
                )
                if existing_extract:
                    continue
                try:
                    extract = await extract_document_artifacts(case, doc)
                    extract_dict = extract.model_dump()
                    extract_dict["created_at"] = extract_dict["created_at"].isoformat()
                    await db.document_extracts.update_one(
                        {"case_id": case_id, "document_id": doc["document_id"], "user_id": user_id},
                        {"$set": extract_dict},
                        upsert=True,
                    )
                    new_extractions += 1
                except Exception as ex:
                    logger.warning(f"Pipeline extract failed for doc {doc.get('document_id')}: {ex}")
            if new_extractions > 0:
                logger.info(f"Pipeline extracted {new_extractions} new documents for {case_id}")
        except Exception as e:
            logger.warning(f"Pipeline re-extraction failed for {case_id}: {e}")

    except Exception as e:
        logger.warning(f"Background auto-generate failed for {case_id}: {e}")


@router.get("/cases/{case_id}/documents", response_model=List[dict])
async def get_documents(case_id: str, request: Request):
    """Get all documents for a case"""
    user = await get_current_user(request)
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    documents = await db.documents.find({"case_id": case_id}, {"_id": 0, "file_data": 0}).sort("uploaded_at", -1).to_list(500)
    return documents


@router.post("/cases/{case_id}/documents", response_model=dict)
async def upload_document(
    case_id: str, request: Request, background_tasks: BackgroundTasks,
    file: UploadFile = File(...), category: str = Form(...),
    description: str = Form(None), event_date: str = Form(None)
):
    """Upload a document to a case"""
    user = await get_current_user(request)
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    file_content = await file.read()
    file_base64 = base64.b64encode(file_content).decode('utf-8')
    content_text = ""
    file_type = file.content_type or "application/octet-stream"
    filename_lower = file.filename.lower() if file.filename else ""
    try:
        if "text" in file_type or filename_lower.endswith('.txt'):
            content_text = file_content.decode('utf-8', errors='ignore')
        elif "pdf" in file_type or filename_lower.endswith('.pdf'):
            try:
                import io
                from pypdf import PdfReader
                pdf_reader = PdfReader(io.BytesIO(file_content))
                text_parts = []
                for page in pdf_reader.pages[:20]:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
        elif filename_lower.endswith('.docx') or "word" in file_type:
            try:
                import io
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(io.BytesIO(file_content))
                text_parts = [para.text for para in docx_doc.paragraphs if para.text.strip()]
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
    except Exception as e:
        logger.warning(f"Text extraction error: {e}")
    parsed_event_date = None
    if event_date:
        try:
            parsed_event_date = datetime.fromisoformat(event_date.replace('Z', '+00:00')).isoformat()
        except ValueError:
            pass
    doc = Document(
        case_id=case_id, user_id=user.user_id,
        filename=file.filename, file_type=file_type, category=category,
        description=description, content_text=content_text, file_data=file_base64
    )
    doc_dict = doc.model_dump()
    doc_dict["uploaded_at"] = doc_dict["uploaded_at"].isoformat()
    if parsed_event_date:
        doc_dict["event_date"] = parsed_event_date
    await db.documents.insert_one(doc_dict)
    await db.cases.update_one({"case_id": case_id}, {"$set": {"updated_at": datetime.now(timezone.utc).isoformat()}})
    # Schedule background tasks for metadata detection and auto-generation
    if content_text and len(content_text) > 200:
        background_tasks.add_task(_background_auto_detect_metadata, case_id, user.user_id, content_text, file.filename)
        background_tasks.add_task(_background_auto_generate, case_id, user.user_id)
    created_doc = await db.documents.find_one({"document_id": doc.document_id}, {"_id": 0, "file_data": 0})
    return created_doc


@router.get("/cases/{case_id}/documents/{document_id}", response_model=dict)
async def get_document(case_id: str, document_id: str, request: Request):
    """Get a specific document"""
    user = await get_current_user(request)
    doc = await db.documents.find_one(
        {"document_id": document_id, "case_id": case_id, "user_id": user.user_id}, {"_id": 0}
    )
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.delete("/cases/{case_id}/documents/{document_id}")
async def delete_document(case_id: str, document_id: str, request: Request):
    """Delete a document"""
    user = await get_current_user(request)
    result = await db.documents.delete_one({"document_id": document_id, "case_id": case_id, "user_id": user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"message": "Document deleted"}


@router.post("/cases/{case_id}/documents/search", response_model=dict)
async def search_documents(case_id: str, search_request: DocumentSearchRequest, request: Request):
    """Search for text across all documents in a case"""
    user = await get_current_user(request)
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    query = search_request.query.strip()
    if not query:
        raise HTTPException(status_code=400, detail="Search query is required")
    if len(query) < 2:
        raise HTTPException(status_code=400, detail="Search query must be at least 2 characters")
    documents = await db.documents.find(
        {"case_id": case_id, "content_text": {"$exists": True, "$ne": ""}}, {"_id": 0, "file_data": 0}
    ).to_list(500)
    results = []
    total_matches = 0
    flags = 0 if search_request.case_sensitive else re.IGNORECASE
    for doc in documents:
        content = doc.get("content_text", "")
        if not content:
            continue
        matches = []
        try:
            pattern = re.compile(re.escape(query), flags)
            for match in pattern.finditer(content):
                start_pos = match.start()
                end_pos = match.end()
                context_start = max(0, start_pos - 100)
                context_end = min(len(content), end_pos + 100)
                context = content[context_start:context_end]
                if context_start > 0:
                    context = "..." + context
                if context_end < len(content):
                    context = context + "..."
                matches.append({"context": context, "position": start_pos, "matched_text": match.group()})
                if len(matches) >= 10:
                    break
        except re.error:
            continue
        if matches:
            total_matches += len(matches)
            results.append({
                "document_id": doc.get("document_id"), "filename": doc.get("filename"),
                "category": doc.get("category"), "matches": matches, "match_count": len(matches)
            })
    results.sort(key=lambda x: x["match_count"], reverse=True)
    return {"query": query, "total_matches": total_matches, "documents_with_matches": len(results), "total_documents_searched": len(documents), "results": results}


@router.post("/cases/{case_id}/documents/{document_id}/ocr", response_model=dict)
async def ocr_document(case_id: str, document_id: str, request: Request):
    """Extract text from a document using OCR"""
    user = await get_current_user(request)
    doc = await db.documents.find_one({"document_id": document_id, "case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not doc.get('file_data'):
        raise HTTPException(status_code=400, detail="No file data available")
    file_content = base64.b64decode(doc['file_data'])
    content_text, ocr_used = extract_text_with_ocr(file_content, doc.get('filename', ''), doc.get('file_type', ''))
    if not content_text.strip():
        return {"document_id": document_id, "filename": doc.get('filename', ''), "success": False, "ocr_used": ocr_used, "message": "No text could be extracted from this document", "content_length": 0}
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {"content_text": content_text, "ocr_extracted": ocr_used, "text_extracted_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"document_id": document_id, "filename": doc.get('filename', ''), "success": True, "ocr_used": ocr_used, "content_length": len(content_text), "content_preview": content_text[:500] + "..." if len(content_text) > 500 else content_text}


@router.post("/cases/{case_id}/ocr-all", response_model=dict)
async def ocr_all_documents(case_id: str, request: Request):
    """Run OCR on all documents in a case that don't have extracted text"""
    user = await get_current_user(request)
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    documents = await db.documents.find({"case_id": case_id}, {"_id": 0}).to_list(500)
    results = []
    successful_ocr = 0
    for doc in documents:
        existing_text = doc.get('content_text', '')
        if len(existing_text.strip()) > 100:
            results.append({"document_id": doc['document_id'], "filename": doc.get('filename'), "status": "skipped", "reason": "Already has extracted text"})
            continue
        if not doc.get('file_data'):
            results.append({"document_id": doc['document_id'], "filename": doc.get('filename'), "status": "skipped", "reason": "No file data"})
            continue
        file_content = base64.b64decode(doc['file_data'])
        content_text, ocr_used = extract_text_with_ocr(file_content, doc.get('filename', ''), doc.get('file_type', ''))
        if content_text.strip():
            await db.documents.update_one(
                {"document_id": doc['document_id']},
                {"$set": {"content_text": content_text, "ocr_extracted": ocr_used, "text_extracted_at": datetime.now(timezone.utc).isoformat()}}
            )
            successful_ocr += 1
            results.append({"document_id": doc['document_id'], "filename": doc.get('filename'), "status": "success", "ocr_used": ocr_used, "content_length": len(content_text)})
        else:
            results.append({"document_id": doc['document_id'], "filename": doc.get('filename'), "status": "failed", "reason": "No text could be extracted"})
    return {"total_documents": len(documents), "successful_extractions": successful_ocr, "results": results}


@router.post("/cases/{case_id}/documents/{document_id}/extract-text", response_model=dict)
async def extract_document_text(case_id: str, document_id: str, request: Request):
    """Re-extract text from a document"""
    user = await get_current_user(request)
    doc = await db.documents.find_one({"document_id": document_id, "case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if not doc.get('file_data'):
        raise HTTPException(status_code=400, detail="No file data available")
    file_content = base64.b64decode(doc['file_data'])
    filename_lower = doc.get('filename', '').lower()
    file_type = doc.get('file_type', '')
    content_text = ""
    try:
        if "text" in file_type or filename_lower.endswith('.txt'):
            content_text = file_content.decode('utf-8', errors='ignore')
        elif "pdf" in file_type or filename_lower.endswith('.pdf'):
            try:
                import io
                from pypdf import PdfReader
                pdf_reader = PdfReader(io.BytesIO(file_content))
                text_parts = [page.extract_text() for page in pdf_reader.pages[:30] if page.extract_text()]
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PDF extraction failed: {e}")
        elif filename_lower.endswith('.docx') or "word" in file_type:
            try:
                import io
                from docx import Document as DocxDocument
                docx_doc = DocxDocument(io.BytesIO(file_content))
                text_parts = [para.text for para in docx_doc.paragraphs if para.text.strip()]
                content_text = "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"DOCX extraction failed: {e}")
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Text extraction failed: {str(e)}")
    await db.documents.update_one(
        {"document_id": document_id},
        {"$set": {"content_text": content_text, "text_extracted_at": datetime.now(timezone.utc).isoformat()}}
    )
    return {"document_id": document_id, "filename": doc.get('filename'), "content_length": len(content_text), "content_preview": content_text[:500] + "..." if len(content_text) > 500 else content_text, "success": bool(content_text)}


@router.post("/cases/{case_id}/extract-all-text", response_model=dict)
async def extract_all_documents_text(case_id: str, request: Request):
    """Extract text from all documents in a case"""
    user = await get_current_user(request)
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    documents = await db.documents.find({"case_id": case_id}, {"_id": 0}).to_list(500)
    results = []
    for doc in documents:
        if not doc.get('file_data'):
            results.append({"document_id": doc['document_id'], "success": False, "error": "No file data"})
            continue
        file_content = base64.b64decode(doc['file_data'])
        filename_lower = doc.get('filename', '').lower()
        file_type = doc.get('file_type', '')
        content_text = ""
        error = None
        try:
            if "text" in file_type or filename_lower.endswith('.txt'):
                content_text = file_content.decode('utf-8', errors='ignore')
            elif "pdf" in file_type or filename_lower.endswith('.pdf'):
                try:
                    import io
                    from pypdf import PdfReader
                    pdf_reader = PdfReader(io.BytesIO(file_content))
                    text_parts = [page.extract_text() for page in pdf_reader.pages[:30] if page.extract_text()]
                    content_text = "\n".join(text_parts)
                except Exception as e:
                    error = str(e)
            elif filename_lower.endswith('.docx') or "word" in file_type:
                try:
                    import io
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(io.BytesIO(file_content))
                    text_parts = [para.text for para in docx_doc.paragraphs if para.text.strip()]
                    content_text = "\n".join(text_parts)
                except Exception as e:
                    error = str(e)
        except Exception as e:
            error = str(e)
        if content_text:
            await db.documents.update_one(
                {"document_id": doc['document_id']},
                {"$set": {"content_text": content_text, "text_extracted_at": datetime.now(timezone.utc).isoformat()}}
            )
        results.append({"document_id": doc['document_id'], "filename": doc.get('filename'), "success": bool(content_text), "content_length": len(content_text) if content_text else 0, "error": error})
    successful = sum(1 for r in results if r['success'])
    # AI auto-detect case metadata from extracted text
    detected_metadata = {}
    if successful > 0:
        try:
            combined_text = ""
            for doc in documents:
                ct = doc.get("content_text") or ""
                if not ct:
                    for r in results:
                        if r.get("document_id") == doc["document_id"] and r.get("success"):
                            fresh = await db.documents.find_one({"document_id": doc["document_id"]}, {"_id": 0, "content_text": 1})
                            ct = (fresh or {}).get("content_text", "")
                if ct:
                    combined_text += f"\n--- DOCUMENT: {doc.get('filename','')} ---\n{ct[:6000]}\n"
            if len(combined_text) > 200:
                detect_prompt = f"""Analyse the following case documents and extract these metadata fields.
Return ONLY valid JSON with these exact keys (no markdown, no explanation):
{{
  "offence_category": "<one of: homicide, assault, sexual_offences, robbery_theft, drug_offences, fraud_dishonesty, firearms_weapons, domestic_violence, public_order, terrorism, driving_offences>",
  "offence_type": "<specific offence>",
  "sentence": "<exact sentence imposed>",
  "state": "<one of: nsw, vic, qld, sa, wa, tas, nt, act>",
  "court": "<court name if found>",
  "case_number": "<case number if found>",
  "defendant_name": "<defendant full name if found>"
}}
Rules: offence_category MUST be from the provided list. Read ACTUAL content. If unknown, null. state lowercase.

DOCUMENTS:
{combined_text[:30000]}"""
                system_msg = "You are a legal document analyst. Extract factual metadata from Australian criminal case documents. Do NOT invent facts. Do NOT default to NSW or homicide if the document does not explicitly state the jurisdiction or offence type. Use Australian English spelling. Return only valid JSON."
                meta = await call_llm_for_json(system_msg, detect_prompt, f"detect_{case_id}", max_tokens=2000, timeout_seconds=60)
                if not isinstance(meta, dict):
                    meta = {}
                update_fields = {}
                valid_categories = ["homicide","assault","sexual_offences","robbery_theft","drug_offences","fraud_dishonesty","firearms_weapons","domestic_violence","public_order","terrorism","driving_offences"]
                valid_states = ["nsw","vic","qld","sa","wa","tas","nt","act"]
                if meta.get("offence_category") in valid_categories:
                    update_fields["offence_category"] = meta["offence_category"]
                if meta.get("offence_type"):
                    update_fields["offence_type"] = meta["offence_type"]
                if meta.get("sentence"):
                    update_fields["sentence"] = meta["sentence"]
                if meta.get("state") and meta["state"].lower() in valid_states:
                    update_fields["state"] = meta["state"].lower()
                if meta.get("court"):
                    update_fields["court"] = meta["court"]
                if meta.get("case_number"):
                    update_fields["case_number"] = meta["case_number"]
                if meta.get("defendant_name") and not case.get("defendant_name"):
                    update_fields["defendant_name"] = meta["defendant_name"]
                if update_fields:
                    update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
                    await db.cases.update_one({"case_id": case_id}, {"$set": update_fields})
                    detected_metadata = update_fields
                    logger.info(f"Auto-detected case metadata for {case_id}: {update_fields}")
        except Exception as e:
            logger.warning(f"Auto-detect metadata failed for {case_id}: {e}")
    return {"total_documents": len(documents), "successful_extractions": successful, "results": results, "detected_metadata": detected_metadata}
