# DO NOT UNDO — server.py. All endpoints in this file are approved and must be preserved.
# ===========================================================================
# Criminal Appeal AI - Main Server (Refactored)
# ADDITIVE HARDENING PATCH
# ===========================================================================
# After extraction, this file contains:
#   - App factory, CORS, health check
#   - Report quality/dedup utilities
#   - AI report generation (analyze_case_with_ai)
#   - Barrister view generation
#   - Report CRUD endpoints
#   - PDF/DOCX export endpoints
#   - Router inclusion and startup/shutdown
# ===========================================================================

import os
import re
import uuid
import asyncio
from datetime import datetime, timezone, timedelta
from typing import List

from fastapi import FastAPI, APIRouter, HTTPException, Request, Response
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
from starlette.responses import StreamingResponse

from config import db, logger, client, get_admin_emails
from auth_utils import get_current_user
from models import (
    ReportRequest, FEATURE_PRICES, feature_type_variants,
    ReportMetadata,
)
from services.llm_service import call_llm_with_fallback, call_llm_for_json, call_llm_structured
from services.offence_helpers import get_offence_context, get_offence_system_prompt
from services.document_helpers import build_document_context
from offence_framework import OFFENCE_CATEGORIES, AUSTRALIAN_STATES

from routers.cases import router as cases_router
from routers.auth import router as auth_router
from routers.password_reset import router as password_reset_router
from routers.admin import router as admin_router
from routers.utilities import router as utilities_router
from routers.analytics import router as analytics_router
from routers.statistics import router as statistics_router
from routers.compare import router as compare_router
from routers.contradictions import router as contradictions_router
from routers.export import router as export_router
from routers.export import translate_router
from routers.collaboration import router as collaboration_router
from routers.documents import router as documents_router
from routers.timeline import router as timeline_router
from routers.deadlines import router as deadlines_router
from routers.notes import router as notes_router
from routers.grounds import router as grounds_router
from routers.payments import router as payments_router
from routers.resources import router as resources_router
from routers.analysis import router as analysis_router
from routers.pipeline import router as pipeline_router
from routers.pipeline_staged import router as pipeline_staged_router
from routers.caselaw import router as caselaw_router
from routers.pipeline_staged import (
    _ensure_document_extracts as _staged_ensure_extracts,
    _refresh_case_extract as _staged_refresh_case,
    _classify_issues as _staged_classify,
    _verify_top_issues as _staged_verify,
    _sync_pipeline_projection_to_grounds as _staged_sync_grounds,
)
from services.pipeline import (
    extract_document_artifacts,
    classify_case_issues,
    verify_issue,
)
from services.pipeline_models import CaseExtract

# ── Admin helpers ──
def is_admin_user(email: str) -> bool:
    admin_emails = get_admin_emails()
    normalized = (email or "").strip().lower()
    allowed = {(e or "").strip().lower() for e in admin_emails}
    return normalized in allowed


# ── FastAPI app ──
app = FastAPI(title="Criminal Appeal AI", version="2.0.0")
api_router = APIRouter(prefix="/api")


# ── Security Headers Middleware ──
class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    """Adds security headers and rate-limits auth endpoints."""
    _auth_attempts: dict = {}  # IP -> [(timestamp, ...)]

    async def dispatch(self, request: Request, call_next):
        # Rate limit auth endpoints: 10 requests per minute per IP
        path = request.url.path
        if request.method == "POST" and any(p in path for p in ["/auth/login", "/auth/register", "/auth/forgot-password"]):
            ip = request.client.host if request.client else "unknown"
            now = datetime.now(timezone.utc)
            cutoff = now - timedelta(minutes=1)
            attempts = self._auth_attempts.get(ip, [])
            attempts = [t for t in attempts if t > cutoff]
            if len(attempts) >= 10:
                return Response(
                    content='{"detail":"Too many attempts. Please wait 1 minute."}',
                    status_code=429,
                    media_type="application/json"
                )
            attempts.append(now)
            self._auth_attempts[ip] = attempts

        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
        if request.url.scheme == "https":
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response

app.add_middleware(SecurityHeadersMiddleware)


@app.get("/health")
@app.get("/api/health")
async def health_check():
    try:
        await db.command("ping")
        db_status = "connected"
    except Exception:
        db_status = "disconnected"
    return {"status": "healthy", "database": db_status, "timestamp": datetime.now(timezone.utc).isoformat()}


@app.get("/api/ready")
async def readiness_check():
    """Readiness probe — returns 503 if MongoDB is not connected."""
    try:
        await db.command("ping")
        return {"ready": True, "timestamp": datetime.now(timezone.utc).isoformat()}
    except Exception:
        return Response(
            content='{"ready":false,"reason":"database_unavailable"}',
            status_code=503,
            media_type="application/json"
        )


@app.get("/api/health/deep")
async def deep_health_check():
    """Deep health check — verifies MongoDB, LLM key, and email service."""
    from config import EMERGENT_LLM_KEY
    checks = {}
    # MongoDB
    try:
        await db.command("ping")
        user_count = await db.users.count_documents({})
        checks["mongodb"] = {"status": "ok", "users": user_count}
    except Exception as e:
        checks["mongodb"] = {"status": "error", "detail": str(e)}
    # LLM Key
    checks["llm_key"] = {"status": "ok" if EMERGENT_LLM_KEY else "missing"}
    # Resend
    resend_key = os.environ.get("RESEND_API_KEY", "")
    checks["email"] = {"status": "ok" if resend_key else "not_configured"}
    all_ok = all(c.get("status") == "ok" for c in checks.values())
    return {"healthy": all_ok, "checks": checks, "timestamp": datetime.now(timezone.utc).isoformat()}


# ============ AI ANALYSIS & REPORTS ============


def _normalise_text(value: str) -> str:
    cleaned = re.sub(r"[^a-z0-9\s]", " ", (value or "").lower())
    return re.sub(r"\s+", " ", cleaned).strip()


def _token_set(value: str) -> set:
    return {token for token in _normalise_text(value).split(" ") if len(token) > 3}


def _jaccard_similarity(set_a: set, set_b: set) -> float:
    if not set_a or not set_b:
        return 0.0
    intersection = len(set_a & set_b)
    union = len(set_a | set_b)
    return 0.0 if union == 0 else intersection / union


def _build_anchor_terms(case: dict, documents: list, timeline: list, grounds: list) -> set:
    terms = set()

    def add_terms(value):
        if not value:
            return
        for token in re.split(r"[^a-z0-9]+", str(value).lower()):
            if len(token) > 3:
                terms.add(token)

    for field in ["title", "defendant_name", "case_number", "court", "judge", "sentence", "state", "offence_type", "offence_category"]:
        add_terms(case.get(field))

    for doc in documents or []:
        add_terms(doc.get("filename"))
        add_terms(doc.get("category"))
        add_terms(doc.get("document_type"))

    for event in timeline or []:
        add_terms(event.get("title"))
        add_terms(event.get("event_type"))
        add_terms(event.get("event_date"))

    for ground in grounds or []:
        add_terms(ground.get("title"))
        add_terms(ground.get("ground_type"))

    return terms


def _paragraph_quality_score(paragraph: str, anchor_terms: set) -> float:
    if not paragraph:
        return 0.0
    score = 0.0
    if re.search(r"\b\d{2,}\b", paragraph):
        score += 1.1
    if re.search(r"\b(s\.|section)\s*\d+", paragraph, re.I):
        score += 1.2
    if re.search(r"\bAct\s+\d{4}\b", paragraph, re.I):
        score += 1.0
    if re.search(r"\bR\s+v\b", paragraph) or re.search(r"\bNSWCCA\b", paragraph):
        score += 0.9

    word_set = _token_set(paragraph)
    if anchor_terms and word_set:
        anchor_hits = sum(1 for word in word_set if word in anchor_terms)
        score += min(2.0, anchor_hits * 0.25)

    score += min(1.4, len(paragraph) / 1200)
    return score


def _split_report_sections(text: str) -> list:
    parts = re.split(r"(^##\s+\d+\.\s+.+$)", text, flags=re.M)
    if len(parts) <= 1:
        return [("", text.strip())]

    sections = []
    lead = parts[0].strip()
    if lead:
        sections.append(("", lead))

    for index in range(1, len(parts), 2):
        heading = parts[index].strip()
        content = parts[index + 1] if index + 1 < len(parts) else ""
        sections.append((heading, content.strip()))

    return sections


def _dedupe_report_content(text: str, report_type: str, anchor_terms: set) -> str:
    """DO_NOT_UNDO — Dedup thresholds are FINAL. 0.97 for multi-pass, 0.90 for single-pass.
    Previous 0.82 threshold stripped ~50% of valid report content. NEVER lower below 0.90."""
    if not text:
        return text

    # DO_NOT_UNDO — Multi-pass reports must NOT be deduped aggressively
    if report_type in ("full_detailed", "extensive_log"):
        similarity_threshold = 0.97
    else:
        similarity_threshold = 0.90

    sections = _split_report_sections(text)
    cleaned_sections = []

    for heading, content in sections:
        paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", content) if p.strip()]
        kept = []
        seen_sets = []
        seen_texts = []

        for paragraph in paragraphs:
            if len(paragraph) < 40:
                kept.append(paragraph)
                continue

            paragraph_set = _token_set(paragraph)
            is_duplicate = False
            for existing_set, existing_text in zip(seen_sets, seen_texts):
                # Only strip if one paragraph is a substring of another (exact containment)
                if paragraph in existing_text or existing_text in paragraph:
                    is_duplicate = True
                    break
                if _jaccard_similarity(paragraph_set, existing_set) > similarity_threshold:
                    is_duplicate = True
                    break

            if is_duplicate:
                continue

            kept.append(paragraph)
            seen_sets.append(paragraph_set)
            seen_texts.append(paragraph)

        cleaned_sections.append((heading, "\n\n".join(kept).strip()))

    rebuilt = []
    for heading, content in cleaned_sections:
        if heading:
            rebuilt.append(heading)
        if content:
            rebuilt.append(content)

    return "\n\n".join(rebuilt).strip()


def _strip_report_placeholders(text: str) -> str:
    if not text:
        return text
    cleaned_lines = []
    for line in text.splitlines():
        if re.search(r"\[Your Name\]|\[Your Legal Organisation/Team\]", line, re.I):
            continue
        if re.search(r"Best regards|Prepared by|Prepared for|This report was generated by|Criminal Appeal AI|Justitia AI", line, re.I):
            continue
        if re.search(r"\bDO NOT UNDO\.?\b", line, re.I):
            continue
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    # Strip \1 artifacts
    cleaned = cleaned.replace("\\1", "")
    cleaned = cleaned.replace("\x01", "")
    # Strip prompt instruction text from section headings
    cleaned = re.sub(r'\s*—\s*keep ALL[^\n]*', '', cleaned)
    cleaned = re.sub(r'\s*—\s*DETAILED PATHWAY ANALYSIS[^\n]*', '', cleaned)
    cleaned = re.sub(r'\s*\(\d+\+?\s*words[^)]*\)', '', cleaned)  # e.g. "(900+ words per ground, flowing paragraphs)"
    cleaned = re.sub(r'\s*\(\d+\+?\s*CASES[^)]*\)', '', cleaned)  # e.g. "(12+ CASES with full citations)"
    cleaned = re.sub(r'(GROUNDS OF MERIT)\s*—\s*DEEP ANALYSIS', r'\1', cleaned)
    # Sanitise "we/us/our" language — convert to third-person educational tone
    we_us_replacements = [
        (r'\bWe are arguing\b', 'The applicant argues'),
        (r'\bwe are arguing\b', 'the applicant argues'),
        (r'\bWe are aiming\b', 'The appeal aims'),
        (r'\bwe are aiming\b', 'the appeal aims'),
        (r'\bWe are filing\b', 'The legal professional is filing'),
        (r'\bwe are filing\b', 'the legal professional is filing'),
        (r'\bWe are taking\b', 'The legal professional is taking'),
        (r'\bwe are taking\b', 'the legal professional is taking'),
        (r'\bWe succeed\b', 'The appeal succeeds'),
        (r'\bwe succeed\b', 'the appeal succeeds'),
        (r'\bwe will gather\b', 'the legal professional will gather'),
        (r'\bWe will gather\b', 'The legal professional will gather'),
        (r'\bwe will craft\b', 'the legal professional will craft'),
        (r'\bWe will craft\b', 'The legal professional will craft'),
        (r'\bwe will file\b', 'the legal professional will file'),
        (r'\bWe will file\b', 'The legal professional will file'),
        (r'\bwe will prepare\b', 'the legal professional will prepare'),
        (r'\bWe will prepare\b', 'The legal professional will prepare'),
        (r'\bwe will submit\b', 'the legal professional will submit'),
        (r'\bWe will submit\b', 'The legal professional will submit'),
        (r'\bwe will seek\b', 'the applicant will seek'),
        (r'\bWe will seek\b', 'The applicant will seek'),
        (r'\bwe will argue\b', 'the applicant will argue'),
        (r'\bWe will argue\b', 'The applicant will argue'),
        (r'\bwe will demonstrate\b', 'the appeal will demonstrate'),
        (r'\bWe will demonstrate\b', 'The appeal will demonstrate'),
        (r'\bwe will show\b', 'the appeal will show'),
        (r'\bWe will show\b', 'The appeal will show'),
        (r'\bcontact with us\b', 'contact with the legal professional'),
        (r'\bContact us\b', 'Contact the legal professional'),
        (r'\bcontact us\b', 'contact the legal professional'),
        (r'\bour submissions\b', 'the submissions'),
        (r'\bOur submissions\b', 'The submissions'),
        (r'\bour claims\b', "the applicant's claims"),
        (r'\bOur claims\b', "The applicant's claims"),
        (r'\bour arguments\b', "the applicant's arguments"),
        (r'\bOur arguments\b', "The applicant's arguments"),
        (r'\bour position\b', "the applicant's position"),
        (r'\bOur position\b', "The applicant's position"),
        (r'\bour case\b', "the applicant's case"),
        (r'\bOur case\b', "The applicant's case"),
        (r'\bour strategy\b', 'the legal strategy'),
        (r'\bOur strategy\b', 'The legal strategy'),
        (r'\bour analysis\b', 'this analysis'),
        (r'\bOur analysis\b', 'This analysis'),
        (r'\bon our behalf\b', 'on behalf of the applicant'),
        (r'\bback our\b', "support the applicant's"),
        (r'\bensuring our\b', 'ensuring the'),
        (r', we are\b', ', the legal professional is'),
        (r', we will\b', ', the legal professional will'),
        (r', we have\b', ', the legal professional has'),
        (r'\bWe have identified\b', 'This analysis has identified'),
        (r'\bwe have identified\b', 'this analysis has identified'),
        (r'\bWe have reviewed\b', 'This analysis has reviewed'),
        (r'\bwe have reviewed\b', 'this analysis has reviewed'),
        (r'\bWe have analysed\b', 'This analysis has examined'),
        (r'\bwe have analysed\b', 'this analysis has examined'),
        (r'\bWe have analyzed\b', 'This analysis has examined'),
        (r'\bwe have analyzed\b', 'this analysis has examined'),
    ]
    for pattern, replacement in we_us_replacements:
        cleaned = re.sub(pattern, replacement, cleaned)
    
    # Catch-all: replace remaining "we " patterns at sentence boundaries
    cleaned = re.sub(r'\bwe facilitate\b', 'this analysis facilitates', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe recommend\b', 'it is recommended', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe note\b', 'it is noted', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe observe\b', 'it is observed', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe consider\b', 'it is considered', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe conclude\b', 'it is concluded', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe suggest\b', 'it is suggested', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe assess\b', 'this analysis assesses', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe examine\b', 'this analysis examines', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe highlight\b', 'this analysis highlights', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe identify\b', 'this analysis identifies', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe believe\b', 'it is assessed', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe find\b', 'this analysis finds', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe can see\b', 'it can be seen', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe can\b', 'the applicant can', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe must\b', 'the legal professional must', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe should\b', 'the legal professional should', cleaned, flags=re.I)
    cleaned = re.sub(r'\bwe need\b', 'the legal professional needs', cleaned, flags=re.I)
    cleaned = re.sub(r', we \b', ', the analysis ', cleaned)
    cleaned = re.sub(r'\. We \b', '. This analysis ', cleaned)
    # Catch remaining "our" possessives
    cleaned = re.sub(r'\bour client\b', 'the applicant', cleaned, flags=re.I)
    cleaned = re.sub(r'\bour focus\b', 'the focus', cleaned, flags=re.I)
    cleaned = re.sub(r'\bour team\b', 'the legal professional', cleaned, flags=re.I)
    cleaned = re.sub(r'\bour review\b', 'this review', cleaned, flags=re.I)
    cleaned = re.sub(r'\bour assessment\b', 'this assessment', cleaned, flags=re.I)
    cleaned = re.sub(r'\bour examination\b', 'this examination', cleaned, flags=re.I)
    
    # Strip "you/your" language — third-person only
    you_your_replacements = [
        (r'\byour opportunity\b', 'the opportunity'),
        (r'\bYour opportunity\b', 'The opportunity'),
        (r'\byour conviction\b', 'the conviction'),
        (r'\bYour conviction\b', 'The conviction'),
        (r'\bgiven to you\b', 'imposed'),
        (r'\bGiven to you\b', 'Imposed'),
        (r'\byour sentence\b', 'the sentence'),
        (r'\bYour sentence\b', 'The sentence'),
        (r'\byour appeal\b', 'the appeal'),
        (r'\bYour appeal\b', 'The appeal'),
        (r'\byour case\b', 'the case'),
        (r'\bYour case\b', 'The case'),
        (r'\byour trial\b', 'the trial'),
        (r'\bYour trial\b', 'The trial'),
        (r'\byour lawyer\b', 'the legal representative'),
        (r'\bYour lawyer\b', 'The legal representative'),
        (r'\byour barrister\b', 'the barrister'),
        (r'\bYour barrister\b', 'The barrister'),
        (r'\byour solicitor\b', 'the solicitor'),
        (r'\bYour solicitor\b', 'The solicitor'),
        (r'\byour legal team\b', 'the legal team'),
        (r'\bYour legal team\b', 'The legal team'),
        (r'\byour legal representative\b', 'the legal representative'),
        (r'\bYour legal representative\b', 'The legal representative'),
        (r'\byour defence\b', 'the defence'),
        (r'\bYour defence\b', 'The defence'),
        (r'\byour rights\b', 'the rights of the applicant'),
        (r'\bYour rights\b', 'The rights of the applicant'),
        (r'\byour circumstances\b', 'the circumstances'),
        (r'\bYour circumstances\b', 'The circumstances'),
        (r'\byour situation\b', 'the situation'),
        (r'\bYour situation\b', 'The situation'),
        (r'\byour matter\b', 'this matter'),
        (r'\bYour matter\b', 'This matter'),
        (r'\byour prospects\b', 'the prospects'),
        (r'\bYour prospects\b', 'The prospects'),
        (r'\byour grounds\b', 'the grounds'),
        (r'\bYour grounds\b', 'The grounds'),
        (r'\byour documents\b', 'the documents'),
        (r'\bYour documents\b', 'The documents'),
        (r'\byou may\b', 'the applicant may'),
        (r'\bYou may\b', 'The applicant may'),
        (r'\byou can\b', 'the applicant can'),
        (r'\bYou can\b', 'The applicant can'),
        (r'\byou should\b', 'the applicant should'),
        (r'\bYou should\b', 'The applicant should'),
        (r'\byou must\b', 'the applicant must'),
        (r'\bYou must\b', 'The applicant must'),
        (r'\byou will\b', 'the applicant will'),
        (r'\bYou will\b', 'The applicant will'),
        (r'\byou need\b', 'the applicant needs'),
        (r'\bYou need\b', 'The applicant needs'),
        (r'\byou have\b', 'the applicant has'),
        (r'\bYou have\b', 'The applicant has'),
        (r'\byou are\b', 'the applicant is'),
        (r'\bYou are\b', 'The applicant is'),
        (r'\byou were\b', 'the applicant was'),
        (r'\bYou were\b', 'The applicant was'),
        (r'\bfor you\b', 'for the applicant'),
        (r'\bFor you\b', 'For the applicant'),
        (r'\bto you\b', 'to the applicant'),
        (r'\bTo you\b', 'To the applicant'),
        (r'\bagainst you\b', 'against the applicant'),
        (r'\bAgainst you\b', 'Against the applicant'),
        (r'\bif you\b', 'if the applicant'),
        (r'\bIf you\b', 'If the applicant'),
        (r'\bwhen you\b', 'when the applicant'),
        (r'\bWhen you\b', 'When the applicant'),
    ]
    for pattern, replacement in you_your_replacements:
        cleaned = re.sub(pattern, replacement, cleaned)
    # Catch remaining "your" as possessive — broad catch-all
    cleaned = re.sub(r'\byour\b', "the applicant's", cleaned)
    cleaned = re.sub(r'\bYour\b', "The applicant's", cleaned)
    
    # Strip placeholder/future-tense filler text
    cleaned = re.sub(r'(?:The |This )?(?:comparative sentencing |sentencing )?table (?:below )?will (?:reference|provide|include|contain|show|list|detail|present|outline|cover)\b', 'The table references', cleaned, flags=re.I)
    cleaned = re.sub(r'(?:Details|Information|Data|Analysis|Content),?\s*(?:including [^,]+,?\s*)?will be provided\b', 'Details are provided', cleaned, flags=re.I)
    
    # Fix broken markdown links: <Text> [Text](url) → just [Text](url)
    cleaned = re.sub(r'<([^>]+)>\s*\[([^\]]+)\]\(([^)]+)\)', r'[\2](\3)', cleaned)
    # Fix raw angle-bracket links: <Text> without markdown
    cleaned = re.sub(r'<(Search [^>]+)>', r'\1', cleaned)
    
    return cleaned


def _sanitise_suspect_authorities(text: str) -> str:
    """Post-processing gate that flags or strips suspect LLM-fabricated case citations.
    
    DO NOT UNDO — Prevents fabricated authorities from landing in saved/exported reports.
    Catches:
    - Placeholder citations: [Surname] [Year], [Year] NSWCCA [Number]
    - Obviously templated citations: "R v [Name]", "citation verification needed"
    - Suspiciously round paragraph numbers: at [100], at [200]
    - "section not provided" placeholders
    
    Does NOT strip citations that look real (e.g. "R v Smith [2015] NSWCCA 123 at [45]").
    Instead of removing suspect citations entirely, wraps them in a caveat so the user
    knows to verify, preserving the analytical context.
    """
    if not text:
        return text
    
    # Strip obvious placeholder citations
    text = re.sub(
        r'\[Surname\]|\[Name\]|\[Year\]|\[Number\]|\[Citation\]|\[Court\]',
        '[citation verification required]',
        text,
        flags=re.I
    )
    # Strip "citation verification needed" or "citation pending" as standalone lines
    text = re.sub(
        r'\n[^\n]*(?:citation verification needed|citation pending|citation to be confirmed|citation unavailable)[^\n]*\n',
        '\n',
        text,
        flags=re.I
    )
    # Strip "section not provided" placeholders
    text = re.sub(
        r'(?:section|s\.?)\s+not\s+provided',
        '[section number to be confirmed]',
        text,
        flags=re.I
    )
    # Flag obviously templated case references: "R v [anything in brackets]"
    text = re.sub(
        r'R\s+v\s+\[([A-Z][a-z]+)\]\s*\[(\d{4})\]',
        r'R v \1 [\2] [citation unverified]',
        text
    )
    # Flag suspiciously perfect paragraph references like "at [100]", "at [200]", "at [300]"
    text = re.sub(
        r'at\s+\[([1-9]00)\]',
        r'at [\1] [paragraph number unverified]',
        text
    )
    # Strip AustLII links that are clearly templated (contain placeholder segments)
    text = re.sub(
        r'https?://www\.austlii\.edu\.au/[^\s\)]*\[(?:Year|Number|Citation)\][^\s\)]*',
        '[AustLII link to be verified]',
        text,
        flags=re.I
    )
    
    return text


# ============================================================================
# PIPELINE HELPERS FOR REPORT GENERATION
# ============================================================================

async def _ensure_document_extracts_for_case(case: dict, documents: list) -> int:
    """Ensure every uploaded document has a staged extraction record."""
    created = 0
    for document in documents:
        existing = await db.document_extracts.find_one(
            {"case_id": case["case_id"], "document_id": document["document_id"], "user_id": case["user_id"]},
            {"_id": 0}
        )
        if existing:
            continue
        extract = await extract_document_artifacts(case, document)
        extract_dict = extract.model_dump()
        extract_dict["created_at"] = extract_dict["created_at"].isoformat()
        await db.document_extracts.update_one(
            {"case_id": case["case_id"], "document_id": document["document_id"], "user_id": case["user_id"]},
            {"$set": extract_dict},
            upsert=True,
        )
        created += 1
    return created


async def _refresh_case_extract_for_case(case: dict) -> dict:
    """Merge all document extracts into a case-level extract."""
    extracts = await db.document_extracts.find(
        {"case_id": case["case_id"], "user_id": case["user_id"]}, {"_id": 0}
    ).to_list(500)
    merged_facts, merged_events, merged_findings, extract_ids = [], [], [], []
    for ext in extracts:
        extract_ids.append(ext.get("extract_id"))
        merged_facts.extend(ext.get("facts", []))
        merged_events.extend(ext.get("events", []))
        merged_findings.extend(ext.get("findings", []))
    case_extract = CaseExtract(
        case_id=case["case_id"],
        user_id=case["user_id"],
        metadata={
            "state": case.get("state"),
            "offence_category": case.get("offence_category"),
            "offence_type": case.get("offence_type"),
            "sentence": case.get("sentence"),
            "court": case.get("court"),
        },
        merged_facts=merged_facts,
        merged_events=merged_events,
        merged_findings=merged_findings,
        document_extract_ids=extract_ids,
    )
    case_extract_dict = case_extract.model_dump()
    case_extract_dict["created_at"] = case_extract_dict["created_at"].isoformat()
    await db.case_extracts.update_one(
        {"case_id": case["case_id"], "user_id": case["user_id"]},
        {"$set": case_extract_dict},
        upsert=True,
    )
    return case_extract_dict


async def _ensure_issue_classifications(case: dict, case_extract: dict) -> list[dict]:
    """Run staged issue classification and persist results.
    
    DO_NOT_UNDO — 3 Apr 2026: If issues already exist for this case, DO NOT re-classify.
    Re-classification was the ROOT CAUSE of grounds multiplying — every LLM call generates
    slightly different titles which slip past dedup and create new grounds.
    Only classify if zero issues exist (first-time analysis).
    """
    from services.ground_dedup import is_ground_duplicate, normalise_au_spelling
    
    # Check if issues already exist — if so, return them without re-classifying
    existing_issues = await db.issue_classifications.find(
        {"case_id": case["case_id"], "user_id": case["user_id"]},
        {"_id": 0}
    ).to_list(500)
    
    # Check if enough issues already exist — only skip if 8+ exist (healthy count)
    # If fewer than 8, re-classify to find missing grounds (previous dedup may have been too aggressive)
    if len(existing_issues) >= 8:
        logger.info(f"Skipping re-classification for case {case['case_id']}: {len(existing_issues)} issues already exist (>= 8)")
        return existing_issues
    
    logger.info(f"Re-classifying case {case['case_id']}: only {len(existing_issues)} issues exist (< 8), looking for more")
    issues = await classify_case_issues(case, case_extract)
    
    persisted = list(existing_issues)  # Start with existing issues for dedup comparison
    for issue in issues:
        issue_dict = issue.model_dump()
        issue_dict["created_at"] = issue_dict["created_at"].isoformat()
        issue_title = normalise_au_spelling((issue.title or "").strip())
        issue_dict["title"] = issue_title
        
        # Fuzzy match against ALL existing + newly-persisted issues
        matched = None
        for ei in persisted:
            ei_title = (ei.get("title") or "").strip()
            if is_ground_duplicate(issue_title, ei_title):
                matched = ei
                break
        
        if matched:
            await db.issue_classifications.update_one(
                {"issue_id": matched["issue_id"]},
                {"$set": issue_dict},
            )
            saved = await db.issue_classifications.find_one(
                {"issue_id": matched["issue_id"]}, {"_id": 0}
            )
        else:
            await db.issue_classifications.update_one(
                {"case_id": case["case_id"], "user_id": case["user_id"], "title": issue_title, "ground_type": issue.ground_type},
                {"$set": issue_dict},
                upsert=True,
            )
            saved = await db.issue_classifications.find_one(
                {"case_id": case["case_id"], "user_id": case["user_id"], "title": issue_title, "ground_type": issue.ground_type},
                {"_id": 0}
            )
        
        if saved:
            persisted.append(saved)
    return persisted


async def _pipeline_artifacts_missing_or_stale(case: dict, documents: list) -> bool:
    """
    Conservative staleness check.
    Returns True if:
    - no case_extract exists, or
    - no issue classifications exist, or
    - any uploaded document lacks a document_extract
    """
    case_extract = await db.case_extracts.find_one(
        {"case_id": case["case_id"], "user_id": case["user_id"]},
        {"_id": 0}
    )
    if not case_extract:
        return True

    issue_count = await db.issue_classifications.count_documents(
        {"case_id": case["case_id"], "user_id": case["user_id"]}
    )
    if issue_count == 0:
        return True

    for document in documents:
        existing = await db.document_extracts.find_one(
            {
                "case_id": case["case_id"],
                "document_id": document["document_id"],
                "user_id": case["user_id"],
            },
            {"_id": 0}
        )
        if not existing:
            return True

    return False


async def _sync_pipeline_projection_to_grounds(case: dict) -> int:
    """
    Sync staged issues/verifications into grounds_of_merit so report consumers
    and existing frontend views remain aligned.

    DO_NOT_UNDO — 3 Apr 2026: HARD CAP on ground creation.
    If grounds already exist, NEVER create more than existing_count + 2.
    This prevents the recurring multiplication bug permanently.
    """
    from services.ground_dedup import is_ground_duplicate, normalise_au_spelling

    issues = await db.issue_classifications.find(
        {"case_id": case["case_id"], "user_id": case["user_id"]},
        {"_id": 0}
    ).to_list(500)

    # Pre-load existing grounds for fuzzy matching
    all_existing_grounds = await db.grounds_of_merit.find(
        {"case_id": case["case_id"], "user_id": case["user_id"]},
        {"_id": 0, "ground_id": 1, "title": 1},
    ).to_list(200)

    # DO_NOT_UNDO — HARD CAP enforcement. If grounds already exist,
    # NEVER create more than existing_count + 2. This prevents the
    # recurring multiplication bug permanently.
    initial_ground_count = len(all_existing_grounds)
    max_new_grounds = 2 if initial_ground_count > 0 else 50
    new_grounds_created = 0

    synced = 0
    for issue in issues:
        verification = await db.issue_verifications.find_one(
            {"case_id": case["case_id"], "issue_id": issue["issue_id"], "user_id": case["user_id"]},
            {"_id": 0}
        )

        issue_title = normalise_au_spelling((issue.get("title") or "").strip())

        # Fuzzy match against all existing grounds
        existing_ground = None
        for eg in all_existing_grounds:
            eg_title = (eg.get("title") or "").strip()
            if is_ground_duplicate(issue_title, eg_title):
                existing_ground = eg
                break

        if not existing_ground:
            logger.info(f"Sync: new ground '{issue_title[:50]}'")

        ground_doc = {
            "case_id": case["case_id"],
            "user_id": case["user_id"],
            "title": existing_ground["title"] if existing_ground else issue_title,
            "ground_type": issue.get("ground_type", "other"),
            "description": issue.get("description", ""),
            "strength": (verification or {}).get("legitimacy_scores", {}).get(
                "rating", issue.get("classification_confidence", "moderate")
            ),
            "status": "investigated" if verification else "identified",
            "supporting_evidence": (verification or {}).get("supporting_items", []),
            "law_sections": (verification or {}).get("law_sections", []),
            "similar_cases": (verification or {}).get("similar_cases", []),
            "legitimacy_scores": (verification or {}).get("legitimacy_scores", {}),
            "source_mode": "derived",
            "requires_human_review": (verification or {}).get("requires_human_review", True),
            "verification_status": (verification or {}).get("verification_status", "unverified"),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }

        if existing_ground:
            await db.grounds_of_merit.update_one(
                {"ground_id": existing_ground["ground_id"]},
                {"$set": ground_doc},
            )
        else:
            # DO_NOT_UNDO — Hard cap check: skip if at limit
            if new_grounds_created >= max_new_grounds:
                logger.info(f"Sync: HARD CAP reached ({initial_ground_count}+{max_new_grounds}). Skipping '{issue_title[:50]}'")
                continue
            ground_doc["ground_id"] = f"gnd_{uuid.uuid4().hex[:12]}"
            ground_doc["created_at"] = datetime.now(timezone.utc).isoformat()
            await db.grounds_of_merit.insert_one(ground_doc)
            all_existing_grounds.append({"ground_id": ground_doc["ground_id"], "title": issue_title})
            new_grounds_created += 1

        synced += 1
    return synced


async def _refresh_pipeline_for_reporting(case: dict, documents: list, report_type: str = "quick_summary") -> dict:
    """
    Reporting-time pipeline refresh:
    1. ensure document extracts
    2. refresh case extract
    3. ensure issue classifications
    4. optionally auto-verify top issues by report tier
    5. sync projection into grounds
    6. DO_NOT_UNDO — safety net dedup cleanup after sync
    """
    from services.ground_dedup import cleanup_duplicate_grounds, cleanup_duplicate_issues

    extracted_count = await _ensure_document_extracts_for_case(case, documents)
    case_extract = await _refresh_case_extract_for_case(case)
    issues = await _ensure_issue_classifications(case, case_extract)

    auto_verify_limit = _auto_verification_limit_for_report_type(report_type)
    selected_issues = await _select_issues_for_auto_verification(case, auto_verify_limit)
    auto_verify_result = await _auto_verify_selected_issues(case, selected_issues)

    synced_count = await _sync_pipeline_projection_to_grounds(case)

    # Safety net: clean up any duplicates that may have slipped through
    await cleanup_duplicate_grounds(db, case["case_id"], case["user_id"])
    await cleanup_duplicate_issues(db, case["case_id"], case["user_id"])

    return {
        "extracted_count": extracted_count,
        "classified_count": len(issues),
        "synced_count": synced_count,
        "auto_verify_limit": auto_verify_limit,
        "auto_verify_result": auto_verify_result,
    }


def _issue_priority_rank(issue: dict) -> tuple:
    """Lower tuple sorts earlier."""
    preferred_ground_order = {
        "judicial_error": 0,
        "procedural_error": 1,
        "miscarriage_of_justice": 2,
        "fresh_evidence": 3,
        "sentencing_error": 4,
        "jury_irregularity": 5,
        "ineffective_counsel": 6,
        "prosecution_misconduct": 7,
        "constitutional_violation": 8,
        "other": 9,
    }
    confidence_order = {
        "strong": 0,
        "moderate": 1,
        "weak": 2,
    }
    return (
        preferred_ground_order.get(issue.get("ground_type", "other"), 9),
        confidence_order.get(issue.get("classification_confidence", "moderate"), 1),
        str(issue.get("title", "")).lower(),
    )


async def _select_issues_for_auto_verification(case: dict, limit: int) -> list[dict]:
    """Select issues that do not yet have verifications."""
    if limit <= 0:
        return []
    issues = await db.issue_classifications.find(
        {"case_id": case["case_id"], "user_id": case["user_id"]}, {"_id": 0}
    ).to_list(500)
    if not issues:
        return []
    eligible = []
    for issue in issues:
        existing_verification = await db.issue_verifications.find_one(
            {"case_id": case["case_id"], "issue_id": issue["issue_id"], "user_id": case["user_id"]},
            {"_id": 0}
        )
        if existing_verification:
            continue
        eligible.append(issue)
    eligible.sort(key=_issue_priority_rank)
    return eligible[:limit]


async def _auto_verify_selected_issues(case: dict, selected_issues: list[dict]) -> dict:
    """Verify selected issues and persist results."""
    if not selected_issues:
        return {"attempted": 0, "verified": 0, "failed": 0}
    case_extract = await db.case_extracts.find_one(
        {"case_id": case["case_id"], "user_id": case["user_id"]}, {"_id": 0}
    )
    if not case_extract:
        return {"attempted": len(selected_issues), "verified": 0, "failed": len(selected_issues)}
    supporting_context = {
        "facts": case_extract.get("merged_facts", []),
        "events": case_extract.get("merged_events", []),
        "findings": case_extract.get("merged_findings", []),
    }
    verified = 0
    failed = 0
    for issue in selected_issues:
        try:
            verification = await verify_issue(case, issue, supporting_context)
            verification_dict = verification.model_dump()
            verification_dict["created_at"] = verification_dict["created_at"].isoformat()
            await db.issue_verifications.update_one(
                {"case_id": case["case_id"], "issue_id": issue["issue_id"], "user_id": case["user_id"]},
                {"$set": verification_dict},
                upsert=True,
            )
            verified += 1
        except Exception as e:
            logger.warning(f"Auto-verification failed for case {case['case_id']} issue {issue.get('issue_id')}: {e}")
            failed += 1
    return {"attempted": len(selected_issues), "verified": verified, "failed": failed}


def _auto_verification_limit_for_report_type(report_type: str) -> int:
    if report_type in ("full_report", "full_detailed"):
        return 3
    if report_type in ("extensive_report", "extensive_log"):
        return 6
    return 0


async def _check_case_pipeline_staleness(case_id: str, user_id: str) -> dict:
    """Check pipeline staleness without needing a Request object."""
    documents = await db.documents.find(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0, "document_id": 1, "uploaded_at": 1}
    ).to_list(1000)

    doc_extracts = await db.document_extracts.find(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0, "document_id": 1, "created_at": 1}
    ).to_list(1000)

    case_extract = await db.case_extracts.find_one(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0, "created_at": 1}
    )

    issue_classifications = await db.issue_classifications.find(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0, "created_at": 1}
    ).to_list(1000)

    issue_verifications = await db.issue_verifications.find(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0, "created_at": 1}
    ).to_list(1000)

    reports = await db.reports.find(
        {"case_id": case_id, "user_id": user_id, "content.draft_source": "pipeline"},
        {"_id": 0, "generated_at": 1}
    ).to_list(1000)

    def _to_iso(value):
        if not value:
            return None
        if isinstance(value, str):
            return value
        try:
            return value.isoformat()
        except Exception:
            return str(value)

    def _latest_iso(items, field):
        values = [i.get(field) for i in items if i.get(field)]
        if not values:
            return None
        values = [v if isinstance(v, str) else v.isoformat() for v in values]
        return max(values)

    latest_document_upload = _latest_iso(documents, "uploaded_at")
    latest_document_extract = _latest_iso(doc_extracts, "created_at")
    latest_issue_classification = _latest_iso(issue_classifications, "created_at")
    latest_issue_verification = _latest_iso(issue_verifications, "created_at")
    latest_pipeline_report = _latest_iso(reports, "generated_at")
    case_extract_created_at = _to_iso(case_extract.get("created_at")) if case_extract else None

    extract_map = {d.get("document_id"): d for d in doc_extracts}
    extract_missing = [d.get("document_id") for d in documents if d.get("document_id") not in extract_map]

    documents_newer = (latest_document_upload > latest_document_extract) if (latest_document_upload and latest_document_extract) else bool(latest_document_upload and not latest_document_extract)
    case_extract_stale = (latest_document_extract > case_extract_created_at) if (latest_document_extract and case_extract_created_at) else bool(latest_document_extract and not case_extract_created_at)
    classifications_stale = (case_extract_created_at > latest_issue_classification) if (case_extract_created_at and latest_issue_classification) else bool(case_extract_created_at and not latest_issue_classification)
    verifications_stale = (latest_issue_classification > latest_issue_verification) if (latest_issue_classification and latest_issue_verification) else bool(latest_issue_classification and not latest_issue_verification)
    reports_stale = (latest_issue_verification > latest_pipeline_report) if (latest_issue_verification and latest_pipeline_report) else bool(latest_issue_verification and not latest_pipeline_report)

    overall_stale = any([len(extract_missing) > 0, documents_newer, case_extract_stale, classifications_stale, verifications_stale, reports_stale])

    return {
        "overall_stale": overall_stale,
        "extract_missing_for_documents": extract_missing,
        "documents_newer_than_extracts": documents_newer,
        "case_extract_stale": case_extract_stale,
        "classifications_stale": classifications_stale,
        "verifications_stale": verifications_stale,
        "reports_stale": reports_stale,
    }


async def _enforce_pipeline_freshness(case_id: str, user_id: str, auto_refresh: bool = False) -> dict:
    """Ensure pipeline artifacts are current before report generation."""
    staleness = await _check_case_pipeline_staleness(case_id, user_id)

    if not staleness.get("overall_stale"):
        return {"status": "fresh", "staleness": staleness}

    if not auto_refresh:
        return {"status": "stale", "staleness": staleness}

    # Auto-refresh pipeline
    case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
    if not case:
        return {"status": "stale", "staleness": staleness}

    docs = await db.documents.find(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0, "file_data": 0}
    ).to_list(1000)

    extract_result = await _staged_ensure_extracts(case, docs)
    case_extract_result = await _staged_refresh_case(case)
    ce = await db.case_extracts.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
    classify_result = await _staged_classify(case, ce)
    verify_result = await _staged_verify(case, 3)
    synced = await _staged_sync_grounds(case_id, user_id)

    return {
        "status": "refreshed",
        "staleness": staleness,
        "refresh_result": {
            "documents": extract_result,
            "case_extract": case_extract_result,
            "classification": classify_result,
            "verification": verify_result,
            "projection": {"synced_count": synced},
        },
    }


async def _load_issue_arguments(case_id: str, user_id: str) -> list:
    return await db.issue_arguments.find(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0}
    ).to_list(500)


async def _load_submission_draft(case_id: str, user_id: str):
    return await db.submissions_drafts.find_one(
        {"case_id": case_id, "user_id": user_id},
        {"_id": 0}
    )


async def analyze_case_with_ai(case_id: str, user_id: str, report_type: str, aggressive_mode: bool = False, report_id: str = None) -> dict:
    """Use AI to analyse case and generate report — HARDENED with structured LLM calls"""
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # ================= PIPELINE GUARD =================
    high_value_report = report_type in ["full_detailed", "extensive_log"]
    pipeline_status = None

    if high_value_report:
        try:
            pipeline_status = await _enforce_pipeline_freshness(
                case_id,
                user_id,
                auto_refresh=True,
            )
            logger.info(f"Pipeline guard for {case_id}/{report_type}: {pipeline_status.get('status')}")
        except Exception as pg_err:
            logger.warning(f"Pipeline guard failed (non-fatal) for {case_id}: {pg_err}")

    # ================= LOAD PIPELINE DATA =================
    structured_context = None
    if high_value_report:
        case_extract_data = await db.case_extracts.find_one(
            {"case_id": case_id, "user_id": user_id}, {"_id": 0}
        )
        issue_verifications_data = await db.issue_verifications.find(
            {"case_id": case_id, "user_id": user_id}, {"_id": 0}
        ).to_list(100)

        structured_context = {
            "facts": case_extract_data.get("merged_facts", []) if case_extract_data else [],
            "events": case_extract_data.get("merged_events", []) if case_extract_data else [],
            "findings": case_extract_data.get("merged_findings", []) if case_extract_data else [],
            "verified_issues": issue_verifications_data,
        }

    # Load issue arguments (if any exist)
    pipeline_arguments = await _load_issue_arguments(case_id, user_id) if high_value_report else []
    submission_draft = await _load_submission_draft(case_id, user_id) if high_value_report else None

    # Get documents with full content
    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0, "file_data": 0}
    ).to_list(500)
    
    # Get timeline
    timeline = await db.timeline_events.find(
        {"case_id": case_id},
        {"_id": 0}
    ).sort("event_date", 1).to_list(500)
    
    # Get notes
    notes = await db.notes.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Get identified grounds
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # ── Auto-detect case metadata if still defaults (HARDENED via call_llm_for_json) ──
    # DO NOT UNDO — Also re-detect if sentence contains crime narrative (e.g. "for murdering")
    _current_sentence = (case.get('sentence') or "").strip()
    _sentence_has_narrative = bool(re.search(r'\bfor\s+(?:murder|kill|assault|robb|stab|rap|kidnap|abus|supplying|dealing)', _current_sentence, re.I)) if _current_sentence else False
    needs_detection = (
        (not case.get('offence_category') or case.get('offence_category') == 'homicide')
        and not case.get('offence_type')
        and (not _current_sentence or _sentence_has_narrative)
    )
    if needs_detection and documents:
        try:
            combined_text = ""
            for d in documents:
                ct = d.get("content_text") or ""
                if ct:
                    combined_text += f"\n--- {d.get('filename','')} ---\n{ct[:6000]}\n"
            if len(combined_text) > 200:
                detect_prompt = f"""Analyse the following case documents and extract metadata.
Return ONLY valid JSON (no markdown, no explanation):
{{
  "offence_category": "<one of: homicide, assault, sexual_offences, robbery_theft, drug_offences, fraud_dishonesty, firearms_weapons, domestic_violence, public_order, terrorism, driving_offences>",
  "offence_type": "<specific offence e.g. Murder, Assault Occasioning ABH, Armed Robbery>",
  "sentence": "<the HEAD SENTENCE and non-parole period ONLY. Format: '[X] years imprisonment with a non-parole period of [Y] years [and Z months]'. Do NOT include the crime description, victim details, or case narrative. If life sentence, state 'Life imprisonment with a non-parole period of [X] years'. If sentence not stated, null.>",
  "state": "<one of: nsw, vic, qld, sa, wa, tas, nt, act — if determinable>",
  "court": "<court name if found>"
}}
Rules:
- offence_category MUST be from the listed values only. Read the ACTUAL documents — do NOT default to homicide.
- If the document is about assault, category is "assault". If robbery, "robbery_theft", etc.
- sentence must be ONLY the numerical sentence and non-parole period. Never include what the crime was or who the victim was.
- If a field cannot be determined, set to null. state must be lowercase.

DOCUMENTS:
{combined_text[:30000]}"""

                def _validate_metadata(parsed: dict) -> bool:
                    valid_cats = ["homicide","assault","sexual_offences","robbery_theft","drug_offences","fraud_dishonesty","firearms_weapons","domestic_violence","public_order","terrorism","driving_offences"]
                    cat = parsed.get("offence_category")
                    return cat in valid_cats if cat else True

                meta = await call_llm_for_json(
                    system_prompt="You are a legal document analyst. Extract factual metadata only. Return valid JSON.",
                    user_prompt=detect_prompt,
                    session_id=f"rpt_detect_{case_id}",
                    task_type="metadata_detection",
                    validation_fn=_validate_metadata,
                )

                valid_cats = ["homicide","assault","sexual_offences","robbery_theft","drug_offences","fraud_dishonesty","firearms_weapons","domestic_violence","public_order","terrorism","driving_offences"]
                valid_sts = ["nsw","vic","qld","sa","wa","tas","nt","act"]
                update_fields = {}
                if meta.get("offence_category") in valid_cats:
                    update_fields["offence_category"] = meta["offence_category"]
                if meta.get("offence_type"):
                    update_fields["offence_type"] = meta["offence_type"]
                if meta.get("sentence"):
                    # DO NOT UNDO — Clean sentence to strip crime descriptions, victim info
                    cleaned_sentence = _clean_sentence_candidate(meta["sentence"])
                    if _is_valid_sentence_candidate(cleaned_sentence):
                        update_fields["sentence"] = cleaned_sentence
                    elif meta["sentence"].strip():
                        update_fields["sentence"] = meta["sentence"].strip()
                if meta.get("state") and str(meta["state"]).lower() in valid_sts:
                    update_fields["state"] = str(meta["state"]).lower()
                if meta.get("court"):
                    update_fields["court"] = meta["court"]
                if update_fields:
                    update_fields["updated_at"] = datetime.now(timezone.utc).isoformat()
                    update_fields["classification_source"] = "ai_detected"
                    update_fields["requires_metadata_review"] = True
                    await db.cases.update_one({"case_id": case_id}, {"$set": update_fields})
                    # Re-read case with updated metadata
                    case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
                    logger.info(f"Report gen auto-detected metadata for {case_id}: {update_fields}")
        except Exception as e:
            logger.warning(f"Report gen auto-detect failed for {case_id}: {e}")
    
    # Get offence-specific context
    offence_category = case.get('offence_category', 'homicide')
    offence_context = get_offence_context(case)
    category_data = OFFENCE_CATEGORIES.get(offence_category, OFFENCE_CATEGORIES.get('homicide'))
    category_name = category_data.get('name', 'criminal')
    state = case.get('state', 'nsw')
    state_info = AUSTRALIAN_STATES.get(state, AUSTRALIAN_STATES.get('nsw'))
    
    # Context limits — FULL DETAIL for quality reports
    context_limits = {
        "quick_summary": {
            "per_doc_chars": 700,
            "total_doc_chars": 6000,
            "timeline_limit": 40,
            "notes_limit": 20,
            "note_chars": 180,
            "grounds_limit": 30,
            "ground_desc_chars": 200,
            "ground_analysis_chars": 200,
            "ground_deep_chars": 0,
        },
        "full_detailed": {
            "per_doc_chars": 2400,
            "total_doc_chars": 24000,
            "timeline_limit": 150,
            "notes_limit": 80,
            "note_chars": 500,
            "grounds_limit": 80,
            "ground_desc_chars": 600,
            "ground_analysis_chars": 500,
            "ground_deep_chars": 0,
        },
        "extensive_log": {
            "per_doc_chars": 3200,
            "total_doc_chars": 32000,
            "timeline_limit": 200,
            "notes_limit": 100,
            "note_chars": 600,
            "grounds_limit": 100,
            "ground_desc_chars": 800,
            "ground_analysis_chars": 700,
            "ground_deep_chars": 700,
        },
    }
    limits = context_limits.get(report_type, context_limits["quick_summary"])

    # Prepare comprehensive context for AI with bounded document content
    case_context = f"""
=== CASE INFORMATION ===
Title: {case.get('title', 'N/A')}
Defendant: {case.get('defendant_name', 'N/A')}
Case Number: {case.get('case_number', 'N/A')}
Court: {case.get('court', 'N/A')}
Judge: {case.get('judge', 'N/A')}
Sentence: {case.get('sentence', 'N/A')}
Summary: {case.get('summary', 'N/A')}

{offence_context}
"""

    doc_context = build_document_context(
        documents,
        per_doc_char_limit=limits["per_doc_chars"],
        total_char_budget=limits["total_doc_chars"],
        include_description=True,
        content_heading="CONTENT"
    )
    
    # Include bounded document content for cross-referencing
    if documents:
        case_context += f"=== CASE DOCUMENTS ({doc_context['included_docs']} included / {len(documents)} total files) ===\n"
        case_context += doc_context["text"] + "\n"
        if doc_context["omitted_docs"] > 0:
            case_context += f"[Note: {doc_context['omitted_docs']} document(s) omitted from prompt for speed optimisation.]\n"
    else:
        case_context += "=== NO DOCUMENTS UPLOADED ===\n"

    if timeline:
        timeline_slice = timeline[:limits["timeline_limit"]]
        case_context += f"\n=== TIMELINE OF EVENTS ({len(timeline_slice)} included / {len(timeline)} total events) ===\n"
        for event in timeline_slice:
            case_context += f"- {event.get('event_date', 'Unknown')}: [{event.get('event_type')}] {event.get('title')}\n"
            if event.get('description'):
                case_context += f"  Details: {event.get('description')}\n"
        if len(timeline) > limits["timeline_limit"]:
            case_context += f"[... {len(timeline) - limits['timeline_limit']} additional events omitted for speed ...]\n"

    if notes:
        notes_slice = notes[:limits["notes_limit"]]
        case_context += f"\n=== LEGAL NOTES ({len(notes_slice)} included / {len(notes)} total notes) ===\n"
        for note in notes_slice:
            case_context += f"- [{note.get('category')}] {note.get('title')}: {note.get('content', '')[:limits['note_chars']]}\n"
        if len(notes) > limits["notes_limit"]:
            case_context += f"[... {len(notes) - limits['notes_limit']} additional notes omitted for speed ...]\n"

    if grounds:
        grounds_slice = grounds[:limits["grounds_limit"]]
        case_context += f"\n=== IDENTIFIED GROUNDS OF MERIT ({len(grounds_slice)} included / {len(grounds)} total grounds) ===\n"
        for g in grounds_slice:
            viability = g.get('legitimacy_scores', {}).get('viability_label', g.get('strength', 'moderate'))
            case_context += f"- [{g.get('ground_type')}] {g.get('title')} (Viability: {viability})\n"
            case_context += f"  {g.get('description', '')[:limits['ground_desc_chars']]}\n"
            if g.get('appellate_pathway'):
                case_context += f"  Appellate Pathway: {g.get('appellate_pathway')}\n"
            if g.get('error_identified'):
                case_context += f"  Error: {g.get('error_identified')}\n"
            if report_type != "quick_summary" and g.get('analysis'):
                case_context += f"  Analysis: {g.get('analysis', '')[:limits['ground_analysis_chars']]}\n"
            if limits["ground_deep_chars"] > 0 and g.get('deep_analysis'):
                deep = g.get('deep_analysis', '')
                if isinstance(deep, str):
                    case_context += f"  Deep Analysis: {deep[:limits['ground_deep_chars']]}\n"
        if len(grounds) > limits["grounds_limit"]:
            case_context += f"[... {len(grounds) - limits['grounds_limit']} additional grounds omitted for speed ...]\n"

    grounds_titles = [g.get('title') for g in grounds if g.get('title')]
    grounds_enumerated = "\n".join([f"{idx + 1}. {title}" for idx, title in enumerate(grounds_titles)])

    # ================= INJECT STRUCTURED PIPELINE DATA =================
    if high_value_report and structured_context:
        facts_text = "\n".join([f"- {f}" if isinstance(f, str) else f"- {f}" for f in structured_context["facts"]]) if structured_context["facts"] else "No extracted facts available."
        events_text = "\n".join([f"- {e}" if isinstance(e, str) else f"- {e}" for e in structured_context["events"]]) if structured_context["events"] else "No extracted events available."
        findings_text = "\n".join([f"- {f}" if isinstance(f, str) else f"- {f}" for f in structured_context["findings"]]) if structured_context["findings"] else "No extracted findings available."

        verified_text = ""
        for vi in structured_context["verified_issues"]:
            verified_text += f"\n--- Issue: {vi.get('title', 'Untitled')} (Type: {vi.get('ground_type', 'unknown')}) ---\n"
            verified_text += f"  Verification status: {vi.get('verification_status', 'unverified')}\n"
            if vi.get("supporting_items"):
                verified_text += f"  Supporting items: {len(vi['supporting_items'])}\n"
            if vi.get("undermining_items"):
                verified_text += f"  Undermining items: {len(vi['undermining_items'])}\n"
            if vi.get("legitimacy_scores"):
                verified_text += f"  Legitimacy: {vi['legitimacy_scores']}\n"

        case_context += f"""
=== STRUCTURED PIPELINE DATA (VERIFIED MATERIAL) ===

CASE FACTS:
{facts_text}

TIMELINE EVENTS:
{events_text}

KEY FINDINGS:
{findings_text}

VERIFIED APPEAL ISSUES:
{verified_text if verified_text else "No verified issues available."}
"""

    # Define prompts based on report type with offence-specific language
    base_system = get_offence_system_prompt(offence_category)
    report_guardrails = """
MANDATORY GUARDRAILS:
- Use a HYBRID tone: court-ready legal analysis + plain-English action notes for the client.
- Use Australian English only.
- Anchor findings to supplied case material; clearly mark uncertainty when evidence is incomplete.
- Include legislation as: section + Act + jurisdiction + year (e.g. s.18 Crimes Act 1900 (NSW)).
- Include sentencing comparisons and precedent outcomes where relevant.
- DO NOT include cost estimates, fee ranges, funding commentary, or budget analysis.
- DO NOT include witness contradiction sections or witness credibility scoring sections.

LANGUAGE RULES — ABSOLUTE AND NON-NEGOTIABLE:
- This report is an EDUCATIONAL TOOL. It is NOT written by a legal team speaking on behalf of the applicant.
- NEVER use the words "we", "us", "our", or "them" when referring to the legal team, analysis team, or report authors.
- NEVER use the words "you" or "your" when addressing the reader or the applicant. WRONG: "The appeal is your opportunity to challenge your conviction." RIGHT: "The appeal represents an opportunity for the applicant to challenge the conviction."
- Instead of "we are arguing" write "the applicant argues" or "this analysis identifies".
- Instead of "we will file" write "the legal professional will file" or "the applicant should file".
- Instead of "our submissions" write "the submissions" or "the applicant's submissions".
- Instead of "contact us" write "contact the legal professional" or "contact the assisting legal practitioner".
- Instead of "we are aiming to show" write "the appeal aims to demonstrate" or "the applicant seeks to establish".
- Instead of "our claims" write "the applicant's claims" or "the claims advanced".
- Instead of "your appeal" write "the appeal". Instead of "your conviction" write "the conviction". Instead of "your sentence" write "the sentence".
- Instead of "you should" write "the applicant should". Instead of "you may" write "the applicant may". Instead of "for you" write "for the applicant".
- The report must read as a neutral educational analysis document, NOT as a first-person legal team communication and NOT as a second-person advisory letter.
- Use third-person references throughout: "the applicant", "the defendant", "the legal professional", "this analysis", "the appeal".
- Violations of this rule make the report legally problematic and unprofessional.

NO PLACEHOLDER TEXT — ABSOLUTE RULE:
- NEVER write future-tense placeholder text like "The table will reference...", "Details will be provided...", "This section will include...", "Content will be added...", or "Analysis will cover...".
- Every sentence must contain ACTUAL analysis, ACTUAL data, or ACTUAL legal content. If a table cannot be populated with real data from the case material, populate it with the best available comparable cases from Australian jurisprudence.
- If information is unavailable, state that explicitly ("No sentencing comparisons could be identified from the supplied material") rather than promising future content.

FORENSIC APPELLATE LANGUAGE — ABSOLUTE AND NON-NEGOTIABLE:
Appellate work is about ARGUABILITY, not declarations. All conclusions must be framed forensically:
- BANNED PHRASES (NEVER use these): "The trial judge erred", "The judge clearly erred", "This clearly shows", "This proves", "This demonstrates conclusively", "The conviction is unsafe", "The sentence is excessive", "The error is established".
- REQUIRED FORENSIC FRAMING (use these instead):
  * "It is arguable that the trial judge erred in..."
  * "It is contended that..."
  * "There is a tenable argument that..."
  * "On one view of the evidence..."
  * "It may be submitted that..."
  * "A reasonable argument exists that..."
  * "The available material supports the contention that..."
  * "It is open to the appellant to argue that..."
- The distinction is critical: an appeal brief identifies ARGUABLE errors, it does not declare findings of fact. The Court of Criminal Appeal makes findings — the brief identifies where findings are open.
- Exception: When citing what a court HAS decided in a precedent case, declarative language is appropriate (e.g. "In R v Smith, the Court held that...").
- For AGGRESSIVE MODE only: stronger language is permitted (e.g. "The Crown's position is untenable", "The error is compelling") but still avoid absolute declarations about the current case.

CONTENT QUALITY — STRICTLY ENFORCED (violations make the report worthless):
- DO the analysis. Do NOT describe what analysis should be done. WRONG: "Delve into aggravating and mitigating factors." RIGHT: "Under s.21A(2) of the Crimes (Sentencing Procedure) Act 1999 (NSW), the aggravating factors in Homann's case include the use of a weapon and the vulnerability of the victim. However, it is arguable that the sentencing judge failed to give adequate weight to the mitigating factor under s.21A(3)(d)..."
- NEVER create filler sections with titles like "URGENCY PRIORITY", "RELEVANCE", "KEY TAKEAWAY", "SUMMARY", "OVERVIEW" as standalone sections. These are padding. Instead, weave relevance and urgency INTO the substantive analysis.
- NEVER write generic consultant-speak like "Leverage legal databases to draw parallels that authenticate excessive sentencing claims through empirical trends." Instead, NAME the specific cases, cite the specific sentencing outcomes, and EXPLAIN the specific parallels.
- Every paragraph MUST reference specific names, dates, section numbers, case citations, or document names from the supplied case material. If a paragraph could apply to ANY appeal case, it is too generic — rewrite it with THIS case's specific facts.
- Avoid repetition across sections. If a point is already covered, deepen it with NEW evidence, dates, or authority rather than restating it.
- Do NOT reuse boilerplate phrases. Every paragraph must read as drafted specifically for this case and this report tier.
- For legislation sections: Do NOT just name the Act and describe what it covers in general terms. APPLY each provision to THIS case's specific facts. WRONG: "s.44 discusses parole periods, directly affecting Homann's sentencing outcomes." RIGHT: "Under s.44 of the Crimes (Sentencing Procedure) Act 1999 (NSW), the non-parole period must reflect the objective seriousness of the offence. In Homann's case, the 22-year non-parole period imposed by Justice McCallum is arguable as disproportionate when compared with R v Loveridge [2014] NSWCCA 120 where a 7-year non-parole period was imposed for a one-punch manslaughter..."
- For precedent/sentencing tables: Include the full case citation, the specific factual similarity to THIS case, the actual sentence imposed, and the specific relevance to the current appeal. NEVER use a one-line vague description.

FORMATTING RULES — STRICTLY ENFORCED:
- DO NOT begin your response with any preamble, greeting, or introduction.
- DO NOT use placeholder notes in brackets like "[Note: Continue...]", "[Insert details...]". Every section MUST contain COMPLETE, REAL content.
- Every section heading MUST be followed by substantive content (minimum 3-4 detailed paragraphs). If a section cannot be substantiated from the case material, omit it entirely.
- Include the year in ALL legislation references (e.g. Crimes Act 1900 (NSW), NOT just Crimes Act (NSW)).
- SECTION HEADINGS: Use ONLY ## for numbered section headings (e.g. ## 1. EXECUTIVE BRIEF). Do NOT create sub-sections with ### headings. Do NOT put bold text on its own line as a sub-heading. Instead, write flowing paragraphs and use bold text inline (e.g. "The **legal threshold** for this ground requires...").
- FOR GROUND ANALYSIS: Write each ground as a continuous series of detailed paragraphs (minimum 300 words in Quick Summary, 800+ words in Full Detailed, 1200+ words in Extensive). Do NOT use bullet points. Cover the legal threshold, case facts, viability, Crown response, defence rebuttal, and impact all within flowing prose.
"""
    
    if report_type == "quick_summary":
        system_prompt = f"""{base_system}
{report_guardrails}
You are generating a FREE Quick Summary — an ISSUE IDENTIFICATION report. Its purpose is to identify and list the potential grounds of appeal, NOT to provide deep legal analysis. Deliver real legal value in a concise overview, then clearly explain what deeper paid reports add. IMPORTANT: Write at least 2000 words. Narrative sections (Case Snapshot, Sentencing Overview, Value Statement) must have 3-5 substantive paragraphs. Structured sections (Issues, Grounds, Legislation) use the specified format (numbered items, lists, or tables). Use forensic appellate language: "It is arguable that...", "It is contended that...", "There is a tenable argument that..." — NOT bare declarations like "The judge erred" or hedging like "may have" or "could potentially"."""
        user_prompt = f"""Analyse this {category_name.lower()} appeal matter and produce a QUICK SUMMARY REPORT (Issue Identification).

{case_context}

Write MINIMUM 2000 words. This is an ISSUE IDENTIFICATION report — identify all potential grounds clearly but do not provide the deep appellate pathway analysis (that is for paid reports). Structure EXACTLY as follows:

## 1. CASE SNAPSHOT
3-4 paragraphs covering: defendant, offence, jurisdiction, sentence imposed, non-parole period, key procedural dates, presiding judge, and what material was reviewed. Be specific to the supplied facts.

MANDATORY: Start this section with: "This analysis is based on [X] documents and [Y] timeline events. [Z] grounds of appeal have been identified." Use EXACT counts from supplied data.

## 2. PRIMARY ISSUES IDENTIFIED
6-10 numbered items. Each must include:
- Issue title framed as an appellate ground (e.g. "Failure to Properly Evaluate Psychiatric Evidence")
- Which document/evidence it comes from
- Why it matters for the appeal

## 3. ALL GROUNDS IDENTIFIED (PREVIEW)
You MUST list EVERY ground identified. For each ground:
- Ground title + type (e.g., Procedural Error, Sentencing Error, Miscarriage of Justice)
- Appellate viability: Arguable — Strong / Arguable — Moderate / Requires Development
- 2-3 sentence legal rationale referencing the specific case facts
- One immediate action step

Do NOT use percentage success rates. Use appellate viability language only.

## 4. KEY LEGISLATION (PREVIEW)
List the most relevant statutory provisions from {state_info.get('name', 'NSW')} / Commonwealth law with ACTUAL section numbers, years, and one-line relevance notes. If the exact section number is not known, do NOT include that entry.

## 5. SENTENCING OVERVIEW
2-3 paragraphs comparing the sentence imposed against typical appellate principles for this offence category. Include a mini comparison table:
| Comparator Case | Original Sentence / NPP | Appeal Outcome | Revised Sentence / NPP | Key Insight |
Include at least 3 rows. Do NOT include cases with "[Surname]" or placeholder citations.

## 6. APPEAL OUTLOOK
Overall appellate viability assessment: Arguable — Strong / Arguable — Moderate / Requires Development.
2-3 paragraphs explaining the reasoning, strongest pathway to relief, and main risk factors. Do NOT use percentages.

## 7. PLAIN ENGLISH GUIDE
Explain the case and appeal in clear, plain English for a non-lawyer. CRITICAL: Use ONLY third-person language ("the applicant", "the legal professional"). ABSOLUTE BAN on "we", "us", "our", "you", "your".

IMPORTANT:
- No cost estimates or funding discussion.
- No witness contradiction section.
- Be specific to the supplied material — do not write generic legal advice.
- Do NOT use percentage success rates anywhere. Use appellate viability language.
- Do NOT include "Similar Cases (AI-Suggested)" with unverified citations. If citing cases, use REAL citations only.

GROUNDS TO COVER (MUST INCLUDE ALL):
{grounds_enumerated}

MATERIAL COUNTS:
- Total documents analysed: {len(documents)}
- Total timeline events: {len(timeline)}
- Total grounds identified: {len(grounds)}"""

    elif report_type == "full_detailed":
        system_prompt = f"""{base_system}
{report_guardrails}
You are generating a PAID Full Detailed Report ($150 AUD). This report is the PRIMARY product the client is paying for.

CRITICAL RULES FOR THIS REPORT:
1. The client ALREADY has a FREE Quick Summary. If ANY section in this report resembles the Quick Summary, the product is WORTHLESS.
2. Every section must be 3x MORE detailed than the equivalent free report section.
3. Every paragraph must cite SPECIFIC case facts: names (Joshua Homann, Kirralee Paepaerei, Justice McCallum), dates (21 September 2015, 25 May 2018), section numbers (s.23A Crimes Act 1900 (NSW)), and document names.
4. NEVER write generic legal explanations. WRONG: "The appeal process involves reviewing the original decision." RIGHT: "Homann's appeal to the NSWCCA challenges the 30-year sentence imposed by Justice McCallum on 25 May 2018 for the murder of Kirralee Paepaerei, arguing that the non-parole period of 22 years and 6 months is manifestly excessive when compared with..."
5. EVERY ground listed in GROUNDS TO COVER MUST be covered in FULL in Sections 4, 7, 11, and 15. If there are 5 grounds, write 5 full analyses. NO OMISSIONS.
6. Write in FLOWING PARAGRAPHS, not bullet points. Bullet points are ONLY acceptable inside tables or checklists.
7. Each section heading must be followed by AT LEAST 4-6 detailed paragraphs of substantive analysis.

Include forensic appellate strategy, professional courtroom framing, and plain-English action notes. Include hyperlinks to AustLII legislation, case databases, and court forms where the URL is known and verifiable. Do NOT fabricate URLs — if the exact AustLII path is not known, reference the legislation by name and section number only.
CRITICAL: NEVER use placeholder text in parentheses like '(Entries will develop...)'. Every section MUST have REAL, SUBSTANTIVE CONTENT with actual legal analysis."""
        user_prompt = f"""Create a FULL DETAILED LEGAL ANALYSIS REPORT for this {category_name.lower()} appeal case.

{case_context}

GROUNDS TO COVER — YOU MUST INCLUDE EVERY SINGLE ONE (if 5 grounds, write 5 full analyses):
{grounds_enumerated}

MATERIAL COUNTS (use these exact numbers in the report):
- Total documents analysed: {len(documents)}
- Total timeline events: {len(timeline)}
- Total grounds identified: {len(grounds)}

TARGET: 15,000-20,000 words minimum. This is a $150 AUD paid product. It must be AT LEAST 3x the depth of the free Quick Summary in EVERY section.

ANTI-LAZINESS RULES — VIOLATIONS MAKE THE REPORT WORTHLESS:
1. Section 1 (Executive Brief): MUST be 4-6 paragraphs with strategic assessment, NOT just a case snapshot rehash. Name the strongest ground, the weakest ground, the most likely outcome, and the recommended immediate action.
2. Section 2 (Forensic Chronology): MUST have 12+ dated events in FULL PARAGRAPH FORMAT (not bullet points). Each entry: date, what happened, source document, legal significance. Write 3-4 sentences per event minimum.
3. Section 3 (Document Evidence Digest): MUST analyse EVERY uploaded document individually with key extracts, reliability, and appellate relevance. One paragraph per document minimum.
4. Section 4 (Grounds Portfolio): MUST cover EVERY ground with 800+ words each. Include: legal threshold, case-specific facts, viability rating with reasoning, Crown response prediction (what will the prosecution argue?), defence rebuttal strategy, practical impact if ground succeeds. Write as flowing prose, NOT bullets.
5. Section 7 (Outcome Options): MUST write 300+ words for EACH outcome pathway (conviction quashed, retrial, downgrade, sentence reduction, appeal dismissed). Reference which specific grounds support each outcome.
6. Section 11 (How to Argue): MUST cover EVERY ground — if there are 5 grounds, write 5 complete argument strategies with lead proposition, supporting authority, prosecution answer, and rebuttal. 400+ words per ground.
7. Section 12 (Submissions Blueprint): Write ACTUAL draft submission paragraphs ready for court. Include specific argument sequences, authority placement, and time allocation. NOT bullet point summaries.
8. Section 15 (Client Brief): MUST cover EVERY ground and EVERY outcome in plain English. Explain what each ground means in simple terms and what happens if it succeeds. 1000+ words minimum for this section.

SECTION ORDERING: Analysis first, then Strategy, then Practical steps, then Client brief at the end.

Use this exact structure:

## 1. EXECUTIVE BRIEF
High-impact strategic summary: 4-6 paragraphs covering strongest grounds, jurisdiction posture, likely pathways to relief, urgency items, recommended strategy, and risk factors. Include case snapshot (defendant, offence, sentence, court, judge) and primary issues. This section must ADD strategic value beyond the free report.

MANDATORY: Start with: "This analysis is based on [X] documents and [Y] timeline events. [Z] grounds of appeal have been identified." Use EXACT counts.

## 2. FORENSIC CASE CHRONOLOGY
Full chronological reconstruction with 12+ dated entries. Each entry must be a FULL PARAGRAPH (3-4 sentences minimum) with: exact date, what occurred, which document/evidence supports it, and legal significance for the appeal. Write this as a narrative, NOT as bullet points.

## 3. DOCUMENT EVIDENCE DIGEST
For EACH document uploaded: title, key extracts (quote directly), reliability context, probative value, and specific appellate relevance. One detailed paragraph per document minimum.

## 4. GROUNDS OF MERIT — APPELLATE PATHWAY ANALYSIS
For EACH ground listed in GROUNDS TO COVER (ALL of them — no omissions):
Write as "Ground X: [Exact Title]" then 800+ words using the MANDATORY appellate structure:

**Appellate Pathway:** State the specific statutory provision engaged (e.g. "Miscarriage of justice under s 6(1) Criminal Appeal Act 1912 (NSW)")

**Ground:** Frame the arguable error forensically: "It is arguable that the trial judge erred in failing to..." or "It is contended that..." — NOT bare declarations.

**Error Identified:** What specifically is arguable as having gone wrong at trial. Reference case facts, dates, evidence. Use forensic language: "It is arguable that...", "It is contended that...".

**Materiality:** Why this error matters. How it affected the verdict or sentence.

**Consequence:** Legal consequence (verdict unsafe, miscarriage of justice, sentence set aside).

**Appellate Viability:**
- Outcome impact: Determinative / Influential / Minor
- Legal alignment: Direct authority / Analogous / Weak
- Evidence support: Strong / Partial / Limited

**Crown Response:** What the prosecution will argue.
**Defence Rebuttal:** Counter with specific authority.

Write as flowing prose. Minimum 800 words per ground.

## 5. COMPARATIVE SENTENCING TABLE (8+ CASES)
CRITICAL: Produce an ACTUAL populated markdown table with real case data — NEVER placeholder text like "The table will reference..." or "Details will be provided...".
Markdown table with 8+ comparable cases, then Detailed Outcome Analysis paragraph for each row.
| Case | Offence | Original Sentence / NPP | Appeal Outcome | Revised Sentence / NPP | Reduction (Years + %) | Key Reason |
Include AustLII search URL.

## 6. COMMON APPEAL GROUNDS FOR THIS OFFENCE TYPE
Markdown table with 8+ rows:
| Common Ground | Frequency in Comparable Appeals | Typical Success Pattern | Best Supporting Evidence |

## 7. OUTCOME OPTIONS AVAILABLE
First, markdown table. Then 300+ words for EACH pathway referencing specific grounds from Section 4:
- **Conviction quashed** — which grounds support this, legal threshold, likelihood
- **Retrial ordered** — implications, timeframes, changes
- **Conviction substituted/downgraded** — legal basis, sentence impact
- **Sentence reduced as manifestly excessive** — before/after with specific numbers
- **Appeal dismissed** — consequences, further options (High Court)

## 8. EVIDENTIARY GAPS + REMEDIATION CHECKLIST
Specific missing material with urgency priority (Critical/Important/Helpful) and exact remediation steps.

## 9. PRECEDENT OUTCOME MATRIX (10-12 CASES)
For each: citation, factual similarity, hearing outcome, extracted legal principle.

## 10. STATUTORY + DOCTRINAL FRAMEWORK MAP
APPLY each provision to THIS case — section number, Act name with year, specific relevance to Homann's appeal.

## 11. HOW TO ARGUE EACH TOP GROUND
For EACH ground (ALL of them — no omissions), write 400+ words covering:
- Lead proposition (core argument in 2 sentences)
- Supporting authority cluster (statute + 2-3 precedent anchors with citations)
- Expected prosecution answer (what will the Crown argue?)
- Rebuttal strategy with specific authority
- How establishing this ground leads to successful appeal outcome

## 12. SUBMISSIONS BLUEPRINT
Write ACTUAL draft submission text ready for court. Not bullet point summaries.
Written submission strategy: full argument sequence paragraphs, authority placement, framing of each ground.
Oral submission strategy: likely bench questions with response lines, time allocation per ground.

## 13. HOW TO START YOUR APPEAL + REQUIRED FORMS
Step-by-step guide specific to {state_info.get('name', 'NSW')} with forms table.

## 14. PRIORITISED ACTION PLAN
72-hour / 7-day / 28-day actions. Each action: what to do, who to contact, objective.

## 15. CLIENT PLAIN-ENGLISH BRIEF
1000+ words explaining in plain English: what the appeal is about, what EACH ground means, the viability of each ground (arguable/moderate/strong or requires development), what happens next, realistic outcomes, and what the applicant should do now. Cover EVERY ground individually. CRITICAL: Use ONLY third-person language. ABSOLUTE BAN on "you", "your", "we", "us", "our". Use "the applicant", "the legal professional". Do NOT use percentage success rates — use appellate viability language instead (arguable, moderate, strong, requires development).

IMPORTANT:
- No cost discussion. No witness contradiction section.
- Every section must have substantive content — no placeholders.
- Keep ALL outcomes within SECTION 7. Keep ALL timeframes within SECTION 14.
- NEVER truncate. Write the ACTUAL full content."""

    else:  # extensive_log
        system_prompt = f"""{base_system}
{report_guardrails}
You are generating the PREMIUM Extensive Report ($200 AUD) — the most detailed and case-specific legal analysis available.

The user has ALREADY received BOTH:
1. A FREE Quick Summary (case overview, top grounds preview, basic sentencing comparison)
2. A PAID Full Detailed Report (executive brief, forensic chronology, evidence digest, grounds portfolio with Crown/defence strategy, 8+ sentencing comparisons, outcome options matrix, evidence gaps checklist, 10-12 precedent cases, statutory framework, argument strategy, submissions blueprint, appeal steps guide, action plan, and plain-English brief)

THIS PREMIUM REPORT MUST BUILD ON — NOT REPEAT — THE FULL DETAILED REPORT. Do NOT re-state the same analysis. Instead:
- Where the Full Detailed analysed each ground with Crown response and rebuttal, THIS report must provide 1200+ word DEEP analysis per ground with fallback positions, additional authorities, and draft submission paragraphs.
- Where the Full Detailed provided 8+ sentencing comparisons, THIS report must provide 12+ with detailed paragraph analysis for EACH case.
- Where the Full Detailed had a precedent matrix of 10-12 cases, THIS report must have 15+ with specific factual comparisons.
- THIS report adds 5 ENTIRELY NEW sections not in the Full Detailed: Hearing Preparation Notes, Conference Preparation Pack, Court Pathway Operations Playbook, Similar Case Search Options, and Risk Assessment + Contingency Planning.
- Every shared section must go SIGNIFICANTLY deeper with fresh analysis, additional authorities, and more detailed strategy.

Every section must directly reference the supplied case material. Include hyperlinks to AustLII legislation, case databases, and court forms where the URL is known and verifiable. Do NOT fabricate URLs — if the exact AustLII path is not known, reference the legislation by name and section number only.
CRITICAL: NEVER use placeholder text. Every section MUST have REAL, SUBSTANTIVE, CASE-SPECIFIC CONTENT. Each ground analysis must be at least 1200 words. Reference specific documents, dates, and facts from the case throughout."""
        user_prompt = f"""Create a PREMIUM EXTENSIVE legal analysis report for this {category_name.lower()} appeal case. This must be the MOST COMPREHENSIVE report — significantly more detailed and case-specific than the Full Detailed tier.

{case_context}

GROUNDS TO COVER (MUST INCLUDE ALL — if 9 grounds, write 9 full analyses):
{grounds_enumerated}

MATERIAL COUNTS (use these exact numbers in the report):
- Total documents analysed: {len(documents)}
- Total timeline events: {len(timeline)}
- Total grounds identified: {len(grounds)}

Target range 25000-35000 words. Every section must reference specific facts, documents, and dates from this case.

CRITICAL — NO REPETITION FROM FULL DETAILED REPORT:
The user already has a Full Detailed Report covering grounds analysis, sentencing table (8 cases), outcome options, evidence gaps, precedent matrix (10-12 cases), statutory framework, argument strategy, submissions blueprint, appeal steps, and action plan. This Premium Extensive report must ADVANCE BEYOND all of that with deeper per-ground analysis (1200+ words each), expanded tables (12+ sentencing, 15+ precedents), and 5 ENTIRELY NEW sections: Hearing Preparation Notes, Conference Preparation Pack, Court Pathway Operations Playbook, Similar Case Search Options, and Risk Assessment + Contingency Planning. Do NOT copy or paraphrase content from the lower-tier reports.

SECTION ORDERING: Case-specific analysis first, then broader legal framework, then strategy, then practical steps, then client brief at the very end.

Use this exact structure:

## 1. EXECUTIVE BRIEF
Confident assessment of the appeal: strongest grounds with specific evidence anchors, jurisdiction posture, pathways to relief, urgency items, and a clear one-paragraph statement of the case. Include a case snapshot paragraph (defendant, offence, sentence, court, judge) and a short list of primary issues at the end.

MANDATORY: Start this section with a summary line: "This analysis is based on [X] documents and [Y] timeline events. [Z] grounds of appeal have been identified." Use the EXACT counts from the supplied case data.

## 2. FORENSIC CASE CHRONOLOGY
Comprehensive chronological reconstruction with at least 15 dated entries. For each entry:
- Date
- Event description (specific to the case facts)
- Source document/evidence
- Legal significance for the appeal
Include pre-trial, trial, sentencing, and post-sentencing events.

## 3. DOCUMENT EVIDENCE DIGEST
For EACH document/source in the case material:
- Key extracts (quote directly where possible)
- Reliability and credibility assessment
- Probative value for the appeal
- Which grounds this evidence supports
- Rating: Critical / Important / Supporting / Peripheral

## 4. GROUNDS OF MERIT — DEEP APPELLATE ANALYSIS
For EACH ground listed in GROUNDS TO COVER above (no omissions), provide a MINIMUM 1200-word analysis using the MANDATORY structure:

**Appellate Pathway:** The specific statutory provision engaged and the basis of the appeal.

**Trial Finding:** What the trial judge found or accepted on this issue.

**Error Identified:** Frame the arguable error using forensic language: "It is arguable that the trial judge erred in..." or "It is contended that...". Use forensic appellate language throughout — identify arguability, not bare declarations.

**Materiality:** Why this error matters to the outcome. How it affected the verdict or sentence.

**Consequence:** The legal consequence — verdict unsafe, miscarriage of justice, sentence must be set aside.

**Appellate Viability:**
- Outcome impact: Determinative / Influential / Minor
- Legal alignment: Direct authority / Analogous / Weak
- Evidence support: Strong / Partial / Limited

**Crown Response and Rebuttal:** Likely Crown prosecution response with detailed defence counter-strategy.

**Fallback Position:** Alternative argument if primary submission is rejected.

**Draft Submission Paragraph:** A paragraph ready to be adapted for written submissions.

Write each ground as a numbered entry starting with "Ground X: [Exact Title]" and then flowing paragraphs (no bullet-only answers). Minimum 1200 words per ground.

## 5. COMPARATIVE SENTENCING TABLE (12+ CASES)
CRITICAL: This section MUST contain an ACTUAL populated markdown table with real case data — NEVER placeholder text like "The table will reference..." or "Details will be provided...". If exact data cannot be found, use the best available comparable cases from Australian jurisprudence.

Markdown table with at least 12 comparable sentencing outcomes:
| Case | Offence | Original Sentence / NPP | Appeal Outcome | Revised Sentence / NPP | Reduction (Years + %) | Key Reason |
Include AustLII search link: [Search {state_info.get('appeal_court', 'NSWCCA')}]({state_info.get('cca_search_url', 'https://www.austlii.edu.au/cgi-bin/viewtoc/au/cases/nsw/NSWCCA/')})
After the table, provide a DETAILED paragraph for EACH case explaining:
- What was the original sentence and sentencing judge's reasoning
- What the appeal court decided and why
- How the reduction was achieved and which grounds succeeded
- How this compares to the current case

## 6. COMMON APPEAL GROUNDS FOR THIS OFFENCE TYPE
| Common Ground | Frequency in Comparable Appeals | Typical Success Pattern | Best Supporting Evidence |
Then provide detailed analysis of how each common ground applies or does not apply to THIS specific case.

## 7. OUTCOME OPTIONS — DETAILED PATHWAY ANALYSIS
Markdown table:
| Option | Legal Threshold | Likelihood in This Matter | Core Evidence Trigger | Practical Result |
Then provide DETAILED analysis (minimum 150 words each) for EVERY pathway (keep ALL outcomes within THIS section — do NOT create separate ## headings for each outcome):
- **Conviction quashed** — what standard must be met, what evidence supports this, what the defendant's position would be, which grounds support this outcome
- **Retrial ordered** — when this happens instead of quashing, what the retrial process involves, timeframes
- **Conviction substituted/downgraded** (e.g., murder to manslaughter) — legal basis, resulting sentence range, how this has worked in comparable cases
- **Sentence reduced as manifestly excessive** — show explicit before/after: Original sentence/NPP → Revised sentence/NPP with reasoning for reduction
- **Appeal dismissed** — consequences, options for special leave to the High Court, time limits

## 8. EVIDENTIARY GAPS + REMEDIATION CHECKLIST
For each gap:
- What is missing from the case file
- Why it matters for the appeal
- Exact steps to obtain it (who to contact, what to request, expected timeframe)
- Priority: Critical / Important / Helpful
- Impact on which grounds if not remediated

## 9. PRECEDENT OUTCOME MATRIX (15+ CASES)
For each of at least 15 cases:
- Full citation
- Factual similarity to this matter (specific comparison)
- Hearing outcome
- Extracted legal principle
- How this principle applies to the current case
Include AustLII link: [Search {state_info.get('appeal_court', 'NSWCCA')}]({state_info.get('cca_search_url', 'https://www.austlii.edu.au/cgi-bin/viewtoc/au/cases/nsw/NSWCCA/')})

## 10. STATUTORY + DOCTRINAL FRAMEWORK MAP
15+ statutory provisions with:
- Section number, Act name with year, jurisdiction
- AustLII link: [AustLII Legislation](https://www.austlii.edu.au/cgi-bin/viewdb/au/legis/)
- How this provision specifically applies to the current case
- Any recent amendments that affect the appeal

## 11. HOW TO ARGUE EACH TOP GROUND — DETAILED STRATEGY
For each priority ground (minimum 200 words each):
- Lead proposition (the core argument stated in 1-2 powerful lines)
- Supporting authority cluster (specific statute section + 3-4 precedent cases with citations)
- Expected Crown prosecution response (specific to this case)
- Detailed rebuttal strategy for the hearing
- Fallback argument if the primary submission is rejected
- **How establishing this ground leads to a specific appeal outcome** (what order to seek)

## 12. SUBMISSIONS BLUEPRINT
Written submission strategy (write as flowing paragraphs, NOT bullet lists):
Discuss the recommended argument sequence and why this ordering maximises impact. Explain the authority placement strategy for each ground. Describe how each ground should be framed in written submissions with key passages to quote from case material.

Oral submission strategy (write as flowing paragraphs, NOT bullet lists):
Discuss the likely bench questions for each ground with prepared responses. Cover time allocation per ground, opening and closing lines, and how to handle judicial scepticism on weaker grounds.

## 13. HEARING PREPARATION NOTES
For each ground, write detailed flowing paragraphs (NOT bullet lists or dot points):
Cover the key arguments and talking points for each ground, the anticipated questions from the bench with suggested responses, which authority to cite first and why, and any visual aids or demonstratives to prepare.
- Concessions to make strategically vs points to contest

## 14. CONFERENCE PREPARATION PACK
For briefing a barrister:
- One-page case summary (lead theory)
- Authorities shortlist with key passages highlighted
- Orders sought (specific orders to request from the court)
- Case strengths (with evidence references)
- Case weaknesses and mitigation strategy
- Client instructions summary
- Key dates and deadlines

## 15. COURT PATHWAY OPERATIONS PLAYBOOK
For EACH relevant court level ({state_info.get('name', 'NSW')} specific):
- Filing sequence and required documents
- Court registry details and contact information
- Deadlines for each filing step
- Extension-of-time route if deadlines have passed
- What happens at each stage of the process

## 16. HOW TO START YOUR APPEAL + REQUIRED FORMS
Step-by-step guide specific to {state_info.get('name', 'NSW')}:
1. Obtain trial transcripts and exhibits
2. Identify and finalise grounds of appeal
3. Lodge Notice of Intention to Appeal
4. Prepare detailed written submissions
5. Serve documents on the Crown/DPP
6. Attend the appeal hearing
For each step: plain English explanation, required form, deadline, and link.
Forms table:
| Form/Document | Purpose | Where to Obtain | Filing Deadline |
Links: [Legal Aid {state_info.get('name', 'NSW')}]({state_info.get('legal_aid_url', 'https://www.legalaid.nsw.gov.au/')}) | [AustLII](https://www.austlii.edu.au/) | [Court Forms]({state_info.get('court_forms_url', '#')})

## 17. SIMILAR CASE SEARCH OPTIONS
Tailored AustLII search guidance:
- 5+ query strings specifically designed for this case's offence and grounds profile
- Links: [Search {state_info.get('appeal_court', 'NSWCCA')}]({state_info.get('cca_search_url', 'https://www.austlii.edu.au/cgi-bin/viewtoc/au/cases/nsw/NSWCCA/')}) | [Search all states](https://www.austlii.edu.au/)
- Court-level filtering suggestions
- Keyword alternatives for each ground
- How to narrow results to the most relevant period and jurisdiction

## 18. PRIORITISED ACTION PLAN
72-hour actions (urgent — time-sensitive filings, evidence at risk of loss):
- Specific action, who to contact, deadline, objective

7-day actions (important — evidence gathering, legal research):
- Specific action, resources needed, dependencies

28-day actions (strategic — submission drafting, hearing preparation):
- Specific action, preparation steps, milestones

## 19. OPERATIONS ENGINE — CRITICAL MISSING EVIDENCE
Auto-generated checklist of evidence gaps identified from the case analysis:
For each gap:
- What is missing (e.g. "Psychiatric report addressing intoxication impact")
- Why it matters to the appeal
- How to obtain it
- Priority: Critical / Important / Desirable

## 20. EVIDENCE REQUEST LIST
Specific documents and materials that must be obtained:
| Item | Source | Purpose | Priority | Deadline |
Include: transcripts, expert reports, affidavits, medical records, CCTV, forensic reports.

## 21. SUGGESTED EXPERT REPORTS
For each identified evidentiary gap:
- Type of expert required (psychiatrist, forensic pathologist, etc.)
- What opinion is needed
- How the opinion would strengthen the appeal
- Suggested brief outline for instructing the expert

## 22. FILING PATHWAY
Step-by-step filing requirements for {state_info.get('name', 'NSW')}:
- Notice of Intention to Appeal (if not filed) — deadline and form
- Detailed grounds of appeal — content requirements
- Written submissions — structure and page limits
- Application for leave (if required) — criteria
- Supporting affidavit requirements
- Service requirements on DPP/Crown

## 23. RISK ASSESSMENT + CONTINGENCY PLANNING
For each ground:
- Appellate viability assessment (arguable/moderate/strong or requires development)
- Main risk factor
- Contingency if this ground fails
- Impact on overall appeal if this ground is excluded

Overall appeal risk assessment:
- Best case scenario
- Most likely outcome
- Worst case scenario and mitigation

Do NOT use percentage success rates. Use appellate viability language only.

## 24. CLIENT PLAIN-ENGLISH BRIEF
THIS MUST BE THE FINAL SECTION. Write this as if explaining the case to the defendant in everyday language, BUT STRICTLY IN THIRD PERSON:
- What the appeal is about and why it matters
- What are the strongest arguments in the applicant's favour (reference specific facts)
- What are the realistic prospects for the appeal (use viability language: arguable, moderate, strong — NOT percentages)
- What the different possible outcomes mean for the applicant personally
- What the applicant needs to do right now, this week, and this month
- What to expect at the hearing
- Honest assessment of risks alongside the opportunities
CRITICAL: This section is an educational tool. Use ONLY third-person language ("the applicant", "the legal professional", "this analysis"). NEVER use "we", "us", "our", "you", "your", or "them" when referring to the legal team or the applicant. WRONG: "The appeal is your opportunity to challenge your conviction." RIGHT: "The appeal represents an opportunity for the applicant to challenge the conviction."

IMPORTANT:
- Use markdown headings and tables exactly where specified.
- Working hyperlinks to AustLII and court websites throughout.
- Every section MUST reference specific facts from the supplied case material — no generic legal advice.
- This report must be VISIBLY more comprehensive than the Full Detailed tier.
- Australian English throughout.
- No cost discussion. No witness contradiction section.
- Quote directly from supplied documents where relevant.
- Every conclusion must be tied to case material or clearly marked as an assumption.
- NEVER use continuation markers like "... (continue with further analysis)", "... (continue analysis of other cases)", or similar truncation. Write the ACTUAL content for every section.
- Do NOT insert meta-commentary about the document itself (e.g., "This truncated document provides..."). Only output the report content.
- Keep ALL outcome pathways (conviction quashed, retrial, etc.) within SECTION 7 — do NOT create separate ## headings for individual outcomes.
- Keep ALL action items (72-hour, 7-day, 28-day) within SECTION 18 — do NOT create separate ## headings for each timeframe."""

    if aggressive_mode:
        aggressive_directive = """

AGGRESSIVE ADVOCACY MODE (USER-REQUESTED) — THIS FUNDAMENTALLY CHANGES YOUR APPROACH:

TONE SHIFT — You are no longer a cautious analyst. You are a senior criminal appeal barrister preparing to ARGUE this case in court. Write as if you are personally invested in winning this appeal:
- In aggressive mode, stronger forensic language is permitted: "This ground is compelling", "The Crown's position is untenable", "It is forcefully contended that the sentencing judge erred"
- NEVER hedge with "may", "could potentially", "it is possible that". Instead: "The evidence establishes", "It is contended that this constitutes a clear error", "There is a compelling argument that the conviction cannot stand"
- Frame EVERY ground as an argument TO BE WON, not a possibility to be explored
- Attack prosecution weaknesses directly: "The Crown's reliance on [X] is fatally undermined by [Y]"
- Draft ACTUAL submission paragraphs that could be read to the bench word-for-word
- For each ground, write the opening line you would say to the Court of Appeal judges
- Assume a standard (non-aggressive) version already exists. This aggressive report must add materially NEW arguments, authorities, and case-specific detail — do NOT rephrase or recycle the standard report.
- If a point is already made earlier in this report, deepen it with NEW evidence, dates, or authority rather than repeating it.
- This aggressive report must be approximately DOUBLE the depth of the standard report. Do not compress or summarise.

EXPANDED SCOPE:
1. DOUBLE the word count target for the entire report.
2. For EVERY ground of appeal, provide:
   - The STRONGEST possible legal argument as if arguing before the bench
   - Minimum 3 specific case citations that directly support this ground
   - A draft submission paragraph ready to be read in court
   - The specific orders to seek if this ground succeeds
   - Fallback position with alternative argument if primary is challenged
3. COMPARATIVE SENTENCING TABLE: 15+ cases minimum. For each, explain HOW the reduction was achieved.
4. OUTCOME OPTIONS: 250+ words per pathway. Reference ALL identified grounds for each pathway.
5. SUBMISSIONS: Write FULL draft argument paragraphs, not outlines. Include opening, authority chains, and closing for each ground.
6. PRECEDENT MATRIX: 20+ cases with detailed factual comparison.
7. Every conclusion must state the SPECIFIC order sought from the court.
8. Write as if the appeal WILL succeed — identify the path to victory, not just the obstacles.
"""
        system_prompt = f"{system_prompt}\n{aggressive_directive}"
        user_prompt = f"""{user_prompt}

AGGRESSIVE ADVOCACY MODE IS ON. Write as a senior barrister who believes in this appeal.
- DOUBLE the word count target.
- Use confident, forensic advocacy language throughout — frame as arguable, contended, tenable — never bare declarations, and never hedging with "may have" or "could potentially".
- Draft actual submission paragraphs for each ground that could be read to the bench.
- Frame the analysis as building the STRONGEST possible case for the appellant.
- Every section must reference specific case facts and documents.
- Attack Crown weaknesses directly and confidently.
- Write the opening line you would use to address the Court of Appeal for each key argument."""

    # Call AI — HARDENED via call_llm_structured with enhanced report-generation timeouts
    api_key = os.environ.get('EMERGENT_LLM_KEY')
    if not api_key:
        raise HTTPException(status_code=500, detail="AI service not configured — EMERGENT_LLM_KEY missing")

    _report_generation_metadata = {"models_used": [], "fallback_used": False}

    async def _subprocess_llm(prompt_text):
        """DO_NOT_UNDO — Run LLM call with pass-level retry for 502/service errors.
        
        The 4 model fallbacks inside call_llm_structured handle per-call retries.
        This outer retry handles cases where ALL 4 models fail on the same pass
        (e.g., upstream proxy 502 storm). Waits 30-60s between outer retries.
        """
        call_timeout = 420 if report_type in ("full_detailed", "extensive_log") else 300
        max_pass_retries = 3

        for retry in range(max_pass_retries):
            response = await call_llm_structured(
                system_prompt=system_prompt,
                user_prompt=prompt_text,
                session_id=f"rpt_gen_{case_id}",
                task_type="report_generation",
                max_tokens=16384,
                timeout_seconds=call_timeout,
                require_json=False,
                validation_fn=None,
            )

            if response["ok"]:
                if response.get("model"):
                    _report_generation_metadata["models_used"].append(response["model"])
                if response.get("fallback_used"):
                    _report_generation_metadata["fallback_used"] = True
                return response["content"]

            # All 4 models failed — wait and retry the entire pass
            if retry < max_pass_retries - 1:
                backoff = 10 + retry * 10  # 10s, 20s
                logger.warning(f"All models failed for pass (attempt {retry+1}/{max_pass_retries}). Waiting {backoff}s before retry. Error: {response['error']}")
                await asyncio.sleep(backoff)

        raise Exception(f"All LLM attempts failed after {max_pass_retries} pass retries. Last error: {response['error']}")

    response = None
    last_error = None
    try:
        # DO_NOT_UNDO — Condensed prompt for multi-pass generation.
        # Full document text is only needed for Pass 1-2 (Executive Brief, Chronology, Document Digest).
        # Later passes (Grounds analysis, Sentencing, Strategy, etc.) need grounds data, case metadata,
        # and pipeline data but NOT 100k+ of raw document text. This eliminates 502 proxy timeouts
        # on large prompts WITHOUT changing report depth or content.
        condensed_prompt = f"""
=== CASE INFORMATION ===
Title: {case.get('title', 'N/A')}
Defendant: {case.get('defendant_name', 'N/A')}
Case Number: {case.get('case_number', 'N/A')}
Court: {case.get('court', 'N/A')}
Judge: {case.get('judge', 'N/A')}
Sentence: {case.get('sentence', 'N/A')}
Summary: {case.get('summary', 'N/A')}

{offence_context}

=== CASE DOCUMENTS ({len(documents)} files — full text was provided in earlier passes) ===
"""
        for d in documents:
            condensed_prompt += f"- {d.get('filename', 'Unknown')} ({d.get('document_type', 'unknown')})\n"
        condensed_prompt += "\n"

        if timeline:
            condensed_prompt += f"=== TIMELINE OF EVENTS ({len(timeline)} events) ===\n"
            for event in timeline[:20]:
                condensed_prompt += f"- {event.get('event_date', 'Unknown')}: [{event.get('event_type')}] {event.get('title')}\n"
            condensed_prompt += "\n"

        if grounds:
            condensed_prompt += f"=== IDENTIFIED GROUNDS OF MERIT ({len(grounds)} grounds) ===\n"
            for g in grounds:
                condensed_prompt += f"- [{g.get('ground_type')}] {g.get('title')} (Strength: {g.get('strength')})\n"
                desc = (g.get('description') or '')[:500]
                if desc:
                    condensed_prompt += f"  {desc}\n"
                analysis = (g.get('analysis') or '')[:500]
                if analysis:
                    condensed_prompt += f"  Analysis: {analysis}\n"
            condensed_prompt += "\n"

        if structured_context:
            facts = structured_context.get("facts", [])
            findings = structured_context.get("findings", [])
            verified = structured_context.get("verified_issues", [])
            condensed_prompt += "=== STRUCTURED PIPELINE DATA ===\n"
            if facts:
                condensed_prompt += "CASE FACTS:\n"
                for f in facts[:15]:
                    condensed_prompt += f"- {f}\n" if isinstance(f, str) else f"- {f}\n"
            if findings:
                condensed_prompt += "\nKEY FINDINGS:\n"
                for f in findings[:10]:
                    condensed_prompt += f"- {f}\n" if isinstance(f, str) else f"- {f}\n"
            if verified:
                condensed_prompt += "\nVERIFIED APPEAL ISSUES:\n"
                for vi in verified:
                    condensed_prompt += f"- {vi.get('title', 'Untitled')} ({vi.get('ground_type', 'unknown')})\n"
            condensed_prompt += "\n"

        logger.info(f"Report prompt sizes: full={len(user_prompt)} chars, condensed={len(condensed_prompt)} chars")

        if report_type == "full_detailed":
            # DO_NOT_UNDO — 8-pass generation is FINAL. Previous 5-pass produced only ~7,890 words.
            # Current 8-pass produces ~12,000+ words. NEVER reduce pass count.
            half_grounds = max(1, len(grounds) // 2)
            # DO_NOT_UNDO — grounds_enumerated is ONE line per ground (e.g. "1. Title").
            # Split by actual ground count, not by arbitrary line multiplier.
            all_ground_lines = grounds_enumerated.split('\n')
            first_grounds = all_ground_lines[:half_grounds]
            second_grounds = all_ground_lines[half_grounds:]
            first_grounds_text = '\n'.join(first_grounds) if first_grounds else grounds_enumerated
            second_grounds_text = '\n'.join(second_grounds) if second_grounds else ""
            
            passes = [
                ("PASS 1/8", f"""
NOW GENERATE ONLY SECTIONS 1-2. Write 3000+ WORDS for this pass. Every paragraph must reference specific case facts.

## 1. EXECUTIVE BRIEF (1200+ words)
Write 6 FULL paragraphs — each paragraph must be 150+ words. This is a COUNSEL BRIEFING, not a case summary rehash.
Think: "what does counsel need to know in 5 minutes?"

- Paragraph 1: "This analysis is based on {len(documents)} documents and {len(timeline)} timeline events. {len(grounds)} grounds of appeal have been identified." Then strategic overview of the appeal's overall appellate position. Do NOT use percentage success rates — use viability language (arguable, moderate, strong).
- Paragraph 2: Key appellate themes (e.g. psychiatric evidence conflict, intoxication + intent interaction, procedural fairness concerns). NOT a list of grounds — identify the 2-3 overarching legal themes.
- Paragraph 3: Priority grounds (strongest first) with specific evidence anchors and why they have the best appellate viability.
- Paragraph 4: Strategic risks — what the prosecution will exploit, evidence gaps, weak factual anchors.
- Paragraph 5: Recommended immediate actions with specific deadlines and who to contact.
- Paragraph 6: Overall appellate position summary — "There are arguable grounds capable of attracting appellate consideration, subject to refinement and evidentiary development" or equivalent honest assessment.

CRITICAL: Do NOT repeat content from the 3 underlying reports. This must be SYNTHESIS only. Use forensic appellate language: "It is arguable that...", "It is contended that...", "There is a tenable argument that..." — NOT bare declarations like "The judge erred" and NOT hedging like "may have" or "could potentially".

## 2. FORENSIC CASE CHRONOLOGY (1500+ words)
Write 15+ dated events as FULL PARAGRAPHS (4-5 sentences each). NOT bullet points. Each event:
"On [DATE], [WHAT HAPPENED]. This is evidenced by [SOURCE DOCUMENT]. The legal significance is [WHY IT MATTERS FOR APPEAL]. This event [IMPACT]."
Cover ALL of: offence date and circumstances, arrest, charges, bail, committal, psychiatric assessments, trial dates, key evidence, verdict, sentencing, appeal filing.

STOP after section 2. Write every paragraph in full."""),

                ("PASS 2/8", f"""
NOW GENERATE ONLY SECTION 3. Write 2000+ WORDS for this pass. Analyse EVERY single uploaded document.

## 3. DOCUMENT EVIDENCE DIGEST (2000+ words)
For EACH of the {len(documents)} uploaded documents, write a FULL PARAGRAPH (150+ words each) covering:
- Document title and type
- Key extracts — quote directly from the content where possible
- Reliability and credibility assessment
- Probative value — what does this document prove or disprove?
- Which specific grounds of appeal does this document support?
- Rating: Critical / Important / Supporting / Peripheral

If there are {len(documents)} documents, write {len(documents)} detailed paragraphs minimum.

STOP after section 3."""),

                ("PASS 3/8", f"""
NOW GENERATE ONLY SECTION 4 (FIRST HALF OF GROUNDS). Write 4000+ WORDS for this pass.

## 4. GROUNDS OF MERIT — APPELLATE PATHWAY ANALYSIS (Part 1)
Write detailed appellate analyses for the FIRST {half_grounds} grounds. Each ground MUST follow this structure:

GROUNDS TO COVER IN THIS PASS:
{first_grounds_text}

For EACH ground, write as "Ground X: [Exact Title]" then 800+ words using the MANDATORY structure:
1. **Appellate Pathway:** The specific statutory provision engaged (e.g. "Miscarriage of justice under the Criminal Appeal Act")
2. **Ground:** Frame the arguable error — "It is arguable that the trial judge erred in failing to..." or "It is contended that..."
3. **Error Identified:** What is arguable as having gone wrong. Reference documents, dates, witnesses by name. Use forensic language: "It is arguable that...", "It is contended that...".
4. **Materiality:** Why this error matters. How it affected the verdict or sentence.
5. **Consequence:** Legal consequence (verdict unsafe, miscarriage of justice, sentence set aside).
6. **Appellate Viability:** Outcome impact (Determinative/Influential/Minor), Legal alignment (Direct/Analogous/Weak), Evidence support (Strong/Partial/Limited).
7. **Crown Response:** What the prosecution will argue (4-5 sentences). Reference authority ONLY if a real, verifiable citation is known — otherwise argue on principle without citing a specific case.
8. **Defence Rebuttal:** How to counter the Crown's likely position (4-5 sentences). Reference authority ONLY if a real, verifiable citation is known — otherwise frame the rebuttal on legal principle.

Write 800+ words per ground. Do NOT compress into bullet points. Use forensic appellate language throughout — frame as arguable, not declarations.

STOP after covering the first {half_grounds} grounds."""),

                ("PASS 4/8", f"""
NOW GENERATE ONLY SECTION 4 (SECOND HALF OF GROUNDS) AND SECTION 5. Write 4000+ WORDS for this pass.

## 4. GROUNDS OF MERIT — APPELLATE PATHWAY ANALYSIS (Part 2)
Continue with the REMAINING grounds. Each must follow the same mandatory structure as Part 1.

GROUNDS TO COVER IN THIS PASS:
{second_grounds_text if second_grounds_text.strip() else grounds_enumerated}

For EACH ground, use the same structure: Appellate Pathway → Ground → Error Identified → Materiality → Consequence → Appellate Viability → Crown Response → Defence Rebuttal.

## 5. COMPARATIVE SENTENCING TABLE (8+ CASES)
PRODUCE AN ACTUAL POPULATED MARKDOWN TABLE with case data. Then write a Detailed Outcome Analysis paragraph for EACH case in the table.
| Case | Offence | Original Sentence / NPP | Appeal Outcome | Revised Sentence / NPP | Reduction | Key Reason |
CRITICAL: Only include cases with REAL, VERIFIABLE citations (e.g. "R v Smith [2015] NSWCCA 123"). If fewer than 8 verified cases are known for this offence type, include only those that ARE known — do NOT fabricate citations to reach the target count. It is better to have 4 verified cases than 8 fabricated ones.

STOP after section 5."""),

                ("PASS 5/8", """
NOW GENERATE ONLY SECTIONS 6-7. Write 3000+ WORDS for this pass.

## 6. COMMON APPEAL GROUNDS FOR THIS OFFENCE TYPE
Markdown table with 8+ rows:
| Common Ground | Frequency in Comparable Appeals | Typical Success Pattern | Best Supporting Evidence |
Then for each common ground, explain how it applies or does not apply to THIS specific case.

## 7. OUTCOME OPTIONS AVAILABLE
Summary table first, then write 400+ WORDS for EACH of these 5 pathways:
- **Conviction quashed** — which grounds support this, legal threshold, appellate viability assessment (300+ words)
- **Retrial ordered** — triggers, what changes, timeframes (200+ words)
- **Conviction substituted/downgraded** — legal basis, sentence impact (200+ words)
- **Sentence reduced as manifestly excessive** — show exact before/after numbers with reasoning. State the appellate threshold: "Intervention required where the sentence is manifestly excessive, OR relevant mitigating factors were not properly weighed." (300+ words)
- **Appeal dismissed** — consequences, High Court special leave options (200+ words)
Do NOT use percentage success rates. Use appellate viability language.

STOP after section 7."""),

                ("PASS 6/8", """
NOW GENERATE ONLY SECTIONS 8-10. Write 3000+ WORDS for this pass.

## 8. OPERATIONS ENGINE — EVIDENTIARY GAPS + REMEDIATION (800+ words)
List 8+ specific gaps with urgency ratings. For each:
- What's missing (e.g. "Psychiatric report addressing intoxication impact")
- Why it matters and which ground it affects
- How to obtain it (specific steps, who to contact)
- Priority: Critical / Important / Desirable
- Impact on appeal if not obtained

Include: evidence request list, missing affidavit prompts, suggested expert reports.

## 9. PRECEDENT OUTCOME MATRIX (10-12 CASES)
For each case write a FULL PARAGRAPH: citation, factual similarity, outcome, extracted principle, how it applies to THIS case.
Do NOT include cases with "[Surname]" or unverified citations.

## 10. STATUTORY + DOCTRINAL FRAMEWORK MAP (1000+ words)
For EACH relevant provision write a paragraph APPLYING it to THIS case — section + Act + year + jurisdiction + specific relevance. Not generic descriptions.
Law sections MUST have ACTUAL section numbers. If the section number is not known, do NOT include that entry.

STOP after section 10."""),

                ("PASS 7/8", f"""
NOW GENERATE ONLY SECTIONS 11-12. Write 3000+ WORDS for this pass.

## 11. COUNSEL BRIEFING NOTE — HOW TO ARGUE EACH GROUND
Think: "what does counsel need in a 5-minute briefing on each ground?"
You MUST cover EVERY ground — ALL {len(grounds)} of them. For EACH ground write 400+ words:
GROUNDS TO COVER:
{grounds_enumerated}

For each ground — NO repetition of content from underlying reports. SYNTHESIS ONLY:
- **Lead Proposition**: Core argument in 2 powerful forensically-framed sentences ("It is arguable that...", "It is contended that...")
- **Appellate Pathway**: Which statutory provision is engaged
- **Supporting Authority Cluster**: Statute + 3 precedent cases (REAL citations only)
- **Expected Prosecution Answer**: 4-5 sentences with specific authorities
- **Rebuttal Strategy**: 4-5 sentences with counter-authorities
- **Strategic Risk**: Main weakness or gap
- **If Established**: Specific court order to seek

## 12. SUBMISSIONS BLUEPRINT (1200+ words)
Write ACTUAL DRAFT SUBMISSION TEXT ready for court.
**Written Submissions**: Full argument sequence paragraphs with authority placement. Write the opening paragraph and closing paragraph.
**Oral Submissions**: Bench questions with prepared responses. Time allocation per ground. Opening and closing strategies.

STOP after section 12."""),

                ("PASS 8/8", f"""
NOW GENERATE ONLY SECTIONS 13-15. Write 4000+ WORDS for this pass. Section 15 is critical.

## 13. FILING PATHWAY + REQUIRED FORMS (800+ words)
Step-by-step filing requirements:
- Notice of Intention to Appeal — deadline and form
- Detailed grounds of appeal — content requirements
- Written submissions — structure and page limits
- Application for leave (if required) — criteria
- Supporting affidavit requirements
- Service requirements on DPP/Crown
Each step: plain English explanation, required form, deadline.

## 14. PRIORITISED ACTION PLAN (800+ words)
72-hour actions (at least 5): exact action, who to contact, deadline, objective.
7-day actions (at least 5): exact action, resources needed, dependencies.
28-day actions (at least 5): exact action, preparation steps, milestones.

## 15. PLAIN-ENGLISH BRIEF (2000+ words)
For EACH of the {len(grounds)} grounds individually:
- What this ground means in simple terms (3-4 sentences)
- Why it matters (2-3 sentences)
- Appellate viability (honest assessment — arguable/moderate/strong, NOT percentages)
- What happens if it succeeds (specific outcome)

Then cover:
- The overall appeal: what, why, timeline
- Overall position: "There are arguable grounds capable of attracting appellate consideration, subject to refinement and evidentiary development." (or equivalent honest assessment)
- Each possible outcome and what it means
- What the applicant should do right now, this week, this month
- ABSOLUTE BAN: NEVER use "we", "us", "our", "you", "your". Use "the applicant", "the legal professional".

Do NOT truncate. Write ALL content."""),
            ]
            
            parts = []
            resume_from = 0
            if report_id:
                existing_report = await db.reports.find_one(
                    {"report_id": report_id},
                    {"_id": 0, "content.partial_analysis": 1, "content.passes_completed": 1},
                )
                existing_content = (existing_report or {}).get("content") or {}
                existing_partial = existing_content.get("partial_analysis") or ""
                existing_passes_completed = int(existing_content.get("passes_completed") or 0)
                if existing_partial and existing_passes_completed > 0:
                    parts = [existing_partial]
                    resume_from = min(existing_passes_completed, len(passes))
                    logger.info(f"Resuming full_detailed report {report_id} from pass {resume_from + 1}")

            # DO_NOT_UNDO — Same condensed context optimisation for full_detailed.
            # Pass 1 (Exec Brief + Chronology) and Pass 2 (Document Digest) get full context.
            # Passes 3-8 (Grounds analysis, Sentencing, Strategy) get condensed context.
            for pass_index, (label, instruction) in enumerate(passes[resume_from:], start=resume_from + 1):
                base_prompt = user_prompt if pass_index <= 2 else condensed_prompt
                pass_prompt = base_prompt + instruction
                logger.info(f"Full detailed {label} prompt size: system={len(system_prompt)}, user={len(pass_prompt)}")
                part = await _subprocess_llm(pass_prompt)
                parts.append(part)
                logger.info(f"Full detailed {label} response: {len(part)} chars")
                # Partial save after each pass to prevent data loss on server restart
                if report_id:
                    partial_content = "\n\n".join(parts)
                    await db.reports.update_one(
                        {"report_id": report_id},
                        {"$set": {"content.analysis": partial_content, "content.partial": True, "content.partial_analysis": partial_content, "content.passes_completed": pass_index}}
                    )
            
            response = "\n\n".join(parts)
            logger.info(f"Full detailed combined: {len(response)} chars")

        elif report_type == "extensive_log":
            # Ten-pass generation for extensive_log (24 sections)
            passes = [
                ("PASS 1/10", f"""

NOW GENERATE ONLY SECTIONS 1-3. Write 5000+ WORDS for this pass. This is a $200 PREMIUM report — every paragraph must be packed with case-specific facts, not generic legal education.

## 1. EXECUTIVE BRIEF (1200+ words)
Write 6-8 FULL paragraphs (NOT bullet points):
- Paragraph 1: Document/timeline/grounds counts from supplied data, then strategic overview of the appeal's overall appellate position (do NOT use percentage success rates — use viability language: arguable, moderate, strong)
- Paragraph 2: The 2-3 STRONGEST grounds with specific evidence anchors and legal tests that support them
- Paragraph 3: The weakest ground and why it's still worth pursuing (or should be abandoned)
- Paragraph 4: Jurisdiction-specific posture — what the {state_info.get('appeal_court', 'NSWCCA')} typically does with this type of appeal
- Paragraph 5: Most likely outcome pathway and realistic assessment of relief
- Paragraph 6: Key risks the prosecution will exploit and how to counter them
- Paragraph 7: Immediate actions required with specific deadlines
- Paragraph 8: Summary of 8+ primary issues identified with document references

CRITICAL: Do NOT use percentage probabilities or success rates anywhere in this report. Use appellate viability language only (arguable, moderate, strong, requires development). Use forensic appellate language throughout: "It is arguable that...", "It is contended that...", "There is a tenable argument that..." — NOT bare declarations like "The judge erred" and NOT hedging like "may have" or "could potentially".

## 2. FORENSIC CASE CHRONOLOGY (1500+ words)
Write 18+ dated events as FULL PARAGRAPHS (4-5 sentences each). NOT bullet points. Each event:
"On [EXACT DATE], [WHAT HAPPENED in detail]. This is established by [SOURCE DOCUMENT with specific reference]. The legal significance for the appeal is [SIGNIFICANCE]. The prosecution's likely treatment of this event is [PROSECUTION VIEW], while the defence can argue [DEFENCE POSITION]."
Cover: offence date and circumstances, arrest, charges laid, bail, committal, psychiatric/forensic reports, plea, trial dates, key witness testimony, jury directions, verdict, sentencing submissions, sentence, appeal filing.

## 3. DOCUMENT EVIDENCE DIGEST (1200+ words)
For EACH of the {len(documents)} uploaded documents, write 2-3 FULL PARAGRAPHS analysing:
- Document title and date
- Key extracts (QUOTE directly from the document content where possible)
- Reliability and credibility assessment
- Probative value — what it proves or disproves
- Which specific grounds this document supports
- Rating: Critical / Important / Supporting / Peripheral
If there are 10 documents, write 20-30 paragraphs.

STOP after section 3. Write ALL content — do NOT truncate."""),
                ("PASS 2/10", f"""

NOW GENERATE ONLY SECTIONS 4-5. Write 6000+ WORDS for this pass. Section 4 is the HEART of the $200 report — each ground must be a mini-essay.

## 4. GROUNDS OF MERIT — DEEP ANALYSIS
You MUST write about EVERY ground listed below. If there are 5 grounds, write 5 complete analyses. NO OMISSIONS.
GROUNDS TO COVER:
{grounds_enumerated}

For EACH ground, write 1200+ words as flowing paragraphs (NOT bullet points) covering:
1. "Ground X: [Exact Title from above]" as the heading
2. LEGAL THRESHOLD: The specific test for this ground — cite the statutory provision (section + Act + year) and the leading case that established the test. Explain what must be proved.
3. FACTUAL BASIS: How THIS case's specific facts satisfy the test. Reference documents, dates, witness statements. Quote from case material where possible.
4. VIABILITY RATING: Strong / Moderate / Weak — with 5-6 sentences explaining WHY, citing comparable cases where this type of ground succeeded or failed.
5. CROWN RESPONSE PREDICTION: Write 4-5 sentences predicting EXACTLY what the prosecution will argue. Be specific — name the authorities they'll rely on.
6. DEFENCE REBUTTAL STRATEGY: Write 4-5 sentences with the specific counter-argument. Name authorities that trump the Crown's position.
7. FALLBACK POSITION: If the primary argument fails, what's the fallback? (2-3 sentences)
8. APPEAL IMPACT: If this ground succeeds — conviction quashed? Sentence reduced? Retrial? What specific court order should be sought?
9. KEY AUTHORITY: Name the single most important case, provide the citation, and explain in 3-4 sentences exactly how it applies to this ground in this case.

## 5. COMPARATIVE SENTENCING TABLE (12+ CASES)
CRITICAL: Produce an ACTUAL populated markdown table with real case data — NEVER placeholder text like "The table will reference..." or "Details will be provided...".
Markdown table with 12+ rows. After the table, write a DETAILED PARAGRAPH for EACH of the 12+ cases (200+ words each) explaining: original sentencing reasoning, appeal court's reasoning, how the reduction was achieved, which grounds succeeded, and how this specifically compares to the current case.

STOP after section 5."""),
                ("PASS 3/10", f"""

NOW GENERATE ONLY SECTIONS 6-8. Write 4000+ WORDS for this pass.

## 6. COMMON APPEAL GROUNDS FOR THIS OFFENCE TYPE
Markdown table with 10+ rows, then for EACH common ground write 100+ words explaining how it does or does not apply to THIS specific case. Reference the actual grounds identified in this case where they overlap.

## 7. OUTCOME OPTIONS — DETAILED PATHWAY ANALYSIS
First provide summary table, then write 400+ WORDS for EACH of these 5 pathways (ALL within this one section):
- **Conviction quashed entirely**: Which of the {len(grounds)} grounds support this? What's the legal standard (e.g., miscarriage of justice)? What evidence is strongest? What is the appellate viability of this pathway?
- **Retrial ordered**: What triggers a retrial? What changes? Timeframes? Risks? What happens if the same evidence is presented?
- **Conviction substituted/downgraded**: Could the charge be reduced? Under what legal basis? What would the new sentence range be? Which grounds support this?
- **Sentence reduced as manifestly excessive**: Show EXACT before/after — Original sentence/NPP → realistic revised sentence/NPP with reasoning. Which sentencing comparisons from Section 5 support this?
- **Appeal dismissed**: What happens? Consequences for the applicant? Special leave to High Court — threshold, timeframe, realistic prospects?

## 8. EVIDENTIARY GAPS + REMEDIATION CHECKLIST (800+ words)
List 10+ specific gaps. For each: what's missing, why it matters, exact steps to obtain it (who to contact, what to request, expected timeframe), priority (Critical/Important/Helpful), impact on which grounds if not remediated.

STOP after section 8."""),
                ("PASS 4/10", """

NOW GENERATE ONLY SECTIONS 9-10. Write 3200+ WORDS for this pass.

## 9. PRECEDENT OUTCOME MATRIX (12+ CASES)
For each of 12+ cases, write a FULL PARAGRAPH (not just a table row):
- Full citation
- Factual similarity to THIS matter (be specific — offence type, relationship, circumstances)
- Hearing outcome
- Extracted legal principle
- How this principle applies to the current case (3-4 sentences)

## 10. STATUTORY + DOCTRINAL FRAMEWORK MAP (1000+ words)
12+ statutory provisions. For EACH provision, write a FULL PARAGRAPH:
- Section number, Act name with year, jurisdiction
- What the provision covers (1 sentence)
- How it SPECIFICALLY APPLIES to THIS case (3-4 sentences) — name the defendant, the offence, the ground it relates to
- Any recent amendments or judicial interpretation that affects the appeal

STOP after section 10."""),
                ("PASS 5/10", f"""

NOW GENERATE ONLY SECTION 11. Write 2800+ WORDS for this pass. This pass exists so EVERY ground can be fully argued without truncation.

## 11. HOW TO ARGUE EACH TOP GROUND — DETAILED STRATEGY
MUST cover ALL {len(grounds)} grounds. For EACH ground write 500+ words:
GROUNDS TO COVER:
{grounds_enumerated}

For each ground:
- **Lead Proposition**: The core argument in 2 powerful sentences
- **Supporting Authority Cluster**: Specific statute + 3-4 precedent cases with full citations and 2-3 sentences explaining each
- **Expected Crown Response**: 4-5 sentences predicting what the prosecution will argue
- **Rebuttal Strategy**: 4-5 sentences with specific counter-authorities
- **Fallback Position**: If the primary argument is rejected, what's the alternative? (3-4 sentences)
- **If Established**: What specific court order should be sought?

STOP after section 11."""),
                ("PASS 6/10", f"""

NOW GENERATE ONLY SECTIONS 12-14. Write 5000+ WORDS for this pass. These are the sections that make this $200 report UNIQUE. Sections 13 and 14 DO NOT exist in the $150 report.

## 12. SUBMISSIONS BLUEPRINT (1500+ words)
**Written Submission Strategy**: Write ACTUAL DRAFT PARAGRAPHS that could be filed with the court. For each major ground, write 2-3 paragraphs of draft submission text. Include argument sequence, authority placement, and framing. Write the opening paragraph and closing paragraph of the written submissions.

**Oral Submission Strategy** (write as flowing paragraphs, NOT bullet lists): For each ground, discuss the opening line to say to the bench, the likely bench questions with prepared response lines, time allocation recommendations, and how to handle judicial scepticism on weaker grounds. Write this as continuous prose.

## 13. HEARING PREPARATION NOTES (1200+ words — NEW SECTION NOT IN $150 REPORT)
For EACH of the {len(grounds)} grounds, write detailed flowing paragraphs (NOT bullet lists or dot points):
Cover the key arguments and talking points for each ground, the top 3 anticipated bench questions with suggested answers woven into the prose, which authority to cite first and why, concessions to make strategically vs points to contest, and any visual aids or demonstratives to prepare. Write as continuous analytical prose.

## 14. CONFERENCE PREPARATION PACK (1200+ words — NEW SECTION NOT IN $150 REPORT)
For briefing a barrister — write as an actual document:
- One-page case summary (write the ACTUAL summary — 300+ words covering lead theory, key facts, strongest grounds)
- Authorities shortlist (10+ cases with key passages identified)
- Orders sought (list specific court orders to request)
- Case strengths with evidence references (6+ points)
- Case weaknesses and mitigation strategy (4+ points with specific counter-arguments)
- Client instructions summary

STOP after section 14."""),
                ("PASS 7/10", f"""

NOW GENERATE ONLY SECTIONS 15-17. Write 4000+ WORDS for this pass. These sections are UNIQUE to the $200 report.

## 15. COURT PATHWAY OPERATIONS PLAYBOOK (1200+ words — NEW SECTION NOT IN $150 REPORT)
For EACH relevant court level in {state_info.get('name', 'NSW')}:
- Court name and jurisdiction
- Filing sequence with specific documents required at each stage
- Court registry details (address, phone, email where available)
- Deadlines for each filing step
- Extension-of-time route if deadlines have passed — specific provisions and what needs to be demonstrated
- What happens at each stage (directions hearing, callover, full hearing)
- Estimated timeframes from filing to hearing
- Costs considerations

## 16. HOW TO START YOUR APPEAL + REQUIRED FORMS (800+ words)
Step-by-step guide specific to {state_info.get('name', 'NSW')} with FULL DETAIL per step:
1. Obtain trial transcripts and exhibits — from which court registry, expected cost and timeframe
2. Identify and finalise grounds of appeal — how to narrow from identified grounds to filed grounds
3. Lodge Notice of Intention to Appeal — form name, deadline, where to lodge
4. Prepare detailed written submissions — structure, length, authorities
5. Serve documents on the Crown/DPP — who, where, method
6. Attend the appeal hearing — what to expect, who attends
Forms table: | Form/Document | Purpose | Where to Obtain | Filing Deadline |

## 17. SIMILAR CASE SEARCH OPTIONS (800+ words — NEW SECTION NOT IN $150 REPORT)
5+ tailored AustLII search queries specifically designed for this case:
For each query:
- The exact search string
- What it's designed to find
- Expected number of results and how to filter them
- Key cases to look for in results
- How to use the results to strengthen the appeal
Court-level filtering suggestions and keyword alternatives for each ground.

STOP after section 17."""),
                ("PASS 8/10", """

NOW GENERATE ONLY SECTIONS 18-19. Write 3000+ WORDS for this pass.

## 18. PRIORITISED ACTION PLAN (1000+ words)
72-hour actions (urgent — at least 6 specific actions):
For each: exact action, who to contact (name the office/registry), deadline, objective, what happens if missed.

7-day actions (important — at least 6 specific actions):
For each: exact action, resources needed, dependencies, expected outcome.

28-day actions (strategic — at least 6 specific actions):
For each: exact action, preparation steps, milestones, how this contributes to the appeal.

## 19. OPERATIONS ENGINE — CRITICAL MISSING EVIDENCE (1200+ words — NEW SECTION NOT IN $150 REPORT)
Auto-generated checklist of evidence gaps identified from the case analysis. For EACH gap, write 100+ words covering:
- What is missing (name the specific document, report, or testimony)
- Why it matters to the appeal — which specific ground(s) does it affect?
- Exact steps to obtain it (who to contact, what to request, expected timeframe)
- Priority: Critical / Important / Desirable
- Impact on appeal viability if NOT obtained

Identify at least 8 evidence gaps. This section is what differentiates the $200 report from the $150 report — it must be operationally specific, not generic.

STOP after section 19."""),
                ("PASS 9/10", f"""

NOW GENERATE ONLY SECTIONS 20-22. Write 3500+ WORDS for this pass. These sections are the OPERATIONS ENGINE that makes the $200 report unique.

## 20. EVIDENCE REQUEST LIST (800+ words — NEW SECTION NOT IN $150 REPORT)
Specific documents and materials that must be obtained. Produce an ACTUAL populated markdown table:
| Item | Source | Purpose | Priority | Deadline |
Include at least 10 rows. After the table, write a paragraph for each Critical-priority item explaining why it is essential.
Include: transcripts, expert reports, affidavits, medical records, CCTV, forensic reports, psychiatric assessments, sentencing submissions.

## 21. SUGGESTED EXPERT REPORTS (800+ words — NEW SECTION NOT IN $150 REPORT)
For each identified evidentiary gap that requires expert opinion, write 150+ words covering:
- Type of expert required (psychiatrist, forensic pathologist, toxicologist, etc.)
- What specific opinion is needed and what question should be put to the expert
- How the opinion would strengthen the appeal — which ground(s) does it support?
- Suggested brief outline for instructing the expert (3-4 key points to include in the brief)
Identify at least 4 expert reports needed.

## 22. FILING PATHWAY (800+ words — NEW SECTION NOT IN $150 REPORT)
Step-by-step filing requirements specific to {state_info.get('name', 'NSW')}:
- Notice of Intention to Appeal (if not filed) — deadline, form name, where to lodge
- Detailed grounds of appeal — content requirements, format, page limits
- Written submissions — structure, length, authorities format
- Application for leave (if required) — criteria, threshold, what must be demonstrated
- Supporting affidavit requirements — who swears, what content
- Service requirements on DPP/Crown — method, timing, proof of service
- Extension of time (if applicable) — what must be shown, typical outcomes
Produce a timeline table: | Step | Document | Deadline | Filed With | Served On |

STOP after section 22."""),
                ("PASS 10/10", f"""

NOW GENERATE ONLY SECTIONS 23-24. Write 5000+ WORDS for this pass. Section 24 is the CLIENT BRIEF — it must be thorough and cover EVERY ground.

## 23. RISK ASSESSMENT + CONTINGENCY PLANNING (1200+ words — NEW SECTION NOT IN $150 REPORT)
For EACH of the {len(grounds)} grounds, write 150+ words covering:
- Appellate viability assessment (arguable/moderate/strong — NOT percentages)
- Main risk factor (what could go wrong?)
- Contingency if this ground fails (what's the backup?)
- Impact on overall appeal if this ground is excluded

Overall appeal risk assessment (500+ words):
- Best case scenario, likelihood, and path to get there
- Most likely outcome with realistic assessment
- Worst case scenario and mitigation strategy
- How grounds interact — if Ground 1 fails, does Ground 3 become stronger?
- Whether grounds should be argued independently or as a package

## 24. CLIENT PLAIN-ENGLISH BRIEF (2000+ words)
THIS IS THE FINAL SECTION. Write in plain, everyday English that explains the case in THIRD PERSON.

For EACH of the {len(grounds)} grounds individually:
- What this ground means in simple terms (2-3 sentences)
- Why it matters for the appeal (2-3 sentences)
- What the chances are (honest assessment)
- What happens if it succeeds (specific outcome)

Then cover:
- The overall appeal: what it is, why it's happening, and the realistic timeline
- Each possible outcome and what it means personally for the applicant
- Exactly what the applicant should do right now, this week, and this month
- What to expect at the hearing — how long, who's there, what happens
- Honest assessment of risks alongside the opportunities
- ABSOLUTE BAN: NEVER use "we", "us", "our", "you", "your". Use "the applicant", "the legal professional", "this analysis". WRONG: "The appeal is your opportunity." RIGHT: "The appeal represents an opportunity for the applicant."

Do NOT truncate. Write ALL content for both sections."""),
            ]
            
            parts = []
            resume_from = 0
            if report_id:
                existing_report = await db.reports.find_one(
                    {"report_id": report_id},
                    {"_id": 0, "content.partial_analysis": 1, "content.passes_completed": 1},
                )
                existing_content = (existing_report or {}).get("content") or {}
                existing_partial = existing_content.get("partial_analysis") or ""
                existing_passes_completed = int(existing_content.get("passes_completed") or 0)
                if existing_partial and existing_passes_completed > 0:
                    parts = [existing_partial]
                    resume_from = min(existing_passes_completed, len(passes))
                    logger.info(f"Resuming extensive_log report {report_id} from pass {resume_from + 1}")

            for pass_index, (label, instruction) in enumerate(passes[resume_from:], start=resume_from + 1):
                # Pass 1 gets full document context; passes 2-8 get condensed context
                base_prompt = user_prompt if pass_index == 1 else condensed_prompt
                pass_prompt = base_prompt + instruction
                logger.info(f"Extensive log {label} prompt size: system={len(system_prompt)}, user={len(pass_prompt)}")
                part = await _subprocess_llm(pass_prompt)
                parts.append(part)
                logger.info(f"Extensive log {label} response: {len(part)} chars")
                # Save partial progress after each pass so server restarts don't lose work
                partial_response = "\n\n".join(parts)
                await db.reports.update_one(
                    {"report_id": report_id},
                    {"$set": {"content.partial_analysis": partial_response, "content.passes_completed": pass_index}}
                )
            
            response = "\n\n".join(parts)
            logger.info(f"Extensive log combined: {len(response)} chars")

        else:
            logger.info(f"Report prompt size: system={len(system_prompt)}, user={len(user_prompt)}, total={len(system_prompt)+len(user_prompt)}")
            response = await _subprocess_llm(user_prompt)
    except Exception as e:
        last_error = e
        logger.error(f"Report generation failed: {e}")
    
    if response is None:
        logger.error(f"All report generation attempts failed: {last_error}")
        raise HTTPException(status_code=503, detail=f"AI report generation failed after retries: {str(last_error)}")

    anchor_terms = _build_anchor_terms(case, documents, timeline, grounds)
    response = _dedupe_report_content(response, report_type, anchor_terms)

    # DO_NOT_UNDO — Missing section repair. The resume logic can produce reports with
    # missing sections if a pass's LLM output truncated. This step detects missing
    # sections and generates them in-place BEFORE the expansion step runs.
    if report_type in ("full_detailed", "extensive_log"):
        # Define expected main sections for each report type
        expected_sections_map = {
            "full_detailed": {
                "## 1.": "EXECUTIVE BRIEF",
                "## 2.": "FORENSIC CASE CHRONOLOGY",
                "## 3.": "DOCUMENT EVIDENCE DIGEST",
                "## 4.": "GROUNDS OF MERIT",
                "## 5.": "COMPARATIVE SENTENCING TABLE",
                "## 7.": "OUTCOME OPTIONS",
            },
            "extensive_log": {
                "## 1.": "EXECUTIVE BRIEF",
                "## 2.": "FORENSIC CASE CHRONOLOGY",
                "## 3.": "DOCUMENT EVIDENCE DIGEST",
                "## 4.": "GROUNDS OF MERIT",
                "## 5.": "COMPARATIVE SENTENCING TABLE",
                "## 9.": "PRECEDENT OUTCOME MATRIX",
                "## 11.": "HOW TO ARGUE EACH TOP GROUND",
                "## 18.": "PRIORITISED ACTION PLAN",
                "## 19.": "OPERATIONS ENGINE — CRITICAL MISSING EVIDENCE",
                "## 20.": "EVIDENCE REQUEST LIST",
                "## 21.": "SUGGESTED EXPERT REPORTS",
                "## 22.": "FILING PATHWAY",
                "## 23.": "RISK ASSESSMENT + CONTINGENCY PLANNING",
                "## 24.": "CLIENT PLAIN-ENGLISH BRIEF",
            },
        }
        expected = expected_sections_map.get(report_type, {})
        sections_in_report = _split_report_sections(response)

        for prefix, name in expected.items():
            # Check if this section exists in the report (match "## N." prefix)
            if not any(h.startswith(prefix.rstrip()) for h, _ in sections_in_report if h):
                logger.warning(f"Missing section detected: {prefix} {name}. Generating repair.")
                try:
                    doc_list = "\n".join([f"- {d.get('original_filename', d.get('filename', 'Unknown'))}" for d in documents[:20]])
                    repair_prompt = f"""Generate the MISSING section for a legal appeal report.

CASE: {case.get('title', 'Unknown')} ({case.get('state', 'NSW')})
SENTENCE: {case.get('sentence', 'Not specified')}
DOCUMENTS ({len(documents)}):
{doc_list}

GROUNDS:
{grounds_enumerated}

WRITE THIS SECTION:
{prefix} {name}

INSTRUCTIONS:
- Write 4000-6000 characters of substantive, case-specific content
- Reference specific case facts, dates, documents, and legal authorities
- Use flowing paragraphs, NOT bullet points
- Australian English, strict third-person only (NEVER "you", "your", "we", "us")
- Start the section with exactly: {prefix} {name}
- Do NOT include any other section headings
"""
                    repaired = await _subprocess_llm(repair_prompt)
                    if repaired and len(repaired.strip()) > 500:
                        # Find the insertion point — insert before the NEXT numbered section
                        section_num = int(prefix.replace("##", "").replace(".", "").strip())
                        insert_point = len(response)
                        for h, _ in sections_in_report:
                            if h:
                                try:
                                    h_num = int(re.search(r"(\d+)", h).group(1))
                                    if h_num > section_num:
                                        insert_point = response.find(h)
                                        break
                                except (ValueError, AttributeError):
                                    continue
                        # Clean up the repair output
                        cleaned_repair = repaired.strip()
                        if not cleaned_repair.startswith("##"):
                            cleaned_repair = f"{prefix} {name}\n\n{cleaned_repair}"
                        response = response[:insert_point] + "\n\n" + cleaned_repair + "\n\n" + response[insert_point:]
                        logger.info(f"Repaired missing section: {prefix} {name} ({len(cleaned_repair)} chars)")
                        # Re-split after insertion
                        sections_in_report = _split_report_sections(response)
                except Exception as exc:
                    logger.warning(f"Section repair failed for {prefix} {name}: {exc}")

    # DO_NOT_UNDO — Minimum character targets. NEVER lower these values.
    min_lengths = {
        "quick_summary": 14000,
        "full_detailed": 80000,
        "extensive_log": 150000
    }
    target_length = min_lengths.get(report_type, 12000)
    if aggressive_mode:
        target_length = int(target_length * 2.0)

    # Expand ANY report that falls below target — multi-pass reports can still be short
    if len(response) < target_length:
        logger.info(f"Report below target ({len(response)} < {target_length}), running expansion pass(es)")
        
        if report_type in ("full_detailed", "extensive_log"):
            # For multi-pass reports, expand the THINNEST sections individually to avoid 502 errors
            sections = _split_report_sections(response)
            # Sort sections by content length (shortest first) and expand the thin ones
            thin_sections = [(heading, content) for heading, content in sections if heading and len(content) < 3000]
            thin_sections.sort(key=lambda x: len(x[1]))
            
            for heading, content in thin_sections[:6]:  # Expand up to 6 thin sections
                try:
                    section_expand_prompt = f"""You are expanding one section of a legal report for a criminal appeal case.

CASE CONTEXT:
{case_context[:15000]}

SECTION TO EXPAND: {heading}
Current content ({len(content)} chars — this is TOO SHORT, must be 4000+ chars):

{content}

INSTRUCTIONS:
- Rewrite this ENTIRE section with MUCH MORE detail. Target 4000-6000 characters minimum.
- Reference specific case facts: names, dates, section numbers, case citations.
- Write in flowing paragraphs, NOT bullet points.
- Use Australian English and strict third-person language.
- Include statutory references with full citations (section + Act + year + jurisdiction).
- The section heading must remain exactly: {heading}
- Do NOT include any other section headings. Write ONLY this one section.
- Make the content case-specific — avoid generic legal advice.

GROUNDS COVERED IN THIS CASE:
{grounds_enumerated}
"""
                    expanded_section = await _subprocess_llm(section_expand_prompt)
                    if expanded_section and len(expanded_section) > len(content) * 1.3:
                        # Replace the thin section with the expanded version
                        # Find the section in the response and replace it
                        escaped_heading = re.escape(heading)
                        section_pattern = rf"({escaped_heading})\n([\s\S]*?)(?=\n## \d+\.|$)"
                        match = re.search(section_pattern, response)
                        if match:
                            # Extract just the new content (strip the heading if the LLM repeated it)
                            new_content = expanded_section
                            if new_content.strip().startswith(heading):
                                new_content = new_content.strip()[len(heading):].strip()
                            response = response[:match.start(2)] + new_content + response[match.end(2):]
                            logger.info(f"Section expanded: {heading} ({len(content)} → {len(new_content)} chars)")
                except Exception as exc:
                    logger.warning(f"Section expansion failed for {heading}: {exc}")
                    continue
        else:
            # For single-pass reports (quick_summary), try a full expansion
            expansion_attempts = 0
            while len(response) < target_length and expansion_attempts < 2:
                expansion_attempts += 1
                try:
                    expansion_prompt = f"""{case_context}

The report below is {len(response)} characters but must be at least {target_length} characters. EXPAND IT SUBSTANTIALLY. 

RULES:
- Keep ALL existing ## section headings exactly as-is. Do NOT rename or reorder sections.
- Do NOT remove or rewrite any existing text — only ADD new paragraphs.
- Add case-specific details: names, dates, section numbers, case citations, document references.
- Every added paragraph must reference specific facts from THIS case.
- Avoid repetition — add NEW analysis, additional authorities, deeper reasoning.
- Write in FLOWING PARAGRAPHS, not bullet points.
- Maintain Australian English and third-person language throughout.

REPORT TO EXPAND:
{response}
"""
                    expanded = await _subprocess_llm(expansion_prompt)
                    if expanded and len(expanded) > len(response) * 1.05:
                        response = expanded
                        logger.info(f"Expansion attempt {expansion_attempts}: {len(response)} chars")
                    else:
                        logger.warning(f"Expansion attempt {expansion_attempts} did not add enough content")
                        break
                except Exception as exc:
                    logger.warning(f"Report expansion attempt {expansion_attempts} failed: {exc}")
                    break

    response = _strip_report_placeholders(response)
    response = _sanitise_suspect_authorities(response)
    response = response.strip()

    # Preserve the actual grounds linked to the case so reports reflect the real ground count.
    grounds_of_merit = [
        {
            "ground_id": ground.get("ground_id"),
            "title": ground.get("title", "Untitled ground"),
            "description": ground.get("description", ""),
            "strength": ground.get("strength", "moderate"),
            "ground_type": ground.get("ground_type", "other"),
        }
        for ground in grounds
    ]
    
    if aggressive_mode:
        response += """

---

**IMPORTANT DISCLAIMER:** This report is for educational and research purposes only. It does not constitute legal advice. Always consult a qualified legal practitioner before taking any action.
"""

    # Build provenance metadata
    report_metadata = ReportMetadata(
        generated_by_model=_report_generation_metadata["models_used"][-1] if _report_generation_metadata["models_used"] else None,
        fallback_used=_report_generation_metadata["fallback_used"],
        documents_analyzed=len(documents),
        timeline_events_analyzed=len(timeline),
        grounds_considered=len(grounds),
        verification_status="draft",
        confidence_note="AI-generated output requiring legal review",
    )

    return {
        "analysis": response,
        "grounds_of_merit": grounds_of_merit,
        "case_data": case,
        "document_count": len(documents),
        "event_count": len(timeline),
        "metadata": report_metadata.model_dump(),
        "pipeline_metadata": {
            "status": pipeline_status.get("status") if pipeline_status else "not_checked",
            "staleness": pipeline_status.get("staleness") if pipeline_status else None,
            "auto_refreshed": pipeline_status.get("status") == "refreshed" if pipeline_status else False,
            "pipeline_argument_count": len(pipeline_arguments),
            "submission_draft_present": bool(submission_draft),
        },
    }


BARRISTER_SOURCE_TYPES = ["quick_summary", "full_detailed", "extensive_log"]
BARRISTER_GENERATION_TIMEOUT_MINUTES = 40


def _coerce_utc_datetime(value):
    if isinstance(value, datetime):
        return value if value.tzinfo else value.replace(tzinfo=timezone.utc)
    if isinstance(value, str) and value:
        try:
            parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
            return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)
        except ValueError:
            return None
    return None


def _generated_at_sort_value(report: dict) -> str:
    return str(report.get("generated_at") or "")


async def _get_latest_standard_reports(case_id: str, user_id: str) -> List[dict]:
    reports = await db.reports.find(
        {
            "case_id": case_id,
            "user_id": user_id,
            "report_type": {"$in": BARRISTER_SOURCE_TYPES},
            "status": "completed",
            "content.aggressive_mode": {"$ne": True},
        },
        {"_id": 0},
    ).sort("generated_at", -1).to_list(50)

    latest_by_type = {}
    for report in reports:
        report_type = report.get("report_type")
        if report_type not in latest_by_type:
            latest_by_type[report_type] = report

    missing_types = [report_type for report_type in BARRISTER_SOURCE_TYPES if report_type not in latest_by_type]
    if missing_types:
        missing_labels = ", ".join(missing_types)
        raise HTTPException(status_code=409, detail=f"Barrister View remains locked until all 3 standard reports are completed. Missing: {missing_labels}")

    selected = [latest_by_type[report_type] for report_type in BARRISTER_SOURCE_TYPES]
    selected.sort(key=_generated_at_sort_value, reverse=True)
    return selected


def _build_barrister_source_signature(reports: List[dict]) -> str:
    ordered = sorted(
        reports,
        key=lambda report: BARRISTER_SOURCE_TYPES.index(report.get("report_type")) if report.get("report_type") in BARRISTER_SOURCE_TYPES else 99,
    )
    return "|".join(
        f"{report.get('report_type')}::{report.get('report_id')}::{report.get('generated_at')}"
        for report in ordered
    )


def _build_barrister_report_source_text(reports: List[dict]) -> str:
    type_labels = {
        "quick_summary": "Quick Summary",
        "full_detailed": "Full Detailed Report",
        "extensive_log": "Extensive Log Report",
    }
    blocks = []
    per_report_limit = 100000
    for report in sorted(
        reports,
        key=lambda item: BARRISTER_SOURCE_TYPES.index(item.get("report_type")) if item.get("report_type") in BARRISTER_SOURCE_TYPES else 99,
    ):
        analysis = _strip_report_placeholders((report.get("content") or {}).get("analysis", ""))
        analysis = re.sub(r"\n{3,}", "\n\n", analysis).strip()
        blocks.append(
            f"===== {type_labels.get(report.get('report_type'), report.get('report_type', 'Report'))} =====\n"
            f"Report ID: {report.get('report_id')}\n"
            f"Generated: {report.get('generated_at')}\n\n"
            f"{analysis[:per_report_limit]}"
        )
    return "\n\n".join(blocks)


def _trim_text_preserving_ends(text: str, max_chars: int) -> str:
    if not text or len(text) <= max_chars:
        return text
    head_chars = int(max_chars * 0.65)
    tail_chars = max_chars - head_chars
    return (
        text[:head_chars].rstrip()
        + "\n\n[Additional detailed source material omitted here for prompt length control]\n\n"
        + text[-tail_chars:].lstrip()
    )


def _build_barrister_group_source_text(reports: List[dict], max_chars_by_type: dict) -> str:
    type_labels = {
        "quick_summary": "Quick Summary",
        "full_detailed": "Full Detailed Report",
        "extensive_log": "Extensive Log Report",
    }
    blocks = []
    for report in sorted(
        reports,
        key=lambda item: BARRISTER_SOURCE_TYPES.index(item.get("report_type")) if item.get("report_type") in BARRISTER_SOURCE_TYPES else 99,
    ):
        report_type = report.get("report_type")
        analysis = _strip_report_placeholders((report.get("content") or {}).get("analysis", ""))
        analysis = re.sub(r"\n{3,}", "\n\n", analysis).strip()
        limited_analysis = _trim_text_preserving_ends(analysis, max_chars_by_type.get(report_type, 18000))
        blocks.append(
            f"===== {type_labels.get(report_type, report_type or 'Report')} =====\n"
            f"Report ID: {report.get('report_id')}\n"
            f"Generated: {report.get('generated_at')}\n\n"
            f"{limited_analysis}"
        )
    return "\n\n".join(blocks)


def _dedupe_barrister_ground_subsections(text: str) -> str:
    section_match = re.search(r"(## Grounds of Merit\n)([\s\S]*?)(?=\n## Statutory Framework)", text)
    if not section_match:
        return text

    grounds_body = section_match.group(2).strip()
    block_pattern = re.compile(r"(### Ground \d+: [^\n]+\n[\s\S]*?)(?=(?:\n### Ground \d+: )|\Z)")
    blocks = [block.strip() for block in block_pattern.findall(grounds_body)]
    if not blocks:
        return text

    deduped_blocks = []
    seen_titles = set()
    for block in blocks:
        heading = block.split("\n", 1)[0].strip()
        title_match = re.match(r"^### Ground \d+: (.+)$", heading)
        title = (title_match.group(1).strip().lower() if title_match else heading.lower())
        if title in seen_titles:
            continue
        seen_titles.add(title)
        deduped_blocks.append(block)

    renumbered_blocks = []
    for index, block in enumerate(deduped_blocks, start=1):
        heading, remainder = (block.split("\n", 1) + [""])[:2]
        title_match = re.match(r"^### Ground \d+: (.+)$", heading.strip())
        title = title_match.group(1).strip() if title_match else heading.replace("### ", "").strip()
        rebuilt = f"### Ground {index}: {title}"
        if remainder.strip():
            rebuilt += f"\n{remainder.strip()}"
        renumbered_blocks.append(rebuilt.strip())

    replacement = "\n\n".join(renumbered_blocks)
    return text[:section_match.start(2)] + replacement + text[section_match.end(2):]


def _normalise_barrister_table_titles(text: str) -> str:
    replacements = {
        "### Table: Comparative Authorities": "### Comparative Authorities Table",
        "### Table: Sentencing Comparison": "### Sentencing Comparison Table",
        "### Table: Evidentiary Pressure Points": "### Evidentiary Pressure Points Table",
        "### Table: Relief Pathways": "### Relief Pathways Matrix",
    }
    for source, target in replacements.items():
        text = text.replace(source, target)
    return text


# DO NOT UNDO — Barrister View multi-pass generation engine
async def generate_barrister_brief(case_id: str, user_id: str, report_id: str | None = None) -> dict:
    case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    documents = await db.documents.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(300)
    timeline = await db.timeline_events.find({"case_id": case_id}, {"_id": 0}).sort("event_date", 1).to_list(400)
    grounds = await db.grounds_of_merit.find({"case_id": case_id}, {"_id": 0}).to_list(100)
    source_reports = await _get_latest_standard_reports(case_id, user_id)

    timeline_text = "\n".join(
        f"- {event.get('event_date', '')}: {event.get('title', '')} — {event.get('description', '')[:200]}"
        for event in timeline[:25]
    ) or "- No timeline events recorded"
    grounds_text = "\n".join(
        f"- {ground.get('title', 'Untitled ground')} [{ground.get('ground_type', 'other')} / {ground.get('strength', 'unrated')}] — {ground.get('description', '')[:220]}"
        for ground in grounds[:20]
    ) or "- No structured grounds recorded"
    grounds_heading_text = "\n".join(
        f"### Ground {idx}: {ground.get('title', 'Untitled ground')}"
        for idx, ground in enumerate(grounds[:20], start=1)
    ) or "### Ground 1: Grounds identified from source reports"
    documents_text = "\n".join(
        f"- {document.get('filename', 'Untitled document')} ({document.get('category', 'other')})"
        for document in documents[:30]
    ) or "- No documents uploaded"

    case_profile = f"""
CASE TITLE: {case.get('title', 'Unknown case')}
DEFENDANT: {case.get('defendant_name', 'Not recorded')}
CASE NUMBER: {case.get('case_number', 'Not recorded')}
COURT: {case.get('court', 'Not recorded')}
JURISDICTION: {(case.get('state') or 'nsw').upper()}
OFFENCE: {case.get('offence_type') or case.get('offence_category') or 'Not recorded'}
SENTENCE: {case.get('sentence', 'Not recorded')}
DOCUMENT COUNT: {len(documents)}
TIMELINE EVENT COUNT: {len(timeline)}
GROUNDS COUNT: {len(grounds)}
""".strip()

    # DO NOT UNDO — Barrister View system prompt with strict depth requirements and FORENSIC language
    system_prompt = """You are a senior Australian criminal appeal barrister preparing the definitive barrister brief for a criminal appeal matter. This is the CAPSTONE document that synthesises and BUILDS UPON three earlier analytical reports (Quick Summary, Full Detailed Report, and Extensive Log). The output must read like one coherent, authoritative legal document written by a careful appellate specialist who has thoroughly digested all three source reports and is now producing a comprehensive counsel-ready working brief.

MANDATORY RULES:
- Australian English only.
- Strict third-person educational tone only.
- Never use first-person or second-person language, including: we, us, our, you, your.
- Do not mention that reports were merged, combined, synthesised, tiered, paid, unlocked, or generated by AI.
- Do not include bullet-heavy exposition. Detailed reasoning must be written in flowing paragraphs with concrete factual anchors.
- Minimal bullet points are permitted only for short authority lists, procedural checklists, or document inventories where compact presentation improves clarity.
- No duplication. If the same issue appears across multiple source reports, discuss it once in the most logical section but with GREATER depth than any individual source report achieved.
- No placeholders, meta-commentary, drafting notes, or future-tense filler such as 'will be provided'.
- If the materials are uncertain on a point, say that the available materials indicate or suggest the point rather than asserting unsupported fact.
- Use markdown headings only, with ## for main sections and ### for sub-sections.

FORENSIC APPELLATE LANGUAGE — CRITICAL:
- Appellate work is about ARGUABILITY, not declarations. All conclusions about the current case must be framed forensically.
- Use: "It is arguable that...", "It is contended that...", "There is a tenable argument that...", "On one view of the evidence...", "It is open to the appellant to argue that..."
- Do NOT use bare declarations: "The trial judge erred", "The conviction is unsafe", "The sentence is excessive". These are too definitive at the preparation stage.
- Do NOT use weak hedging: "may have", "could potentially", "it is possible that". These lack the confidence needed for counsel preparation.
- Exception: when citing what a court HAS decided in a precedent case, declarative language is appropriate.

GROUND FRAMING RULES:
- Where psychiatric/mental state evidence is involved, frame the ground as a CONVICTION SAFETY attack on mens rea — whether the requisite mental state (intent to kill, intent to cause GBH, reckless indifference) was properly determined. This is more powerful than mere evidentiary criticism.
- Where multiple jury/trial procedure issues exist (judge-alone refusal, jury reduction, juror conduct), present as sub-particulars under a single procedural unfairness ground. The CCA prefers "one ground, multiple particulars."
- Where sentencing is challenged, frame around proportionality and moral culpability — whether the sentence reflects true culpability — not merely "the judge got it wrong."
- Where ineffective counsel is advanced, mark clearly as CONTINGENT — requiring evidentiary support (affidavit, transcript confirmation). Note that without this foundation, the ground risks weakening overall appellate credibility.

DEPTH AND QUALITY MANDATE:
- This brief MUST be substantially MORE detailed than any individual source report. It is NOT a summary — it is an EXPANSION and SYNTHESIS.
- Every section must incorporate the best material from all 3 source reports AND add original barrister-level analysis that none of the source reports contain: strategic framing, submission structure, cross-examination angles, evidential deployment strategy, likely judicial concerns, and fallback positions.
- Each ground of appeal must be discussed with: factual foundation (citing specific evidence and documents), the applicable legal test (with statutory references), decisive authorities (with paragraph-level citations where available), the strongest argument for the appellant, the likely prosecution response, the reply to that response, fallback positions if the primary argument fails, and how this ground interacts with other grounds.
- The authorities section must not merely list cases. It must explain what each authority establishes, why it applies, how the prosecution might distinguish it, and how counsel should deploy it.
- Sentencing analysis must include specific comparator cases with sentence outcomes, statistical context where available, and analysis of parity, manifest excess, or other applicable principles.
- The strategy sections must be written as if counsel is preparing for a conference with the instructing solicitor: specific questions to address, documents to prepare, hearing structure, time estimates, and priority ordering of grounds.
- Generic observations such as 'this ground has merit' or 'further investigation is recommended' are UNACCEPTABLE without substantial supporting detail.
- The brief is for counsel. Avoid consumer-style explanation and avoid shrinking the material into a simplified note.
"""

    shared_context = f"""CASE PROFILE
{case_profile}

STRUCTURED GROUNDS
{grounds_text}

MANDATORY GROUND LIST
{grounds_heading_text}

TIMELINE SNAPSHOT
{timeline_text}

DOCUMENT INVENTORY
{documents_text}
"""

    # DO NOT UNDO — Section groups with calibrated source limits (larger limits cause 502 proxy errors)
    section_groups = [
        {
            "slug": "counsel-synthesis",
            "target_chars": 6000,
            "source_limits": {"quick_summary": 10000, "full_detailed": 20000, "extensive_log": 28000},
            "required_headings": [
                "## Counsel Synthesis",
            ],
            "instructions": f"""Write ONLY the ## Counsel Synthesis section. This section appears at the VERY TOP of the barrister brief and tells a barrister in 30 seconds: "Where do I attack?"

The section MUST follow this EXACT structure:

## Counsel Synthesis

### Primary Issue
One paragraph (3-4 sentences) identifying the single most important attack vector for the appeal. This should be the ground with the highest appellate viability. Frame it forensically: "It is arguable that..." or "The primary contention is that..."

### Secondary Issue
One paragraph (3-4 sentences) identifying the second strongest line of attack.

### Tertiary Issue
One paragraph (3-4 sentences) identifying the third line of attack.

### Priority Order
A numbered list of ALL grounds in order of priority, with a one-line forensic assessment of each:
1. [Strongest ground] — [one-line assessment]
2. [Second strongest] — [one-line assessment]
3. [Third] — [one-line assessment]
...and so on for all grounds.

Mark any contingent grounds (e.g. ineffective counsel) with: "(Contingent — requires evidentiary support)"

### Overall Appellate Position
One paragraph providing an honest, calibrated assessment of the appeal's overall position. Use viability language (arguable, moderate, strong) — NOT percentages. Frame forensically.

CRITICAL FRAMING RULES:
- Where psychiatric/mental state evidence undermines intent (mens rea), this should typically be the PRIMARY ISSUE — frame as a conviction safety attack on mens rea determination.
- Where jury/procedural issues exist, cluster them as a single ground with sub-particulars.
- Where sentencing is challenged, frame around proportionality and moral culpability.
- Where ineffective counsel appears, flag as contingent and place lower in priority unless strong evidentiary support exists.
- Use ONLY forensic appellate language: "It is arguable that...", "It is contended that...", "There is a tenable argument that..."
- This section must reference the specific grounds from the MANDATORY GROUND LIST below.

There are {len(grounds)} grounds identified for this case.""",
        },
        {
            "slug": "source-synthesis",
            "target_chars": 24000,
            "source_limits": {"quick_summary": 14000, "full_detailed": 28000, "extensive_log": 36000},
            "required_headings": [
                "## Executive Overview for Counsel",
                "## Source Report Synthesis",
                "## Case Background and Procedural History",
                "## Conviction, Offence and Sentence Analysis",
            ],
            "instructions": "Write these sections in full depth. Under ## Executive Overview for Counsel, provide a commanding summary of the appeal's landscape — the conviction, sentence, key vulnerabilities, strongest grounds, and the overall strategic posture. This must be dense enough to brief a senior barrister in one reading. Under ## Source Report Synthesis, create exactly these three sub-headings in this order: ### Quick Summary Synthesis, ### Full Detailed Report Synthesis, ### Extensive Log Synthesis. Under each sub-heading, identify what that source report contributes that is unique, critical, or more developed than the others — do not repeat the same point in all 3 sub-sections. Under ## Case Background and Procedural History, set out the full procedural chronology from charge to sentence with specific dates, courts, and procedural steps. Under ## Conviction, Offence and Sentence Analysis, provide a thorough analysis of the offence elements, the evidence supporting conviction, the sentencing judge's reasoning, statutory framework applied, and any sentencing errors apparent from the materials.",
        },
        {
            "slug": "case-analysis",
            "target_chars": 28000,
            "source_limits": {"quick_summary": 12000, "full_detailed": 28000, "extensive_log": 36000},
            "required_headings": [
                "## Evidentiary Tensions and Appeal Pressure Points",
                "## Grounds of Merit",
                "## Statutory Framework and Governing Tests",
                "## Authorities and Comparative Cases",
            ],
            "instructions": "Write these sections at the highest barrister depth. Under ## Evidentiary Tensions and Appeal Pressure Points, identify every contradiction, gap, procedural irregularity, missing evidence, unreliable testimony, and prosecutorial overreach found across all 3 source reports — organise by issue, explain the appellate significance, and indicate how counsel should exploit each tension point. Under ## Grounds of Merit, create one dedicated ### subsection for EVERY item in the mandatory ground list — do not omit, merge, or collapse any listed ground. Each ground must be explained with: (a) substantial factual support citing specific evidence and documents, (b) the applicable legal test with statutory references, (c) decisive authorities with paragraph-level citations, (d) the strongest argument for the appellant, (e) the likely prosecution response, (f) the reply to that response, (g) weaknesses and fallback positions, (h) strategic implications and interaction with other grounds, and (i) why counsel should care about it in conference and submissions. Under ## Statutory Framework and Governing Tests, set out the relevant legislation, appellate jurisdiction provisions, leave requirements, and the legal tests for each ground type. Under ## Authorities and Comparative Cases, meaningfully compare cases — explain what each authority establishes, why it applies, how the prosecution might distinguish it, and how counsel should deploy it. Include a markdown comparative authorities table.",
        },
        {
            "slug": "strategy",
            "target_chars": 26000,
            "source_limits": {"quick_summary": 10000, "full_detailed": 22000, "extensive_log": 32000},
            "required_headings": [
                "## Sentencing Comparison and Relief Pathways",
                "## Proposed Submissions and Hearing Strategy",
                "## Conference Questions, Filing Priorities and Risks",
                "## Final Barrister Briefing Note",
            ],
            "instructions": "Write these sections as an authoritative counsel-facing strategy brief that a barrister would use to prepare for conference and hearing. Under ## Sentencing Comparison and Relief Pathways, analyse specific comparator cases with their sentence outcomes, consider parity, manifest excess, and other sentencing principles, and set out each available relief pathway with its legal basis and likelihood of success. Under ## Proposed Submissions and Hearing Strategy, provide detailed proposed submission themes for each ground, recommended ground ordering, estimated hearing time per ground, oral submission structure, written submission priorities, key documents to include in the appeal book, and how to handle judicial questions. Under ## Conference Questions, Filing Priorities and Risks, list specific questions counsel should address with the instructing solicitor, identify filing deadlines and priority ordering, analyse litigation risk for each ground and the appeal overall, and identify what further material or instructions would strengthen the case. Under ## Final Barrister Briefing Note, provide a comprehensive closing analysis that is still detailed rather than compressed — covering the overall assessment, recommended course of action, fallback strategies, costs considerations, and any urgent steps.",
        },
    ]

    section_outputs = []
    resume_index = 0
    if report_id:
        existing_report = await db.reports.find_one(
            {"report_id": report_id},
            {"_id": 0, "content.partial_sections": 1, "content.partial_stage": 1, "content.partial_analysis": 1},
        )
        existing_content = (existing_report or {}).get("content") or {}
        saved_sections = existing_content.get("partial_sections") or []
        if saved_sections:
            section_outputs = list(saved_sections)
            resume_index = min(len(section_outputs), len(section_groups))
            logger.info(f"Resuming barrister brief {report_id} from group {resume_index + 1}")

    for group_index, group in enumerate(section_groups[resume_index:], start=resume_index + 1):
        headings_text = "\n".join(group["required_headings"])
        group_source_text = _build_barrister_group_source_text(source_reports, group["source_limits"])
        group_prompt = f"""Prepare only the following Barrister Brief sections, using the exact headings below and in the exact order:

{headings_text}

Depth requirements:
- Minimum target length for this response: {group['target_chars']} characters.
- Preserve as much useful detail as possible from the source reports.
- Use flowing paragraphs with concrete facts, authorities, procedural detail, and strategy.
- Avoid generic summary language.
- Do not repeat material unless needed for legal continuity.

Specific drafting instruction:
{group['instructions']}

{shared_context}

SOURCE REPORTS
{group_source_text}
"""

        group_response = await call_llm_with_fallback(
            system_prompt,
            group_prompt,
            session_id=f"barrister-{case_id}-{group['slug']}",
            max_tokens=16384,
            timeout_seconds=300,
            task_type="report_generation",
        )
        group_response = _strip_report_placeholders(group_response)
        group_response = re.sub(r"\n{3,}", "\n\n", group_response).strip()
        section_outputs.append(group_response)
        if report_id:
            partial_analysis = "\n\n".join(part for part in section_outputs if part).strip()
            await db.reports.update_one(
                {"report_id": report_id},
                {"$set": {
                    "content.partial_sections": section_outputs,
                    "content.partial_stage": group.get("slug"),
                    "content.partial_analysis": partial_analysis,
                    "generated_at": datetime.now(timezone.utc).isoformat(),
                }}
            )

    response = "\n\n".join(part for part in section_outputs if part).strip()

    # DO NOT UNDO — General expansion DISABLED (causes 502 proxy errors with large payloads)
    # Section-by-section expansion below is more effective and reliable.
    # DO NOT UNDO — expansion_source_text must remain defined here (used by ground expansion, cross-analysis, strategy, final QA)
    expansion_source_text = _build_barrister_group_source_text(
        source_reports,
        {"quick_summary": 10000, "full_detailed": 24000, "extensive_log": 32000},
    )

    # DO NOT UNDO — Section-by-section expansion for thin sections
    thin_section_targets = {
        "## Counsel Synthesis": 2000,
        "## Executive Overview for Counsel": 5000,
        "## Source Report Synthesis": 6000,
        "## Case Background and Procedural History": 5000,
        "## Conviction, Offence and Sentence Analysis": 5000,
        "## Evidentiary Tensions and Appeal Pressure Points": 6000,
        "## Statutory Framework and Governing Tests": 5000,
        "## Authorities and Comparative Cases": 5000,
        "## Sentencing Comparison and Relief Pathways": 5000,
        "## Proposed Submissions and Hearing Strategy": 5000,
        "## Conference Questions, Filing Priorities and Risks": 4000,
        "## Final Barrister Briefing Note": 4000,
    }
    section_expansion_source = _build_barrister_group_source_text(
        source_reports,
        {"quick_summary": 10000, "full_detailed": 26000, "extensive_log": 32000},
    )
    for section_heading, min_chars in thin_section_targets.items():
        if section_heading not in response:
            # Missing section — generate it fresh
            try:
                missing_prompt = f"""Write only the following section of a Barrister Brief. Output ONLY this one section with its heading:

{section_heading}

Requirements:
- Minimum target length: {min_chars} characters.
- Dense, counsel-facing, case-specific content using the source reports.
- Australian English, strict third-person tone.
- No bullet-heavy exposition — use flowing paragraphs.

SOURCE REPORTS
{section_expansion_source}

CURRENT BARRISTER BRIEF
{response}
"""
                missing_section = await call_llm_with_fallback(
                    system_prompt, missing_prompt,
                    session_id=f"barrister-{case_id}-missing-{section_heading[3:20].lower().replace(' ','-')}",
                    max_tokens=16384, timeout_seconds=300, task_type="report_generation",
                )
                missing_section = _strip_report_placeholders(missing_section)
                missing_section = re.sub(r"\n{3,}", "\n\n", missing_section).strip()
                if missing_section.startswith(section_heading):
                    # Find the right insertion point
                    section_order = list(thin_section_targets.keys())
                    idx = section_order.index(section_heading)
                    inserted = False
                    for later_heading in section_order[idx + 1:]:
                        if later_heading in response:
                            response = response.replace(later_heading, missing_section.rstrip() + "\n\n" + later_heading, 1)
                            inserted = True
                            break
                    if not inserted:
                        # Append before strategy sections
                        if "## Proposed Submissions and Hearing Strategy" in response:
                            response = response.replace("## Proposed Submissions and Hearing Strategy",
                                missing_section.rstrip() + "\n\n## Proposed Submissions and Hearing Strategy", 1)
                        else:
                            response = response.rstrip() + "\n\n" + missing_section
                    logger.info(f"Generated missing section: {section_heading} ({len(missing_section)} chars)")
            except Exception as exc:
                logger.warning(f"Missing section generation skipped for {section_heading}: {exc}")
            continue

        # Existing section — check if it's too thin
        section_match = re.search(
            rf"({re.escape(section_heading)}\n[\s\S]*?)(?=\n## |\Z)",
            response,
        )
        if section_match:
            section_text = section_match.group(1)
            if len(section_text) < min_chars:
                try:
                    expand_prompt = f"""Rewrite and substantially expand the following section of a Barrister Brief. Output ONLY this one section with its heading:

{section_heading}

Current section content (too thin — expand to at least {min_chars} characters):
{section_text}

Requirements:
- Keep ALL existing content and ADD more depth, more factual detail, more legal reasoning.
- Minimum target length: {min_chars} characters.
- Dense, counsel-facing, case-specific content.
- Australian English, strict third-person tone.
- No bullet-heavy exposition — use flowing paragraphs.

SOURCE REPORTS
{section_expansion_source}
"""
                    expanded_section = await call_llm_with_fallback(
                        system_prompt, expand_prompt,
                        session_id=f"barrister-{case_id}-thin-{section_heading[3:20].lower().replace(' ','-')}",
                        max_tokens=16384, timeout_seconds=300, task_type="report_generation",
                    )
                    expanded_section = _strip_report_placeholders(expanded_section)
                    expanded_section = re.sub(r"\n{3,}", "\n\n", expanded_section).strip()
                    if expanded_section.startswith(section_heading) and len(expanded_section) > len(section_text):
                        response = response.replace(section_text, expanded_section, 1)
                        logger.info(f"Expanded thin section: {section_heading} ({len(section_text)} → {len(expanded_section)} chars)")
                except Exception as exc:
                    logger.warning(f"Thin section expansion skipped for {section_heading}: {exc}")

    # DO NOT UNDO — Ground expansion: rewrites each ground as a ### subsection with deep analysis
    if len(response) < 100000 and grounds:
        try:
            rewritten_ground_sections = []
            for ground_index, ground in enumerate(grounds, start=1):
                ground_title = (ground.get("title") or f"Ground {ground_index}").strip()
                ground_prompt = f"""Write only the following Barrister Brief subsection in this exact heading format:

### {ground_title}

Requirements:
- Minimum target length for this subsection alone: 7000 characters.
- Focus on critical, vital, counsel-useful detail rather than generic explanation.
- Set out the factual foundation for the ground, procedural context, legal test, statutory framework, decisive authorities, documentary anchors, timeline anchors, evidentiary vulnerabilities, likely prosecution answer, reply strategy, fallback positions, relief implications, and how counsel should frame the point in written and oral submissions.
- Do not repeat generic appeal language.
- Keep the tone strictly barrister-facing and case-specific.

STRUCTURED GROUND
TITLE: {ground_title}
DESCRIPTION: {ground.get('description', '')}
LEGAL BASIS: {ground.get('legal_basis', '')}
WHY IT MATTERS: {ground.get('strength_reason', '')}
RISK LEVEL: {ground.get('strength', '')}
SUPPORTING DOCUMENT IDS: {', '.join(ground.get('supporting_documents', []) or [])}

MANDATORY GROUND LIST
{grounds_heading_text}

SOURCE REPORTS
{expansion_source_text}

CURRENT BARRISTER BRIEF
{response}
"""
                subsection = await call_llm_with_fallback(
                    system_prompt,
                    ground_prompt,
                    session_id=f"barrister-{case_id}-ground-{ground_index}-expand",
                    max_tokens=16384,
                    timeout_seconds=300,
                    task_type="report_generation",
                )
                subsection = _strip_report_placeholders(subsection)
                subsection = re.sub(r"\n{3,}", "\n\n", subsection).strip()
                # DO NOT UNDO — Strip any heading format the LLM used and enforce ### subsection format
                subsection = re.sub(r'^#{1,4}\s*' + re.escape(ground_title) + r'\s*\n+', '', subsection).strip()
                subsection = f"### {ground_title}\n\n{subsection}"
                rewritten_ground_sections.append(subsection)

            rewritten_grounds = "## Grounds of Merit\n\n" + "\n\n".join(rewritten_ground_sections).strip()
            if rewritten_ground_sections:
                response = re.sub(
                    r"## Grounds of Merit\n[\s\S]*?(?=\n## Statutory Framework and Governing Tests)",
                    rewritten_grounds + "\n\n",
                    response,
                    count=1,
                )
        except Exception as exc:
            logger.warning(f"Barrister grounds expansion skipped for {case_id}: {exc}")

    # DO NOT UNDO — Cross-analysis MUST run BEFORE strategy expansion to avoid content loss
    # Cross-analysis sections insert before Sentencing Comparison, not at end
    # Always run — these are unique sections not generated elsewhere
    if "## Report-to-Report Cross-Analysis" not in response:
        cross_analysis_prompt = f"""Produce only the following sections, in this exact order:

## Report-to-Report Cross-Analysis
## Document and Evidence Deployment for Counsel

Requirements:
- Minimum target length for this response: 28000 characters.
- Use all 3 source reports and identify where they overlap, where they diverge, and what unique critical material each one contributes.
- Do not recycle the same paragraph 3 times. Organise the material by issue and explain which source report adds what.
- Under ## Document and Evidence Deployment for Counsel, explain how specific documents, witnesses, chronology items, psychiatric material, media material, and procedural incidents should be deployed by counsel in conference, written submissions, and oral argument.
- This must be dense, counsel-facing, and case-specific.

SOURCE REPORTS
{expansion_source_text}

CURRENT BARRISTER BRIEF
{response}
"""
        try:
            cross_analysis = await call_llm_with_fallback(
                system_prompt,
                cross_analysis_prompt,
                session_id=f"barrister-{case_id}-cross-analysis",
                max_tokens=16384,
                timeout_seconds=300,
                task_type="report_generation",
            )
            cross_analysis = _strip_report_placeholders(cross_analysis)
            cross_analysis = re.sub(r"\n{3,}", "\n\n", cross_analysis).strip()
            if cross_analysis.startswith("## Report-to-Report Cross-Analysis"):
                # DO NOT UNDO — Insert before Sentencing Comparison, not at the end
                # This ensures cross-analysis isn't eaten by the strategy expansion regex
                if "## Sentencing Comparison and Relief Pathways" in response:
                    response = response.replace(
                        "## Sentencing Comparison and Relief Pathways",
                        cross_analysis.rstrip() + "\n\n## Sentencing Comparison and Relief Pathways",
                        1,
                    )
                else:
                    response = response.rstrip() + "\n\n" + cross_analysis
        except Exception as exc:
            logger.warning(f"Barrister cross-analysis expansion skipped for {case_id}: {exc}")

    # Strategy expansion — always run to deepen strategy sections
    if "## Proposed Submissions and Hearing Strategy" in response:
        strategy_expansion_prompt = f"""Rewrite only the following sections of the Barrister Brief below, keeping the same headings and making them substantially more detailed:

## Proposed Submissions and Hearing Strategy
## Conference Questions, Filing Priorities and Risks
## Final Barrister Briefing Note

Requirements:
- Minimum target length for this rewritten block alone: 30000 characters.
- Expand the oral and written submissions structure greatly.
- Include issue sequencing, fallback positions, framing choices, evidentiary use, likely resistance points, answer lines, and conference questions for counsel.
- The rewritten material must read like a serious barrister working brief, not a summary.
- Avoid repeating the same sentence structure or generic observations.

SOURCE REPORTS
{expansion_source_text}

CURRENT BARRISTER BRIEF
{response}
"""
        try:
            rewritten_strategy = await call_llm_with_fallback(
                system_prompt,
                strategy_expansion_prompt,
                session_id=f"barrister-{case_id}-strategy-expand",
                max_tokens=16384,
                timeout_seconds=300,
                task_type="report_generation",
            )
            rewritten_strategy = _strip_report_placeholders(rewritten_strategy)
            rewritten_strategy = re.sub(r"\n{3,}", "\n\n", rewritten_strategy).strip()
            if rewritten_strategy.startswith("## Proposed Submissions and Hearing Strategy"):
                # DO NOT UNDO — Regex stops at Attachment A/Report-to-Report/Document and Evidence.
                # Old regex ([\s\S]*$) ate cross-analysis sections that were appended after strategy.
                response = re.sub(
                    r"## Proposed Submissions and Hearing Strategy\n[\s\S]*?(?=\n## (?:Attachment A|Report-to-Report|Document and Evidence)|$)",
                    rewritten_strategy + "\n\n",
                    response,
                    count=1,
                )
        except Exception as exc:
            logger.warning(f"Barrister strategy expansion skipped for {case_id}: {exc}")

    try:
        comparison_tables_prompt = f"""Produce only the following markdown blocks, in this exact order, with no introduction or conclusion:

### Evidentiary Pressure Points Table
Create a markdown table with columns: Issue | Supporting Material | Why it matters on appeal | Vulnerability in the prosecution case

### Comparative Authorities Table
Create a markdown table with columns: Authority | Principle | Relevance to this case | Strategic use in submissions

### Sentencing Comparison Table
Create a markdown table with columns: Comparator | Key facts | Sentence outcome | Relevance to this appeal

### Relief Pathways Matrix
Create a markdown table with columns: Relief pathway | Legal basis | Best supporting features | Main risk or limitation

Requirements:
- Use all 3 source reports and the current barrister brief.
- Make the rows detailed and case-specific.
- Do not repeat the main narrative prose.
- Output only the 4 titled blocks above.

SOURCE REPORTS
{expansion_source_text}

CURRENT BARRISTER BRIEF
{response}
"""
        comparison_tables = await call_llm_with_fallback(
            system_prompt,
            comparison_tables_prompt,
            session_id=f"barrister-{case_id}-comparison-tables",
            max_tokens=16384,
            timeout_seconds=300,
            task_type="report_generation",
        )
        comparison_tables = _strip_report_placeholders(comparison_tables)
        comparison_tables = re.sub(r"\n{3,}", "\n\n", comparison_tables).strip()

        evidence_table_match = re.search(
            r"(### Evidentiary Pressure Points Table[\s\S]*?)(?=\n### Comparative Authorities Table|$)",
            comparison_tables,
        )
        authorities_table_match = re.search(
            r"(### Comparative Authorities Table[\s\S]*?)(?=\n### Sentencing Comparison Table|$)",
            comparison_tables,
        )
        sentencing_table_match = re.search(
            r"(### Sentencing Comparison Table[\s\S]*?)(?=\n### Relief Pathways Matrix|$)",
            comparison_tables,
        )
        relief_table_match = re.search(r"(### Relief Pathways Matrix[\s\S]*)$", comparison_tables)

        if evidence_table_match and "### Evidentiary Pressure Points Table" not in response:
            response = re.sub(
                r"(## Evidentiary Tensions and Appeal Pressure Points\n[\s\S]*?)(?=\n## Grounds of Merit)",
                lambda match: match.group(1).rstrip() + "\n\n" + evidence_table_match.group(1).strip() + "\n\n",
                response,
                count=1,
            )

        if authorities_table_match and "### Comparative Authorities Table" not in response:
            response = re.sub(
                r"(## Authorities and Comparative Cases\n[\s\S]*?)(?=\n## Sentencing Comparison and Relief Pathways)",
                lambda match: match.group(1).rstrip() + "\n\n" + authorities_table_match.group(1).strip() + "\n\n",
                response,
                count=1,
            )

        sentencing_blocks = []
        if sentencing_table_match and "### Sentencing Comparison Table" not in response:
            sentencing_blocks.append(sentencing_table_match.group(1).strip())
        if relief_table_match and "### Relief Pathways Matrix" not in response:
            sentencing_blocks.append(relief_table_match.group(1).strip())
        if sentencing_blocks:
            response = re.sub(
                r"(## Sentencing Comparison and Relief Pathways\n[\s\S]*?)(?=\n## Proposed Submissions and Hearing Strategy)",
                lambda match: match.group(1).rstrip() + "\n\n" + "\n\n".join(sentencing_blocks) + "\n\n",
                response,
                count=1,
            )
    except Exception as exc:
        logger.warning(f"Barrister comparison table enrichment skipped for {case_id}: {exc}")

    if grounds:
        try:
            issue_matrix_prompt = f"""Produce only the following attachment, with no introduction or conclusion before it:

## Attachment A — Barrister Issue Matrix

Requirements:
- This must appear as a barrister attachment at the end of the report.
- Use a markdown table with these columns exactly: Ground | Key facts | Main authorities | Risk level | Likely prosecution response | Oral submission angle.
- One row per ground from the mandatory ground list.
- Keep each cell specific, practical, and counsel-facing.
- Do not use generic filler.

MANDATORY GROUND LIST
{grounds_heading_text}

STRUCTURED GROUNDS
{grounds_text}

SOURCE REPORTS
{expansion_source_text}

CURRENT BARRISTER BRIEF
{response}
"""
            issue_matrix = await call_llm_with_fallback(
                system_prompt,
                issue_matrix_prompt,
                session_id=f"barrister-{case_id}-issue-matrix",
                max_tokens=16384,
                timeout_seconds=300,
                task_type="report_generation",
            )
            issue_matrix = _strip_report_placeholders(issue_matrix)
            issue_matrix = re.sub(r"\n{3,}", "\n\n", issue_matrix).strip()
            if issue_matrix.startswith("## Attachment A — Barrister Issue Matrix"):
                response = response.rstrip() + "\n\n" + issue_matrix
        except Exception as exc:
            logger.warning(f"Barrister issue matrix skipped for {case_id}: {exc}")

    # DO NOT UNDO — Final quality validation pass: expand any remaining thin sections
    final_min_chars = {
        "## Counsel Synthesis": 1500,
        "## Executive Overview for Counsel": 5000,
        "## Source Report Synthesis": 5000,
        "## Case Background and Procedural History": 4000,
        "## Conviction, Offence and Sentence Analysis": 4000,
        "## Evidentiary Tensions and Appeal Pressure Points": 5000,
        "## Grounds of Merit": 20000,
        "## Statutory Framework and Governing Tests": 4000,
        "## Authorities and Comparative Cases": 4000,
        "## Sentencing Comparison and Relief Pathways": 5000,
        "## Proposed Submissions and Hearing Strategy": 4000,
        "## Conference Questions, Filing Priorities and Risks": 3000,
        "## Final Barrister Briefing Note": 3000,
    }
    final_source = _build_barrister_group_source_text(
        source_reports,
        {"quick_summary": 8000, "full_detailed": 20000, "extensive_log": 28000},
    )
    for heading, min_c in final_min_chars.items():
        section_match = re.search(
            rf"({re.escape(heading)}\n[\s\S]*?)(?=\n## |\Z)",
            response,
        )
        if section_match and len(section_match.group(1)) < min_c:
            try:
                current_section = section_match.group(1)
                final_expand_prompt = f"""Rewrite and substantially expand the following section. Output ONLY this one section with its heading.

{current_section}

Requirements:
- Keep ALL existing content and ADD significantly more depth, factual detail, legal reasoning, and practical counsel advice.
- Minimum target: {min_c} characters.
- Dense, counsel-facing, Australian English, strict third-person.
- Use flowing paragraphs with concrete factual anchors, not generic observations.

SOURCE REPORTS
{final_source}
"""
                expanded = await call_llm_with_fallback(
                    system_prompt, final_expand_prompt,
                    session_id=f"barrister-{case_id}-final-{heading[3:15].lower().replace(' ','-')}",
                    max_tokens=16384, timeout_seconds=300, task_type="report_generation",
                )
                expanded = _strip_report_placeholders(expanded)
                expanded = re.sub(r"\n{3,}", "\n\n", expanded).strip()
                if expanded.startswith(heading) and len(expanded) > len(current_section):
                    response = response.replace(current_section, expanded, 1)
                    logger.info(f"Final QA expanded: {heading} ({len(current_section)} → {len(expanded)} chars)")
            except Exception as exc:
                logger.warning(f"Final QA expansion skipped for {heading}: {exc}")

    response = _strip_report_placeholders(response)
    response = _normalise_barrister_table_titles(response)
    response = _dedupe_barrister_ground_subsections(response)

    # DO NOT UNDO — Strip rogue H2 sections not in the expected structure
    expected_h2s = {
        "## Counsel Synthesis",
        "## Executive Overview for Counsel",
        "## Source Report Synthesis",
        "## Case Background and Procedural History",
        "## Conviction, Offence and Sentence Analysis",
        "## Evidentiary Tensions and Appeal Pressure Points",
        "## Grounds of Merit",
        "## Statutory Framework and Governing Tests",
        "## Authorities and Comparative Cases",
        "## Report-to-Report Cross-Analysis",
        "## Document and Evidence Deployment for Counsel",
        "## Sentencing Comparison and Relief Pathways",
        "## Proposed Submissions and Hearing Strategy",
        "## Conference Questions, Filing Priorities and Risks",
        "## Conference Questions, Filing Priorities, and Risks",
        "## Final Barrister Briefing Note",
        "## Attachment A — Barrister Issue Matrix",
        "## Attachment A - Barrister Issue Matrix",
    }
    cleaned_lines = []
    skip_section = False
    for line in response.split("\n"):
        if line.startswith("## "):
            if line.strip() in expected_h2s:
                skip_section = False
                cleaned_lines.append(line)
            else:
                skip_section = True
                logger.info(f"Stripped rogue section: {line.strip()}")
        elif skip_section and not line.startswith("## "):
            continue
        else:
            cleaned_lines.append(line)
    response = "\n".join(cleaned_lines)

    response = re.sub(r"\n{3,}", "\n\n", response).strip()

    return {
        "analysis": response,
        "case_data": case,
        "document_count": len(documents),
        "event_count": len(timeline),
        "grounds_of_merit": grounds,
        "source_reports": [
            {
                "report_id": report.get("report_id"),
                "report_type": report.get("report_type"),
                "title": report.get("title"),
                "generated_at": report.get("generated_at"),
            }
            for report in source_reports
        ],
        "source_signature": _build_barrister_source_signature(source_reports),
        "metadata": ReportMetadata(
            documents_analyzed=len(documents),
            timeline_events_analyzed=len(timeline),
            grounds_considered=len(grounds),
            verification_status="draft",
            confidence_note="AI-generated barrister brief requiring legal review",
        ).model_dump(),
    }


async def _run_barrister_report_generation(report_id: str, case_id: str, user_id: str):
    """Background task for barrister brief generation — HARDENED with metadata."""
    try:
        analysis_result = await generate_barrister_brief(case_id, user_id, report_id=report_id)
        await db.reports.update_one(
            {"report_id": report_id},
            {
                "$set": {
                    "status": "completed",
                    "title": "Barrister Brief",
                    "content": {
                        "analysis": analysis_result["analysis"],
                        "case_title": (analysis_result.get("case_data") or {}).get("title", ""),
                        "defendant": (analysis_result.get("case_data") or {}).get("defendant_name", ""),
                        "document_count": analysis_result.get("document_count", 0),
                        "event_count": analysis_result.get("event_count", 0),
                        "source_signature": analysis_result.get("source_signature", ""),
                        "source_reports": analysis_result.get("source_reports", []),
                        "aggressive_mode": False,
                        "partial_sections": [],
                        "partial_stage": None,
                        "partial_analysis": "",
                    },
                    "grounds_of_merit": analysis_result["grounds_of_merit"],
                    "metadata": analysis_result.get("metadata"),
                    "source_mode": "ai_generated",
                    "verification_status": "draft",
                    "error": None,
                    "technical_error": None,
                }
            },
        )
        logger.info(f"Barrister brief {report_id} generated successfully")
    except Exception as exc:
        logger.error(f"Barrister brief {report_id} generation failed: {exc}")
        friendly_error = "Barrister brief generation was interrupted by a temporary AI service error. Please generate again."
        # DO_NOT_UNDO — Restore backup on failure
        backup_doc = await db.reports.find_one(
            {"report_id": report_id},
            {"_id": 0, "content.backup_analysis": 1}
        )
        backup = (backup_doc or {}).get("content", {}).get("backup_analysis", "")
        if backup and len(backup) > 5000:
            await db.reports.update_one(
                {"report_id": report_id},
                {"$set": {
                    "status": "completed",
                    "error": None,
                    "content.analysis": backup,
                },
                "$unset": {"content.backup_analysis": 1}}
            )
            logger.info(f"Barrister {report_id} generation failed but restored from backup ({len(backup)} chars)")
        else:
            await db.reports.update_one(
                {"report_id": report_id},
                {"$set": {"status": "failed", "error": friendly_error, "technical_error": str(exc)}},
            )

async def _run_report_generation(report_id: str, case_id: str, user_id: str, report_type: str, aggressive_mode: bool):
    """Background task that generates the AI report and updates the DB record — HARDENED with metadata."""
    pipeline_refresh_result = {"refreshed": False, "extracted_count": 0, "classified_count": 0, "synced_count": 0, "auto_verify_limit": 0, "auto_verify_result": {"attempted": 0, "verified": 0, "failed": 0}}
    try:
        # Pre-draft pipeline refresh: ensure staged artifacts are current
        case = await db.cases.find_one({"case_id": case_id, "user_id": user_id}, {"_id": 0})
        if case:
            documents = await db.documents.find(
                {"case_id": case_id, "user_id": user_id}, {"_id": 0, "file_data": 0}
            ).to_list(500)
            try:
                needs_refresh = await _pipeline_artifacts_missing_or_stale(case, documents)
                if needs_refresh:
                    refreshed = await _refresh_pipeline_for_reporting(case, documents, report_type)
                    pipeline_refresh_result = {"refreshed": True, **refreshed}
                    logger.info(f"Pipeline refreshed for report {report_id}: {pipeline_refresh_result}")
                else:
                    # Even if artifacts exist, paid report tiers may justify additional auto-verification
                    auto_verify_limit = _auto_verification_limit_for_report_type(report_type)
                    if auto_verify_limit > 0:
                        selected_issues = await _select_issues_for_auto_verification(case, auto_verify_limit)
                        auto_verify_result = await _auto_verify_selected_issues(case, selected_issues)
                        synced_count = await _sync_pipeline_projection_to_grounds(case)
                        pipeline_refresh_result = {
                            "refreshed": False,
                            "extracted_count": 0,
                            "classified_count": 0,
                            "synced_count": synced_count,
                            "auto_verify_limit": auto_verify_limit,
                            "auto_verify_result": auto_verify_result,
                        }
            except Exception as pe:
                logger.warning(f"Pipeline refresh before report {report_id} failed (non-fatal): {pe}")

        report_titles = {
            "quick_summary": "Quick Case Summary",
            "full_detailed": "Full Detailed Legal Analysis",
            "extensive_log": "Extensive Case Log & Analysis"
        }
        analysis_result = await analyze_case_with_ai(case_id, user_id, report_type, aggressive_mode, report_id=report_id)
        new_analysis = analysis_result.get("analysis", "")
        
        # DO_NOT_UNDO — Content protection: never overwrite a longer report with a shorter one.
        # If the new content is less than 50% the length of the backup, something went wrong
        # (e.g. 502 errors truncated the generation). Keep the backup instead.
        backup_doc = await db.reports.find_one(
            {"report_id": report_id},
            {"_id": 0, "content.backup_analysis": 1}
        )
        backup_analysis = (backup_doc or {}).get("content", {}).get("backup_analysis", "")
        if backup_analysis and len(backup_analysis) > 10000 and len(new_analysis) < len(backup_analysis) * 0.5:
            logger.warning(
                f"Report {report_id}: new content ({len(new_analysis)} chars) is less than 50% of backup "
                f"({len(backup_analysis)} chars). Keeping backup to prevent content loss."
            )
            await db.reports.update_one(
                {"report_id": report_id},
                {"$set": {
                    "status": "completed",
                    "error": None,
                    "content.analysis": backup_analysis,
                    "content.partial": False,
                },
                "$unset": {"content.backup_analysis": 1}}
            )
            logger.info(f"Report {report_id}: restored backup instead of thin regeneration")
            return

        title = report_titles.get(report_type, "Report")
        if aggressive_mode:
            title = f"{title} (Aggressive)"
        await db.reports.update_one(
            {"report_id": report_id},
            {"$set": {
                "status": "completed",
                "title": title,
                "content": {
                    "analysis": analysis_result["analysis"],
                    "case_title": (analysis_result.get("case_data") or {}).get("title", ""),
                    "defendant": (analysis_result.get("case_data") or {}).get("defendant_name", ""),
                    "document_count": analysis_result.get("document_count", 0),
                    "event_count": analysis_result.get("event_count", 0),
                    "aggressive_mode": aggressive_mode,
                    "draft_source": "pipeline" if (analysis_result.get("pipeline_metadata") or {}).get("status") in ("fresh", "refreshed") else "legacy",
                },
                "grounds_of_merit": analysis_result.get("grounds_of_merit", []),
                "metadata": {
                    **(analysis_result.get("metadata") or {}),
                    "pipeline_refresh_before_draft": pipeline_refresh_result,
                    "pipeline_issue_count": len(((analysis_result.get("pipeline_metadata") or {}).get("staleness") or {}).get("extract_missing_for_documents", [])) if analysis_result.get("pipeline_metadata") else 0,
                    "pipeline_verification_count": 0,
                    **(analysis_result.get("pipeline_metadata") or {}),
                },
                "source_mode": "ai_generated",
                "verification_status": "draft",
            }}
        )
        logger.info(f"Report {report_id} generated successfully")
        # Clear backup after successful generation
        await db.reports.update_one(
            {"report_id": report_id},
            {"$unset": {"content.backup_analysis": 1}}
        )
    except Exception as exc:
        logger.error(f"Report {report_id} generation failed: {exc}")
        friendly_error = "Report generation was interrupted by a temporary AI service error. Retry resumes from the last completed section."
        # DO_NOT_UNDO — Restore backup on failure. If backup_analysis exists,
        # restore it so the user doesn't lose their previous report.
        backup_doc = await db.reports.find_one(
            {"report_id": report_id},
            {"_id": 0, "content.backup_analysis": 1}
        )
        backup = (backup_doc or {}).get("content", {}).get("backup_analysis", "")
        if backup and len(backup) > 5000:
            await db.reports.update_one(
                {"report_id": report_id},
                {"$set": {
                    "status": "completed",
                    "error": None,
                    "content.analysis": backup,
                },
                "$unset": {"content.backup_analysis": 1}}
            )
            logger.info(f"Report {report_id} generation failed but restored from backup ({len(backup)} chars)")
        else:
            await db.reports.update_one(
                {"report_id": report_id},
                {"$set": {"status": "failed", "error": friendly_error, "technical_error": str(exc)}}
            )


@api_router.post("/cases/{case_id}/reports/generate", response_model=dict)
async def generate_report(case_id: str, report_request: ReportRequest, request: Request):
    """Generate an AI-powered report for a case (background task)"""
    user = await get_current_user(request)
    report_type = report_request.report_type
    
    if report_type not in ["quick_summary", "full_detailed", "extensive_log"]:
        raise HTTPException(status_code=400, detail="Invalid report type")
    
    # DO_NOT_UNDO — Block duplicate generation. If a report of this type is
    # already generating, return its status instead of creating a duplicate.
    existing_generating = await db.reports.find_one(
        {
            "case_id": case_id,
            "user_id": user.user_id,
            "report_type": report_type,
            "status": "generating",
            "content.aggressive_mode": {"$ne": True},
        },
        {"_id": 0, "report_id": 1},
    )
    if existing_generating:
        return {"report_id": existing_generating["report_id"], "status": "generating", "report_type": report_type}
    
    # Check payment for premium reports (admin bypasses all payments)
    is_admin = is_admin_user(user.email)
    
    if report_type == "full_detailed" and not is_admin:
        payment = await db.payments.find_one({
            "case_id": case_id,
            "user_id": user.user_id,
            "feature_type": {"$in": feature_type_variants("full_report")},
            "status": "completed"
        })
        if not payment:
            raise HTTPException(
                status_code=402, 
                detail={
                    "message": "Payment required for Full Detailed Report",
                    "feature_type": "full_report",
                    "price": FEATURE_PRICES["full_report"]["price"],
                    "currency": "AUD"
                }
            )
    
    if report_type == "extensive_log" and not is_admin:
        payment = await db.payments.find_one({
            "case_id": case_id,
            "user_id": user.user_id,
            "feature_type": {"$in": feature_type_variants("extensive_report")},
            "status": "completed"
        })
        if not payment:
            raise HTTPException(
                status_code=402, 
                detail={
                    "message": "Payment required for Extensive Log Report",
                    "feature_type": "extensive_report",
                    "price": FEATURE_PRICES["extensive_report"]["price"],
                    "currency": "AUD"
                }
            )
    
    existing_failed = await db.reports.find(
        {
            "case_id": case_id,
            "user_id": user.user_id,
            "report_type": report_type,
            "status": "failed",
            "content.aggressive_mode": {"$ne": True},
        },
        {"_id": 0},
    ).sort("generated_at", -1).to_list(1)

    if existing_failed:
        resumed_report = existing_failed[0]
        await db.reports.update_one(
            {"report_id": resumed_report["report_id"]},
            {"$set": {"status": "generating", "error": None, "generated_at": datetime.now(timezone.utc).isoformat()}},
        )
        asyncio.create_task(
            _run_report_generation(resumed_report["report_id"], case_id, user.user_id, report_type, False)
        )
        resumed_report["status"] = "generating"
        resumed_report["error"] = None
        return resumed_report

    # If a completed report of this type already exists, reuse its ID (regenerate in-place)
    existing_completed = await db.reports.find_one(
        {
            "case_id": case_id,
            "user_id": user.user_id,
            "report_type": report_type,
            "status": "completed",
            "content.aggressive_mode": {"$ne": True},
        },
        {"_id": 0, "report_id": 1},
    )
    if existing_completed:
        report_id = existing_completed["report_id"]
        # DO_NOT_UNDO — NEVER wipe content.analysis during regeneration.
        # Keep the old report visible to the user while the new one generates.
        # Only clear partial_analysis and passes_completed (used by the generation engine).
        # The old analysis stays visible until the new generation completes and overwrites it.
        existing_doc = await db.reports.find_one({"report_id": report_id}, {"_id": 0, "content.analysis": 1})
        old_analysis = existing_doc.get("content", {}).get("analysis", "") if existing_doc else ""
        await db.reports.update_one(
            {"report_id": report_id},
            {"$set": {
                "status": "generating",
                "error": None,
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "content.partial_analysis": "",
                "content.passes_completed": 0,
                "content.aggressive_mode": False,
                "content.backup_analysis": old_analysis,
            }}
        )
        asyncio.create_task(
            _run_report_generation(report_id, case_id, user.user_id, report_type, False)
        )
        return {"report_id": report_id, "status": "generating", "report_type": report_type}

    # Create a placeholder report with "generating" status and return immediately
    report_id = f"rpt_{uuid.uuid4().hex[:12]}"
    report_titles = {
        "quick_summary": "Quick Case Summary",
        "full_detailed": "Full Detailed Legal Analysis",
        "extensive_log": "Extensive Case Log & Analysis"
    }
    aggressive_mode = False
    title = report_titles.get(report_type, "Report")
    placeholder = {
        "report_id": report_id,
        "case_id": case_id,
        "user_id": user.user_id,
        "report_type": report_type,
        "title": title,
        "content": {"analysis": "", "aggressive_mode": aggressive_mode},
        "grounds_of_merit": [],
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "status": "generating",
        "source_mode": "ai_generated",
        "verification_status": "draft",
    }
    insert_doc = {k: v for k, v in placeholder.items()}
    await db.reports.insert_one(insert_doc)

    # Fire-and-forget background task
    asyncio.create_task(
        _run_report_generation(report_id, case_id, user.user_id, report_type, aggressive_mode)
    )

    return placeholder


@api_router.get("/cases/{case_id}/reports/{report_id}/status")
async def get_report_status(case_id: str, report_id: str, request: Request):
    """Poll endpoint for report generation status"""
    user = await get_current_user(request)
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0, "report_id": 1, "status": 1, "error": 1}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    return {"report_id": report_id, "status": report.get("status", "completed")}

@api_router.get("/cases/{case_id}/reports", response_model=List[dict])
async def get_reports(case_id: str, request: Request):
    """Get all reports for a case"""
    user = await get_current_user(request)
    
    # Auto-fail any report stuck in "generating" for more than 60 minutes
    thirty_min_ago = (datetime.now(timezone.utc) - timedelta(minutes=60)).isoformat()
    await db.reports.update_many(
        {"case_id": case_id, "user_id": user.user_id, "status": "generating", "generated_at": {"$lt": thirty_min_ago}},
        {"$set": {"status": "failed", "error": "Generation timed out"}}
    )
    
    reports = await db.reports.find(
        {
            "case_id": case_id,
            "user_id": user.user_id,
            "content.aggressive_mode": {"$ne": True},
        },
        {"_id": 0}
    ).sort("generated_at", -1).to_list(100)
    
    return reports


@api_router.get("/cases/{case_id}/reports/barrister-view", response_model=dict)
async def get_or_generate_barrister_view(case_id: str, request: Request, regenerate: bool = False):
    """Return the current barrister brief or start a fresh synthesis when explicitly requested."""
    user = await get_current_user(request)
    source_reports = await _get_latest_standard_reports(case_id, user.user_id)
    source_signature = _build_barrister_source_signature(source_reports)

    existing_reports = await db.reports.find(
        {
            "case_id": case_id,
            "user_id": user.user_id,
            "report_type": "barrister_view",
        },
        {"_id": 0},
    ).sort("generated_at", -1).to_list(10)

    # Find current report — prefer one matching current source signature
    current_report = None
    for r in existing_reports:
        if (r.get("content") or {}).get("source_signature") == source_signature:
            current_report = r
            break
    if not current_report and existing_reports:
        current_report = existing_reports[0]

    if current_report and not regenerate:
        current_status = current_report.get("status")
        report_id_cur = current_report.get("report_id")
        if current_status == "completed":
            current_analysis = ((current_report.get("content") or {}).get("analysis") or "").strip()
            if len(current_analysis) < 4000:
                # DO_NOT_UNDO — Backup before auto-regen of thin barrister report
                await db.reports.update_one(
                    {"report_id": report_id_cur},
                    {"$set": {
                        "status": "generating",
                        "error": None,
                        "generated_at": datetime.now(timezone.utc).isoformat(),
                        "content.backup_analysis": current_analysis,
                    }},
                )
                current_report["status"] = "generating"
                current_report["error"] = None
                asyncio.create_task(_run_barrister_report_generation(report_id_cur, case_id, user.user_id))
                return current_report
            return current_report
        if current_status == "failed":
            # Don't auto-retry — return as-is so user can decide to regenerate
            return current_report
        if current_status == "generating":
            generated_at = _coerce_utc_datetime(current_report.get("generated_at"))
            stale_cutoff = datetime.now(timezone.utc) - timedelta(minutes=BARRISTER_GENERATION_TIMEOUT_MINUTES)
            if generated_at and generated_at >= stale_cutoff:
                return current_report

            timeout_message = "Barrister brief generation timed out. Please generate again."
            await db.reports.update_one(
                {"report_id": report_id_cur},
                {"$set": {"status": "failed", "error": timeout_message, "technical_error": timeout_message}},
            )
            current_report["status"] = "failed"
            current_report["error"] = timeout_message
            return current_report

    # If no existing report and not regenerate — return 404 so frontend knows to show "Generate" button
    if not regenerate:
        raise HTTPException(status_code=404, detail="Barrister brief has not been generated yet. Select 'Generate' to create one.")

    # Regenerate requested — create or reuse
    if current_report and current_report.get("status") in ("completed", "failed"):
        report_id_cur = current_report["report_id"]
        await db.reports.update_one(
            {"report_id": report_id_cur},
            {"$set": {"status": "generating", "error": None, "generated_at": datetime.now(timezone.utc).isoformat()}},
        )
        current_report["status"] = "generating"
        current_report["error"] = None
        asyncio.create_task(_run_barrister_report_generation(report_id_cur, case_id, user.user_id))
        return current_report

    report_id = f"rpt_{uuid.uuid4().hex[:12]}"
    placeholder = {
        "report_id": report_id,
        "case_id": case_id,
        "user_id": user.user_id,
        "report_type": "barrister_view",
        "title": "Barrister Brief",
        "content": {
            "analysis": "",
            "document_count": len(source_reports),
            "event_count": 0,
            "source_signature": source_signature,
            "source_reports": [
                {
                    "report_id": report.get("report_id"),
                    "report_type": report.get("report_type"),
                    "title": report.get("title"),
                    "generated_at": report.get("generated_at"),
                }
                for report in source_reports
            ],
            "aggressive_mode": False,
        },
        "grounds_of_merit": [],
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "status": "generating",
    }
    await db.reports.insert_one({k: v for k, v in placeholder.items()})

    asyncio.create_task(_run_barrister_report_generation(report_id, case_id, user.user_id))
    return placeholder


async def _get_latest_completed_barrister_report(case_id: str, user_id: str):
    return await db.reports.find_one(
        {
            "case_id": case_id,
            "user_id": user_id,
            "report_type": "barrister_view",
            "status": "completed",
        },
        {"_id": 0},
        sort=[("generated_at", -1)],
    )


@api_router.get("/cases/{case_id}/reports/barrister-view/export-pdf")
async def export_latest_barrister_view_pdf(case_id: str, request: Request):
    user = await get_current_user(request)
    # Resolve case owner for admin cross-user access
    case = await db.cases.find_one({"case_id": case_id}, {"_id": 0, "user_id": 1})
    owner_user_id = (case or {}).get("user_id", user.user_id)
    report = await _get_latest_completed_barrister_report(case_id, owner_user_id)
    if not report:
        raise HTTPException(status_code=404, detail="Completed barrister report not found")
    return await export_report_pdf(case_id, report["report_id"], request)


@api_router.get("/cases/{case_id}/reports/barrister-view/export-docx")
async def export_latest_barrister_view_docx(case_id: str, request: Request):
    user = await get_current_user(request)
    # Resolve case owner for admin cross-user access
    case = await db.cases.find_one({"case_id": case_id}, {"_id": 0, "user_id": 1})
    owner_user_id = (case or {}).get("user_id", user.user_id)
    report = await _get_latest_completed_barrister_report(case_id, owner_user_id)
    if not report:
        raise HTTPException(status_code=404, detail="Completed barrister report not found")
    return await export_report_docx(case_id, report["report_id"], request)


# DO NOT UNDO — Barrister Quick Brief: 2-page PDF with Counsel Synthesis + Priority Order + top 3 grounds
@api_router.get("/cases/{case_id}/reports/barrister-quick-brief")
async def export_barrister_quick_brief(case_id: str, request: Request):
    """Generate a concise 2-page Barrister Quick Brief PDF.
    Contains Counsel Synthesis, Priority Order, and top 3 grounds only.
    Designed for a barrister to review in under 5 minutes before a conference."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from io import BytesIO

    user = await get_current_user(request)

    # Get case
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        if is_admin_user(user.email):
            case = await db.cases.find_one({"case_id": case_id}, {"_id": 0})
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

    # Get barrister report — use case owner's user_id for admin cross-user access
    report_owner_id = case.get("user_id", user.user_id)
    report = await _get_latest_completed_barrister_report(case_id, report_owner_id)
    if not report:
        raise HTTPException(status_code=404, detail="Completed barrister report not found. Generate the Barrister View first.")

    # Get grounds sorted by priority
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id}, {"_id": 0}
    ).sort([("priority_order", 1), ("strength", 1)]).to_list(100)

    # Extract Counsel Synthesis from report content
    analysis = report.get("content", {}).get("analysis", "")
    counsel_synthesis = ""
    # Find the Counsel Synthesis section in the markdown
    import re as re_mod
    synth_match = re_mod.search(r"##\s*Counsel Synthesis(.*?)(?=\n##\s[^#]|\Z)", analysis, re_mod.DOTALL)
    if synth_match:
        counsel_synthesis = synth_match.group(1).strip()

    # Build PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=18*mm, leftMargin=18*mm,
        topMargin=18*mm, bottomMargin=22*mm
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='QBTitle', fontSize=18, spaceAfter=6, alignment=TA_CENTER, fontName='Helvetica-Bold', textColor=colors.HexColor('#0f172a')))
    styles.add(ParagraphStyle(name='QBSubtitle', fontSize=10, spaceAfter=10, alignment=TA_CENTER, textColor=colors.HexColor('#475569')))
    styles.add(ParagraphStyle(name='QBSection', fontSize=13, spaceBefore=10, spaceAfter=4, fontName='Helvetica-Bold', textColor=colors.HexColor('#1e3a8a')))
    styles.add(ParagraphStyle(name='QBSubSection', fontSize=11, spaceBefore=6, spaceAfter=3, fontName='Helvetica-Bold', textColor=colors.HexColor('#1e293b')))
    styles.add(ParagraphStyle(name='QBBody', fontSize=10, spaceAfter=4, alignment=TA_JUSTIFY, leading=13))
    styles.add(ParagraphStyle(name='QBPriority', fontSize=10, spaceAfter=3, leftIndent=8, leading=13))
    styles.add(ParagraphStyle(name='QBDisclaimer', fontSize=8, fontName='Helvetica-Oblique', textColor=colors.HexColor('#94a3b8'), alignment=TA_CENTER, leading=10))
    styles.add(ParagraphStyle(name='QBGroundTitle', fontSize=11, spaceBefore=6, spaceAfter=3, fontName='Helvetica-Bold', textColor=colors.HexColor('#1e3a8a')))

    def safe_text(text):
        return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []

    # Header
    defendant = case.get("defendant_name", "Unknown")
    state = (case.get("state") or "NSW").upper()
    sentence = case.get("sentence", "")

    story.append(Paragraph("BARRISTER QUICK BRIEF", styles['QBTitle']))
    story.append(Paragraph(f"Appellant: {safe_text(defendant)} | Jurisdiction: {state}", styles['QBSubtitle']))
    if sentence:
        story.append(Paragraph(f"Sentence: {safe_text(sentence)}", styles['QBSubtitle']))
    story.append(Spacer(1, 3*mm))

    # Thin blue line separator
    line_table = Table([[""]],
        colWidths=[doc.width],
        rowHeights=[1]
    )
    line_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#2563eb')),
    ]))
    story.append(line_table)
    story.append(Spacer(1, 4*mm))

    # Counsel Synthesis
    story.append(Paragraph("COUNSEL SYNTHESIS", styles['QBSection']))
    if counsel_synthesis:
        # Enforce 2-page contract: truncate synthesis to ~1200 chars max
        MAX_SYNTHESIS_CHARS = 1200
        truncated_synthesis = counsel_synthesis
        if len(counsel_synthesis) > MAX_SYNTHESIS_CHARS:
            truncated_synthesis = counsel_synthesis[:MAX_SYNTHESIS_CHARS].rsplit(".", 1)[0] + ". [See full Barrister View for complete synthesis.]"

        # Parse the synthesis into sub-sections
        sections = re_mod.split(r"###\s+", truncated_synthesis)
        for section in sections:
            section = section.strip()
            if not section:
                continue
            lines = section.split("\n", 1)
            heading = lines[0].strip()
            body = lines[1].strip() if len(lines) > 1 else ""

            if heading:
                story.append(Paragraph(safe_text(heading), styles['QBSubSection']))
            if body:
                # Handle numbered lists and paragraphs
                for para in body.split("\n"):
                    para = para.strip()
                    if not para:
                        continue
                    # Numbered items
                    if re_mod.match(r"^\d+\.", para):
                        story.append(Paragraph(safe_text(para), styles['QBPriority']))
                    else:
                        # Convert markdown bold
                        clean = safe_text(para)
                        clean = re_mod.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", clean)
                        story.append(Paragraph(clean, styles['QBBody']))
    else:
        story.append(Paragraph("Counsel Synthesis not available. Generate a new Barrister View report with the latest prompt to include this section.", styles['QBBody']))

    story.append(Spacer(1, 4*mm))
    story.append(line_table)
    story.append(Spacer(1, 4*mm))

    # Top 3 Grounds
    top_grounds = grounds[:3]
    story.append(Paragraph("TOP 3 GROUNDS OF APPEAL", styles['QBSection']))
    
    VIABILITY_LABELS = {
        "strong": "Arguable \u2014 Strong",
        "moderate": "Arguable \u2014 Moderate",
        "weak": "Requires Development",
    }

    for idx, ground in enumerate(top_grounds, 1):
        title = ground.get("title", "Untitled Ground")
        strength = ground.get("strength", "moderate")
        # Prefer the verified legitimacy_scores.viability_label over raw classifier strength
        viability = (ground.get("legitimacy_scores") or {}).get("viability_label")
        badge_text = viability or f"Unverified strength: {VIABILITY_LABELS.get(strength, strength)}"
        ground_type = ground.get("ground_type", "")
        appellate_pathway = ground.get("appellate_pathway", "")
        description = ground.get("description", "")

        story.append(Paragraph(f"Ground {idx}: {safe_text(title)}", styles['QBGroundTitle']))
        
        # Viability badge — colour keyed to verified label, not raw strength
        viability_colour = {
            "Arguable \u2014 Strong": "#059669",
            "Arguable \u2014 Moderate": "#2563eb",
            "Requires Development": "#dc2626",
        }.get(viability or "", "#64748b")
        story.append(Paragraph(f'<font color="{viability_colour}"><b>{safe_text(badge_text)}</b></font> | Type: {safe_text(ground_type.replace("_", " ").title())}', styles['QBBody']))
        
        if appellate_pathway:
            story.append(Paragraph(f"<b>Appellate Pathway:</b> {safe_text(appellate_pathway)}", styles['QBBody']))

        # Show description (truncated to enforce 2-page contract)
        desc_clean = description.replace("\n\n", " ").replace("\n", " ")
        if len(desc_clean) > 250:
            desc_clean = desc_clean[:250] + "..."
        story.append(Paragraph(safe_text(desc_clean), styles['QBBody']))

        # Contingent warning — prefer legitimacy_scores flag over ground_type alone
        is_contingent = (ground.get("legitimacy_scores") or {}).get("is_contingent", False) or ground_type == "ineffective_counsel"
        if is_contingent:
            story.append(Paragraph('<font color="#d97706"><b>CONTINGENT</b> \u2014 Requires evidentiary support before advancement</font>', styles['QBBody']))

        story.append(Spacer(1, 2*mm))

    if len(grounds) > 3:
        story.append(Paragraph(f"<i>{len(grounds) - 3} additional ground(s) detailed in the full Barrister View report.</i>", styles['QBBody']))

    # Footer disclaimer
    story.append(Spacer(1, 6*mm))
    story.append(line_table)
    story.append(Spacer(1, 3*mm))
    story.append(Paragraph(
        "This document is an AI-generated educational analysis tool only. "
        "It does not constitute legal advice. Independent legal counsel must be obtained "
        "before taking any action. Created and designed by Deb King.",
        styles['QBDisclaimer']
    ))
    story.append(Paragraph(
        f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y at %H:%M UTC')}",
        styles['QBDisclaimer']
    ))

    doc.build(story)
    buffer.seek(0)

    filename = f"Quick_Brief_{defendant.replace(' ', '_')}_{case_id}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@api_router.get("/reports/embedded-legacy", response_model=dict)
async def get_embedded_legacy_reports(request: Request, limit: int = 3):
    """Return strongest historical reports for embedding/reference in UI."""
    user = await get_current_user(request)

    all_reports = await db.reports.find(
        {"user_id": user.user_id},
        {
            "_id": 0,
            "report_id": 1,
            "case_id": 1,
            "report_type": 1,
            "title": 1,
            "generated_at": 1,
            "content.analysis": 1,
        },
    ).sort("generated_at", -1).to_list(400)

    def analysis_len(item: dict) -> int:
        return len((item.get("content") or {}).get("analysis", ""))

    valid_types = ["quick_summary", "full_detailed", "extensive_log"]
    by_length = [r for r in all_reports if r.get("report_type") in valid_types and analysis_len(r) > 1200]
    by_length.sort(key=analysis_len, reverse=True)

    selected = []
    seen_types = set()

    for report in by_length:
        rtype = report.get("report_type")
        if rtype in seen_types:
            continue
        seen_types.add(rtype)
        selected.append(report)
        if len(selected) >= limit:
            break

    if len(selected) < limit:
        selected_ids = {r.get("report_id") for r in selected}
        for report in by_length:
            if report.get("report_id") in selected_ids:
                continue
            selected.append(report)
            if len(selected) >= limit:
                break

    embedded = []
    for report in selected[:limit]:
        analysis = (report.get("content") or {}).get("analysis", "")
        embedded.append(
            {
                "report_id": report.get("report_id"),
                "case_id": report.get("case_id"),
                "report_type": report.get("report_type"),
                "title": report.get("title"),
                "generated_at": report.get("generated_at"),
                "analysis": analysis,
                "analysis_length": len(analysis),
            }
        )

    return {"reports": embedded}

@api_router.get("/cases/{case_id}/reports/{report_id}", response_model=dict)
async def get_report(case_id: str, report_id: str, request: Request):
    """Get a specific report"""
    user = await get_current_user(request)
    
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return report

@api_router.delete("/cases/{case_id}/reports/{report_id}")
async def delete_report(case_id: str, report_id: str, request: Request):
    """Delete a report"""
    user = await get_current_user(request)
    
    result = await db.reports.delete_one({
        "report_id": report_id,
        "case_id": case_id,
        "user_id": user.user_id
    })
    
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Report not found")
    
    return {"message": "Report deleted"}

# ============ PDF EXPORT ============

def _format_export_display_date(value=None) -> str:
    parsed = _coerce_utc_datetime(value)
    if not parsed:
        parsed = datetime.now(timezone.utc)
    return parsed.astimezone(timezone.utc).strftime("%d/%m/%Y")


def _build_export_footer_label(case: dict, report_label: str, generated_at=None) -> str:
    """DO_NOT_UNDO — Footer label with appellant name, report title, and date"""
    appellant_name = (case.get("defendant_name") or case.get("title") or "Appellant").strip()
    report_name = (report_label or "Legal Report").strip()
    return f"Criminal Appeal Case Management — {report_name} — {appellant_name} — {_format_export_display_date(generated_at)}"


def _build_export_footer_message() -> str:
    return ""


def _truncate_export_footer(text: str, max_chars: int = 118) -> str:
    value = (text or "").strip()
    if len(value) <= max_chars:
        return value
    return value[: max_chars - 1].rstrip() + "…"


def _clean_sentence_candidate(value: str) -> str:
    cleaned = (value or "")
    cleaned = re.sub(r"\s*\[.*$", "", cleaned)
    cleaned = re.sub(r"\s*\(https?:.*$", "", cleaned)
    cleaned = re.sub(r"\s*[|•].*$", "", cleaned)
    cleaned = re.sub(r"[,;:]?\s*(?:appeal|conviction|leave|outcome)\b.*$", "", cleaned, flags=re.I)
    cleaned = re.sub(r"[,;:]?\s*(?:dismissed|upheld|refused|granted)\b.*$", "", cleaned, flags=re.I)
    # DO NOT UNDO — Strip crime narrative like "for murdering pregnant girlfriend"
    cleaned = re.sub(r"\s+for\s+(?:the\s+)?(?:murder|killing|manslaughter|assault|robbery|rape|kidnapping|abuse|stabbing|supplying|dealing|trafficking)[a-z\s,'-]*(?=\s+with\b|,|$)", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\s+for\s+[a-z\s'-]+(?=,|\s+with\b|$)", "", cleaned, flags=re.I)
    cleaned = re.sub(r"\bminimum\s+non[- ]?parole\s+period\b", "non-parole period", cleaned, flags=re.I)
    cleaned = re.sub(r"\bwith\s+a\s+minimum\s+non[- ]?parole\s+period\b", "with a non-parole period", cleaned, flags=re.I)
    cleaned = re.sub(r"imprisonment,\s+with", "imprisonment with", cleaned, flags=re.I)
    return re.sub(r"\s+", " ", cleaned).strip()


def _is_valid_sentence_candidate(value: str) -> bool:
    if not value:
        return False
    if not re.search(r"(life|year|month|non[- ]?parole|imprisonment|gaol|custody|sentence)", value, re.I):
        return False
    if re.search(r"\b(reduced|reduce|precedent|appeal|submissions|could|should|potentially|perhaps|would|adequacy|seek|sought|relief)\b", value, re.I):
        return False
    return len(value) < 140


def _extract_sentence_from_text(case: dict, analysis: str) -> str:
    if case.get('sentence') and str(case.get('sentence')).strip():
        return str(case.get('sentence')).strip()
    patterns = [
        r"(?:sentence\s+imposed\s+was|sentence\s+was|head\s+sentence\s+was|head\s+sentence:|sentenced?\s+to)\s+([^\.\n]{8,160})",
        r"(\d+\s+years?'?\s+with\s+a\s+non[- ]?parole\s+period\s+of\s+\d+\s+years?(?:\s+and\s+\d+\s+months?)?)",
        r"(\d+\s+years?'?(?:\s+and\s+\d+\s+months?)?\s*(?:imprisonment|gaol|jail|custody)?\s*(?:with\s+(?:a\s+)?non[- ]?parole\s+period\s+of\s+\d+\s+years?(?:\s+and\s+\d+\s+months?)?)?)",
        r"(?:was\s+)?sentenced?\s+to\s+(\d+\s*(?:years?|months?)\s+(?:and\s+\d+\s*(?:years?|months?)\s+)?(?:imprisonment|gaol|jail|custody)[^\n\.]{0,80})",
        r"(?:^|\n)\s*(?:Head\s+)?Sentence\s*:\s*(\d+[^\n]{5,100})",
        r"(?:sentenced?\s+to\s+)?(life\s+imprisonment|imprisonment\s+for\s+life|life\s+sentence)[^\n\.]{0,60}",
        r"(\d+[\s-]*(?:years?|months?)(?:'s?)?\s*(?:imprisonment|gaol|jail|custody|sentence|non[- ]?parole)[^\n\.]{0,80})",
        r"sentence\s+of\s+(\d+[^\n\.]{5,80})",
        r"((?:minimum|non[- ]?parole)\s+(?:period\s+)?of\s+\d+[^\n\.]{3,60})",
    ]
    for pattern in patterns:
        match = re.search(pattern, analysis or "", re.I | re.M)
        if match:
            candidate = _clean_sentence_candidate(match.group(1))
            if _is_valid_sentence_candidate(candidate):
                return candidate
    return "See report analysis"


async def _derive_export_sentence(case: dict, case_id: str, user_id: str, fallback_report: dict | None = None) -> str:
    standard_reports = await db.reports.find(
        {
            "case_id": case_id,
            "user_id": user_id,
            "report_type": {"$in": ["quick_summary", "full_detailed", "extensive_log", "barrister_view"]},
            "status": "completed",
        },
        {"_id": 0},
    ).sort("generated_at", -1).to_list(20)

    for report_type in ["quick_summary", "full_detailed", "extensive_log", "barrister_view"]:
        for item in [r for r in standard_reports if r.get("report_type") == report_type]:
            candidate = _extract_sentence_from_text(case, ((item.get("content") or {}).get("analysis") or ""))
            if candidate != "See report analysis":
                return candidate

    if fallback_report:
        candidate = _extract_sentence_from_text(case, ((fallback_report.get("content") or {}).get("analysis") or ""))
        if candidate != "See report analysis":
            return candidate

    return _extract_sentence_from_text(case, "")

@api_router.get("/cases/{case_id}/reports/{report_id}/export-pdf")
async def export_report_pdf(case_id: str, report_id: str, request: Request):
    """Export a report as PDF with Grounds of Merit and Legal References"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from io import BytesIO
    
    user = await get_current_user(request)
    
    # Get report
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get grounds of merit for this case
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Create PDF buffer
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=20*mm,
        leftMargin=20*mm,
        topMargin=20*mm,
        bottomMargin=28*mm
    )
    
    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='ReportTitle',
        fontSize=20,
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    ))
    styles.add(ParagraphStyle(
        name='ReportSubtitle',
        fontSize=11,
        spaceAfter=8,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#475569')
    ))
    styles.add(ParagraphStyle(
        name='SectionHeader',
        fontSize=14,
        spaceBefore=14,
        spaceAfter=6,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#0f172a')
    ))
    styles.add(ParagraphStyle(
        name='SubHeader',
        fontSize=12,
        spaceBefore=8,
        spaceAfter=5,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1e293b')
    ))
    styles.add(ParagraphStyle(
        name='ReportBodyText',
        fontSize=10,
        spaceAfter=5,
        alignment=TA_JUSTIFY,
        leading=14
    ))
    styles.add(ParagraphStyle(
        name='BulletText',
        parent=styles['ReportBodyText'],
        leftIndent=12,
        bulletIndent=6
    ))
    styles.add(ParagraphStyle(
        name='LawSection',
        fontSize=10,
        spaceAfter=4,
        leftIndent=18,
        textColor=colors.HexColor('#1e40af')
    ))
    styles.add(ParagraphStyle(
        name='GroundTitle',
        fontSize=12,
        spaceBefore=8,
        spaceAfter=4,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#1e3a8a')
    ))
    styles.add(ParagraphStyle(
        name='NumberedSectionHeader',
        fontSize=13,
        spaceBefore=12,
        spaceAfter=6,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#0f172a')
    ))
    styles.add(ParagraphStyle(
        name='CoverMetaLabel',
        fontSize=9,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#64748b'),
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        name='CoverMetaValue',
        fontSize=11,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#0f172a'),
        spaceAfter=5,
    ))
    # DO_NOT_UNDO — Bold red disclaimer with white text in PDF
    styles.add(ParagraphStyle(
        name='CoverDisclaimer',
        fontSize=10,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor('#dc2626'),
        alignment=TA_CENTER,
        leading=14,
    ))

    def format_inline(text: str) -> str:
        clean = (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Convert markdown bold
        clean = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", clean)
        # Convert markdown links [text](url) to just text
        clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", clean)
        return clean

    def render_table(lines):
        rows = []
        for line in lines:
            parts = [p.strip() for p in line.strip().strip("|").split("|")]
            if not parts:
                continue
            if all(set(p) <= set("-:") for p in parts):
                continue
            # Escape special characters in table cells for reportlab
            safe_parts = []
            for p in parts:
                cell = re.sub(r"\*\*(.*?)\*\*", r"\1", p)
                cell = cell.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                safe_parts.append(cell)
            rows.append(safe_parts)
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        rows = [r + [""] * (col_count - len(r)) for r in rows]
        try:
            col_width = doc.width / col_count
            para_rows = []
            cell_style = ParagraphStyle(name='CellText', fontSize=10, leading=12, fontName='Helvetica', wordWrap='CJK')
            header_style = ParagraphStyle(name='HeaderCellText', fontSize=10, leading=12, fontName='Helvetica-Bold', textColor=colors.white)
            for ri, row in enumerate(rows):
                style = header_style if ri == 0 else cell_style
                para_rows.append([Paragraph(c[:260], style) for c in row])
            table = Table(para_rows, colWidths=[col_width] * col_count, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e293b')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
                ('LEFTPADDING', (0, 0), (-1, -1), 4),
                ('RIGHTPADDING', (0, 0), (-1, -1), 4),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
            ]))
            story.append(table)
            story.append(Spacer(1, 4*mm))
        except Exception as e:
            logger.warning(f"PDF table render failed: {e}")
            for row in rows:
                story.append(Paragraph(" | ".join(row), styles['ReportBodyText']))

    def render_markdown(text: str):
        lines = (text or "").splitlines()
        buffer = []
        table_lines = []

        def flush_paragraph():
            if buffer:
                paragraph = " ".join(buffer).strip()
                if paragraph:
                    try:
                        story.append(Paragraph(format_inline(paragraph), styles['ReportBodyText']))
                        story.append(Spacer(1, 2*mm))
                    except Exception as e:
                        logger.warning(f"PDF paragraph failed: {e}")
                        # Fallback: strip all XML-like content
                        safe = re.sub(r'<[^>]+>', '', paragraph)
                        story.append(Paragraph(safe, styles['ReportBodyText']))
            buffer.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                flush_paragraph()
                continue
            if stripped.startswith("|") and "|" in stripped:
                flush_paragraph()
                table_lines.append(stripped)
                continue
            if table_lines and not (stripped.startswith("|") and "|" in stripped):
                render_table(table_lines)
                table_lines = []
            if stripped.startswith("## "):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped[3:].strip()), styles['SectionHeader']))
                story.append(Spacer(1, 2*mm))
                continue
            if re.match(r'^\d+\.\s+[A-Z][A-Z0-9\s/&()\-]{4,}$', stripped):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped), styles['NumberedSectionHeader']))
                story.append(Spacer(1, 2*mm))
                continue
            if stripped.startswith("### "):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped[4:].strip()), styles['SubHeader']))
                story.append(Spacer(1, 1*mm))
                continue
            if stripped.startswith("#### "):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped[5:].strip()), styles['SubHeader']))
                story.append(Spacer(1, 1*mm))
                continue
            if re.match(r'^Ground\s+\d+\s*:', stripped, re.I):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped), styles['GroundTitle']))
                story.append(Spacer(1, 1*mm))
                continue
            if stripped.startswith("- ") or stripped.startswith("• "):
                flush_paragraph()
                bullet_text = stripped[2:].strip()
                story.append(Paragraph(format_inline(f"- {bullet_text}"), styles['BulletText']))
                story.append(Spacer(1, 1*mm))
                continue
            # Handle numbered lists
            if re.match(r'^\d+\.\s', stripped):
                flush_paragraph()
                story.append(Paragraph(format_inline(stripped), styles['BulletText']))
                story.append(Spacer(1, 1*mm))
                continue
            buffer.append(stripped)

        flush_paragraph()
        if table_lines:
            render_table(table_lines)

    story = []

    # Report Title
    report_type_labels = {
        'quick_summary': 'Quick Case Summary',
        'full_detailed': 'Full Detailed Legal Analysis ($150 AUD)',
        'extensive_log': 'Extensive Case Log & Analysis ($200 AUD)',
        'barrister_view': 'Barrister Brief'
    }
    title = report_type_labels.get(report.get('report_type'), 'Legal Report')
    resolved_sentence = await _derive_export_sentence(case, case_id, user.user_id, report)
    footer_label = _truncate_export_footer(_build_export_footer_label(case, title, report.get('generated_at')))
    footer_message = _build_export_footer_message()

    cover_meta = [
        ['Case Title', case.get('title', 'N/A')],
        ['Defendant', case.get('defendant_name', 'N/A')],
        ['Court / State', f"{case.get('court', 'Court')} — {(case.get('state', 'NSW') or 'NSW').upper()}"],
        ['Report', title],
        ['Offence', case.get('offence_type') or (case.get('offence_category') or 'Not recorded').replace('_', ' ').title()],
        ['Sentence', resolved_sentence],
    ]

    story.append(Spacer(1, 18*mm))
    story.append(Paragraph("Appeal Case Manager", styles['ReportSubtitle']))
    story.append(Paragraph(title, styles['ReportTitle']))
    story.append(Paragraph("Criminal Law Appeal Case Management", styles['ReportSubtitle']))
    story.append(Spacer(1, 10*mm))

    cover_table_rows = []
    for idx in range(0, len(cover_meta), 2):
        left = cover_meta[idx]
        right = cover_meta[idx + 1] if idx + 1 < len(cover_meta) else ["", ""]
        cover_table_rows.append([
            Paragraph(f"<b>{left[0]}</b><br/>{left[1]}", styles['CoverMetaValue']),
            Paragraph(f"<b>{right[0]}</b><br/>{right[1]}", styles['CoverMetaValue'])
        ])
    cover_table = Table(cover_table_rows, colWidths=[80*mm, 80*mm])
    cover_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(cover_table)
    story.append(Spacer(1, 12*mm))
    story.append(Paragraph(
        "IMPORTANT DISCLAIMER — NOT LEGAL ADVICE — This application is an educational research tool only and does NOT constitute legal advice. The creator is not a lawyer. All analysis and recommendations must be independently verified by a qualified Australian legal professional. Australian law only. No solicitor-client relationship is created.",
        styles['CoverDisclaimer']
    ))
    story.append(PageBreak())

    # Header
    story.append(Paragraph("APPEAL CASE MANAGER", styles['ReportTitle']))
    story.append(Paragraph("Criminal Law Appeal Case Management", styles['ReportSubtitle']))
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph(f"Case: {case.get('title', 'Unknown')}", styles['ReportSubtitle']))
    story.append(Spacer(1, 10*mm))

    def draw_page_footer(canvas_obj, pdf_doc):
        canvas_obj.saveState()
        footer_top = 14 * mm
        footer_mid = 10 * mm
        footer_bottom = 6 * mm
        canvas_obj.setStrokeColor(colors.HexColor('#cbd5e1'))
        canvas_obj.setLineWidth(0.6)
        canvas_obj.line(pdf_doc.leftMargin, footer_top, A4[0] - pdf_doc.rightMargin, footer_top)
        canvas_obj.setFillColor(colors.HexColor('#475569'))
        canvas_obj.setFont('Helvetica', 8)
        canvas_obj.drawString(pdf_doc.leftMargin, footer_mid, footer_label)
        canvas_obj.drawRightString(A4[0] - pdf_doc.rightMargin, footer_mid, f"Page {canvas_obj.getPageNumber()}")
        if footer_message:
            canvas_obj.setFillColor(colors.HexColor('#1e3a5f'))
            canvas_obj.setFont('Helvetica-Bold', 8)
            canvas_obj.drawCentredString(A4[0] / 2, footer_bottom, footer_message)
        canvas_obj.restoreState()

    story.append(Paragraph(title, styles['ReportTitle']))
    story.append(Spacer(1, 5*mm))
    
    # Case Info Table — skip N/A fields
    # Get grounds for PDF header
    pdf_grounds = await db.grounds_of_merit.find({"case_id": case_id}, {"_id": 0}).to_list(100)
    case_data_rows = [
        ['Case Title:', case.get('title', 'N/A')],
        ['Defendant:', case.get('defendant_name', 'N/A')],
    ]
    if case.get('case_number') and case.get('case_number') != 'N/A':
        case_data_rows.append(['Case Number:', case['case_number']])
    if case.get('court') and case.get('court') != 'N/A':
        case_data_rows.append(['Court:', case['court']])
    if case.get('state'):
        case_data_rows.append(['Jurisdiction:', case['state'].upper()])
    if resolved_sentence and resolved_sentence != 'See report analysis':
        case_data_rows.append(['Sentence:', resolved_sentence])
    case_data_rows.append(['Grounds:', f"{len(pdf_grounds)} identified"])
    case_data_rows.append(['Generated:', report.get('generated_at', 'N/A')[:10] if report.get('generated_at') else 'N/A'])
    
    case_table = Table(case_data_rows, colWidths=[40*mm, 120*mm])
    case_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#475569')),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(case_table)
    story.append(Spacer(1, 10*mm))
    
    # Grounds of Merit Section
    if grounds:
        story.append(Paragraph("GROUNDS OF MERIT", styles['SectionHeader']))
        story.append(Spacer(1, 5*mm))
        
        ground_type_labels = {
            'procedural_error': 'Procedural Error',
            'fresh_evidence': 'Fresh Evidence',
            'miscarriage_of_justice': 'Miscarriage of Justice',
            'sentencing_error': 'Sentencing Error',
            'judicial_error': 'Judicial Error',
            'ineffective_counsel': 'Ineffective Counsel',
            'prosecution_misconduct': 'Prosecution Misconduct',
            'jury_irregularity': 'Jury Irregularity',
            'constitutional_violation': 'Constitutional Violation',
            'other': 'Other'
        }
        
        for idx, ground in enumerate(grounds, 1):
            ground_type = ground_type_labels.get(ground.get('ground_type'), 'Other')
            strength = ground.get('strength', 'moderate').capitalize()
            
            # Ground header
            story.append(Paragraph(
                f"{idx}. {ground.get('title', 'Unnamed Ground')} [{ground_type}] - Strength: {strength}",
                styles['GroundTitle']
            ))
            
            # Description
            if ground.get('description'):
                story.append(Paragraph(ground.get('description'), styles['ReportBodyText']))
            
            # Legal References (Law Sections)
            if ground.get('law_sections'):
                story.append(Paragraph("<b>Relevant Law Sections:</b>", styles['ReportBodyText']))
                for section in ground.get('law_sections', []):
                    section_text = f"• s.{section.get('section', '')} {section.get('act', '')} ({section.get('jurisdiction', 'NSW')})"
                    story.append(Paragraph(section_text.replace('•', '-'), styles['LawSection']))
            
            # Similar Cases
            if ground.get('similar_cases'):
                story.append(Paragraph("<b>Similar Cases:</b>", styles['ReportBodyText']))
                for case_ref in ground.get('similar_cases', []):
                    case_text = f"• {case_ref.get('case_name', '')} {case_ref.get('citation', '')}"
                    story.append(Paragraph(case_text.replace('•', '-'), styles['LawSection']))
            
            # Supporting Evidence
            if ground.get('supporting_evidence'):
                story.append(Paragraph("<b>Supporting Evidence:</b>", styles['ReportBodyText']))
                for evidence in ground.get('supporting_evidence', []):
                    story.append(Paragraph(f"- {evidence}", styles['LawSection']))
            
            story.append(Spacer(1, 5*mm))
    
    # Legal Framework Reference
    story.append(Paragraph("LEGAL FRAMEWORK REFERENCE", styles['SectionHeader']))
    story.append(Spacer(1, 2*mm))
    legal_refs = [
        "- Crimes Act 1900 (NSW) - Primary criminal law for NSW",
        "- Criminal Appeal Act 1912 (NSW) - Governs appeals in NSW",
        "- Criminal Code Act 1995 (Cth) - Federal criminal law",
        "- Evidence Act 1995 (NSW & Cth) - Evidence admissibility",
        "- Sentencing Act 1995 (NSW) - Sentencing guidelines"
    ]
    for ref in legal_refs:
        story.append(Paragraph(ref, styles['LawSection']))
    
    story.append(Spacer(1, 10*mm))
    
    # Main Analysis Content
    story.append(Paragraph("DETAILED ANALYSIS", styles['SectionHeader']))

    analysis_text = _strip_report_placeholders(report.get('content', {}).get('analysis', 'No analysis available.'))
    render_markdown(analysis_text)

    # DO_NOT_UNDO — Bold red disclaimer in PDF body
    story.append(Spacer(1, 15*mm))
    story.append(Spacer(1, 5*mm))
    story.append(Paragraph(
        "IMPORTANT DISCLAIMER — NOT LEGAL ADVICE",
        ParagraphStyle(name='DisclaimerTitle', fontSize=12, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=colors.HexColor('#dc2626'), spaceAfter=4)
    ))
    story.append(Paragraph(
        "This application is an educational research tool only and does NOT constitute legal advice. It must NOT be relied upon as such. "
        "The creator of this application is not a lawyer. All analysis, findings, reports, and recommendations generated by this tool must be independently verified "
        "by a qualified Australian legal professional before any action is taken. This tool covers Australian law only. "
        "No solicitor-client relationship is created by using this service.",
        ParagraphStyle(name='DisclaimerBody', fontSize=10, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=colors.HexColor('#dc2626'), leading=14)
    ))
    
    # Build PDF
    try:
        doc.build(story, onFirstPage=draw_page_footer, onLaterPages=draw_page_footer)
    except Exception as e:
        logger.error(f"PDF build failed: {e}")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)[:200]}")
    buffer.seek(0)
    
    # Create filename
    safe_title = "".join(c for c in case.get('title', 'Report')[:30] if c.isalnum() or c in ' -_').strip()
    filename = f"{safe_title}_{report.get('report_type', 'report')}.pdf"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

@api_router.get("/cases/{case_id}/reports/{report_id}/export-docx")
async def export_report_docx(case_id: str, report_id: str, request: Request):
    """Export a report as DOCX (Microsoft Word) with Grounds of Merit and Legal References"""
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from io import BytesIO
    
    user = await get_current_user(request)
    
    # Get report
    report = await db.reports.find_one(
        {"report_id": report_id, "case_id": case_id, "user_id": user.user_id},
        {"_id": 0}
    )
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Get case data
    case = await db.cases.find_one({"case_id": case_id, "user_id": user.user_id}, {"_id": 0})
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    # Get grounds of merit
    grounds = await db.grounds_of_merit.find(
        {"case_id": case_id},
        {"_id": 0}
    ).to_list(100)
    
    # Create DOCX document
    doc = DocxDocument()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(30, 41, 59)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(30, 58, 138)
    
    # Header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run("APPEAL CASE MANAGER")
    header_run.font.size = Pt(12)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(15, 23, 42)

    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_header.add_run("Criminal Law Appeal Case Management")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    case_line = doc.add_paragraph()
    case_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    case_run = case_line.add_run(f"Case: {case.get('title', 'Unknown')}")
    case_run.font.size = Pt(10)
    case_run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()
    
    # Report Title
    report_type_labels = {
        'quick_summary': 'Quick Case Summary',
        'full_detailed': 'Full Detailed Legal Analysis', 
        'extensive_log': 'Extensive Case Log & Analysis',
        'barrister_view': 'Barrister Brief'
    }
    report_title = report_type_labels.get(report.get('report_type'), 'Legal Report')
    resolved_sentence = await _derive_export_sentence(case, case_id, user.user_id, report)
    footer_label = _truncate_export_footer(_build_export_footer_label(case, report_title, report.get('generated_at')))
    footer_message = _build_export_footer_message()

    def add_page_number_field(paragraph):
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' PAGE '
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)

    for section in doc.sections:
        section.footer_distance = Inches(0.35)
        footer = section.footer
        footer_line = footer.paragraphs[0]
        footer_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_line_run = footer_line.add_run(f"{footer_label} — Page ")
        footer_line_run.font.size = Pt(8)
        footer_line_run.font.color.rgb = RGBColor(71, 85, 105)
        add_page_number_field(footer_line)

        if footer_message:
            footer_msg_para = footer.add_paragraph()
            footer_msg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_msg_run = footer_msg_para.add_run(footer_message)
            footer_msg_run.font.size = Pt(8)
            footer_msg_run.font.bold = True
            footer_msg_run.font.color.rgb = RGBColor(30, 58, 95)

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title_run = cover_title.add_run(report_title)
    cover_title_run.font.size = Pt(18)
    cover_title_run.font.bold = True
    cover_title_run.font.color.rgb = RGBColor(15, 23, 42)

    cover_subtitle = doc.add_paragraph()
    cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_subtitle_run = cover_subtitle.add_run("Appeal Case Manager")
    cover_subtitle_run.font.size = Pt(10)
    cover_subtitle_run.font.bold = True
    cover_subtitle_run.font.color.rgb = RGBColor(29, 78, 216)

    cover_case = doc.add_paragraph()
    cover_case.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_case_run = cover_case.add_run(case.get('title', 'Unknown case'))
    cover_case_run.font.size = Pt(11)
    cover_case_run.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph()
    cover_info = [
        ('Case Title', case.get('title', 'N/A')),
        ('Defendant', case.get('defendant_name', 'N/A')),
        ('Court / State', f"{case.get('court', 'Court')} — {(case.get('state', 'NSW') or 'NSW').upper()}"),
        ('Report', report_title),
        ('Offence', case.get('offence_type') or (case.get('offence_category') or 'Not recorded').replace('_', ' ').title()),
        ('Sentence', resolved_sentence),
    ]
    cover_table = doc.add_table(rows=len(cover_info), cols=2)
    cover_table.style = 'Table Grid'
    for row_idx, (label, value) in enumerate(cover_info):
        row = cover_table.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].font.bold = True

    cover_disclaimer = doc.add_paragraph()
    cover_disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_disclaimer_run = cover_disclaimer.add_run(
        "IMPORTANT DISCLAIMER — NOT LEGAL ADVICE — This application is an educational research tool only and does NOT constitute legal advice. The creator is not a lawyer. All analysis and recommendations must be independently verified by a qualified Australian legal professional. Australian law only. No solicitor-client relationship is created."
    )
    # DO_NOT_UNDO — Cover page disclaimer in bold red
    cover_disclaimer_run.font.size = Pt(10)
    cover_disclaimer_run.font.bold = True
    cover_disclaimer_run.font.color.rgb = RGBColor(220, 38, 38)

    doc.add_page_break()

    title = doc.add_heading(report_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Case Information Table — skip N/A fields
    case_info = [
        ('Case Title:', case.get('title', 'N/A')),
        ('Defendant:', case.get('defendant_name', 'N/A')),
    ]
    if case.get('case_number') and case.get('case_number') != 'N/A':
        case_info.append(('Case Number:', case['case_number']))
    if case.get('court') and case.get('court') != 'N/A':
        case_info.append(('Court:', case['court']))
    if case.get('state'):
        case_info.append(('Jurisdiction:', case['state'].upper()))
    if resolved_sentence and resolved_sentence != 'See report analysis':
        case_info.append(('Sentence:', resolved_sentence))
    case_info.append(('Grounds:', f"{len(grounds)} identified"))
    case_info.append(('Generated:', report.get('generated_at', 'N/A')[:10] if report.get('generated_at') else 'N/A'))
    
    case_table = doc.add_table(rows=len(case_info), cols=2)
    case_table.style = 'Table Grid'
    
    for i, (label, value) in enumerate(case_info):
        row = case_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Grounds of Merit Section
    if grounds:
        doc.add_heading('GROUNDS OF MERIT', level=1)
        
        ground_type_labels = {
            'procedural_error': 'Procedural Error',
            'fresh_evidence': 'Fresh Evidence',
            'miscarriage_of_justice': 'Miscarriage of Justice',
            'sentencing_error': 'Sentencing Error',
            'judicial_error': 'Judicial Error',
            'ineffective_counsel': 'Ineffective Counsel',
            'prosecution_misconduct': 'Prosecution Misconduct',
            'jury_irregularity': 'Jury Irregularity',
            'constitutional_violation': 'Constitutional Violation',
            'other': 'Other'
        }
        
        for idx, ground in enumerate(grounds, 1):
            ground_type = ground_type_labels.get(ground.get('ground_type'), 'Other')
            strength = ground.get('strength', 'moderate').capitalize()
            
            # Ground heading
            doc.add_heading(
                f"{idx}. {ground.get('title', 'Unnamed Ground')} [{ground_type}] - Strength: {strength}",
                level=2
            )
            
            # Description
            if ground.get('description'):
                doc.add_paragraph(ground.get('description'))
            
            # Legal References
            if ground.get('law_sections'):
                law_para = doc.add_paragraph()
                law_run = law_para.add_run('Relevant Law Sections:')
                law_run.font.bold = True
                law_run.font.color.rgb = RGBColor(30, 64, 175)
                
                for section in ground.get('law_sections', []):
                    section_text = f"s.{section.get('section', '')} {section.get('act', '')} ({section.get('jurisdiction', 'NSW')})"
                    bullet = doc.add_paragraph(section_text, style='List Bullet')
                    for run in bullet.runs:
                        run.font.color.rgb = RGBColor(30, 64, 175)
            
            # Similar Cases
            if ground.get('similar_cases'):
                cases_para = doc.add_paragraph()
                cases_run = cases_para.add_run('Similar Cases:')
                cases_run.font.bold = True
                cases_run.font.color.rgb = RGBColor(30, 58, 138)
                
                for case_ref in ground.get('similar_cases', []):
                    case_text = f"{case_ref.get('case_name', '')} {case_ref.get('citation', '')}"
                    doc.add_paragraph(case_text, style='List Bullet')
            
            # Supporting Evidence
            if ground.get('supporting_evidence'):
                evidence_para = doc.add_paragraph()
                evidence_run = evidence_para.add_run('Supporting Evidence:')
                evidence_run.font.bold = True
                evidence_run.font.color.rgb = RGBColor(5, 150, 105)
                
                for evidence in ground.get('supporting_evidence', []):
                    doc.add_paragraph(evidence, style='List Bullet')
            
            doc.add_paragraph()
    
    # Legal Framework Reference
    doc.add_heading('LEGAL FRAMEWORK REFERENCE', level=1)
    
    legal_refs = [
        "Crimes Act 1900 (NSW) - Primary criminal law for New South Wales",
        "Criminal Appeal Act 1912 (NSW) - Governs criminal appeals in NSW",
        "Criminal Code Act 1995 (Cth) - Federal criminal law",
        "Evidence Act 1995 (NSW & Cth) - Rules on evidence admissibility",
        "Sentencing Act 1995 (NSW) - Sentencing guidelines and procedures"
    ]
    
    for ref in legal_refs:
        doc.add_paragraph(ref, style='List Bullet')
    
    doc.add_paragraph()
    
    # Detailed Analysis
    doc.add_heading('DETAILED ANALYSIS', level=1)

    def add_table_docx(lines):
        rows = []
        for line in lines:
            parts = [p.strip() for p in line.strip().strip("|").split("|")]
            if not parts:
                continue
            if all(set(p) <= set("-:") for p in parts):
                continue
            rows.append(parts)
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        table.autofit = False
        available_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
        col_width = int(available_width / max(1, len(rows[0])))
        for r_idx, row in enumerate(rows):
            for c_idx, value in enumerate(row):
                cell = table.cell(r_idx, c_idx)
                cell.width = col_width
                cell.text = value.replace('**', '')
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                else:
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
        doc.add_paragraph()

    def render_markdown_docx(text):
        lines = (text or "").splitlines()
        buffer = []
        table_lines = []

        def flush_paragraph():
            if buffer:
                paragraph = " ".join(buffer).replace('**', '').strip()
                if paragraph:
                    doc.add_paragraph(paragraph)
            buffer.clear()

        for line in lines:
            stripped = line.strip()
            if not stripped:
                flush_paragraph()
                continue
            if stripped.startswith("|") and "|" in stripped:
                flush_paragraph()
                table_lines.append(stripped)
                continue
            if table_lines and not (stripped.startswith("|") and "|" in stripped):
                add_table_docx(table_lines)
                table_lines = []
            if stripped.startswith("## "):
                flush_paragraph()
                doc.add_heading(stripped[3:].strip(), level=1)
                continue
            if stripped.startswith("### "):
                flush_paragraph()
                doc.add_heading(stripped[4:].strip(), level=2)
                continue
            if stripped.startswith("- ") or stripped.startswith("• "):
                flush_paragraph()
                doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
                continue
            buffer.append(stripped)

        flush_paragraph()
        if table_lines:
            add_table_docx(table_lines)

    analysis_text = _strip_report_placeholders(report.get('content', {}).get('analysis', 'No analysis available.'))
    render_markdown_docx(analysis_text)
    
    doc.add_paragraph()
    
    # DO_NOT_UNDO — Bold red disclaimer with white text
    disc_title = doc.add_paragraph()
    disc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_run = disc_title.add_run("NOT LEGAL ADVICE")
    disc_run.font.size = Pt(12)
    disc_run.font.bold = True
    disc_run.font.color.rgb = RGBColor(220, 38, 38)
    
    disc_body = doc.add_paragraph()
    disc_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_body_run = disc_body.add_run(
        "IMPORTANT DISCLAIMER — This application is an educational research tool only and does NOT constitute legal advice. It must NOT be relied upon "
        "as such. The creator of this application is not a lawyer. All analysis, findings, reports, and recommendations generated "
        "by this tool must be independently verified by a qualified Australian legal professional before any action is taken. "
        "This tool covers Australian law only. No solicitor-client relationship is created by using this service."
    )
    disc_body_run.font.size = Pt(10)
    disc_body_run.font.bold = True
    disc_body_run.font.color.rgb = RGBColor(220, 38, 38)
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create filename
    safe_title = "".join(c for c in case.get('title', 'Report')[:30] if c.isalnum() or c in ' -_').strip()
    filename = f"{safe_title}_{report.get('report_type', 'report')}.docx"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )

# ============ HEALTH CHECK ============

@api_router.get("/")
async def root():
    return {"message": "Criminal Appeal AI API", "status": "operational"}



# ── Include api_router (report/export endpoints defined above) ──
app.include_router(api_router)

# ── Include routers ──
app.include_router(cases_router)
app.include_router(auth_router)
app.include_router(password_reset_router)
app.include_router(admin_router)
app.include_router(utilities_router)
app.include_router(analytics_router)
app.include_router(statistics_router)
app.include_router(compare_router)
app.include_router(contradictions_router)
app.include_router(export_router)
app.include_router(translate_router)
app.include_router(collaboration_router)
app.include_router(documents_router)
app.include_router(timeline_router)
app.include_router(deadlines_router)
app.include_router(notes_router)
app.include_router(grounds_router)
app.include_router(payments_router)
app.include_router(resources_router)
app.include_router(analysis_router)
app.include_router(pipeline_router)
app.include_router(pipeline_staged_router)
app.include_router(caselaw_router)

# DO_NOT_UNDO — CORS Middleware. Uses CORS_ORIGINS env var for allowed origins.
# Must include ALL domains the frontend is served from (preview, production, custom domain).
_cors_origins_raw = os.environ.get("CORS_ORIGINS", "")
_frontend_url = os.environ.get("FRONTEND_URL", "").replace("/api", "")
_all_origins = set()
if _cors_origins_raw.strip() == "*":
    _all_origins = {"*"}
else:
    for src in [_cors_origins_raw, _frontend_url]:
        for o in src.split(","):
            o = o.strip().strip('"').strip("'")
            if o:
                _all_origins.add(o)
if not _all_origins:
    logger.warning("CORS: No origins configured — set CORS_ORIGINS or FRONTEND_URL in .env")
app.add_middleware(
    CORSMiddleware,
    allow_credentials=True if "*" not in _all_origins else False,
    allow_origins=list(_all_origins),
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "Accept"],
)


@app.on_event("startup")
async def cleanup_orphaned_reports():
    """Auto-fail or recover reports stuck in 'generating' from server restarts.
    
    DO_NOT_UNDO — Recovery uses minimum character targets to decide whether a partial
    report is complete enough to mark as finished, or needs re-generation.
    """
    # ─── Database Initialisation ────────────────────────────────────────
    # Create indexes for ALL collections used by the app.
    # MongoDB creates collections automatically on first write, but indexes
    # must be ensured at startup to guarantee query performance in deployment.
    
    # Core collections
    await db.users.create_index([("user_id", 1)], unique=True)
    await db.users.create_index([("email", 1)], unique=True)
    await db.cases.create_index([("case_id", 1)], unique=True)
    await db.cases.create_index([("user_id", 1)])
    await db.reports.create_index([("report_id", 1)], unique=True)
    await db.reports.create_index([("case_id", 1), ("user_id", 1)])
    await db.reports.create_index([("case_id", 1), ("report_type", 1)])
    await db.documents.create_index([("document_id", 1)], unique=True)
    await db.documents.create_index([("case_id", 1), ("user_id", 1)])
    
    # Grounds and analysis
    await db.grounds_of_merit.create_index([("ground_id", 1)], unique=True)
    await db.grounds_of_merit.create_index([("case_id", 1), ("user_id", 1)])
    await db.grounds_of_merit.create_index([("case_id", 1), ("priority_order", 1)])
    await db.issue_arguments.create_index([("case_id", 1)])
    
    # Pipeline collections
    await db.document_extracts.create_index([("case_id", 1), ("user_id", 1)])
    await db.document_extracts.create_index([("extract_id", 1)], unique=True)
    await db.document_extracts.create_index([("document_id", 1), ("case_id", 1)])
    await db.case_extracts.create_index([("case_id", 1), ("user_id", 1)])
    await db.case_extracts.create_index([("case_extract_id", 1)], unique=True)
    await db.issue_classifications.create_index([("case_id", 1), ("user_id", 1)])
    await db.issue_classifications.create_index([("issue_id", 1)], unique=True)
    await db.issue_verifications.create_index([("issue_id", 1), ("case_id", 1)])
    await db.issue_verifications.create_index([("verification_id", 1)], unique=True)
    await db.pipeline_tasks.create_index([("case_id", 1), ("user_id", 1), ("task_type", 1)])
    await db.pipeline_tasks.create_index([("task_id", 1)], unique=True)
    
    # Auth and sessions
    await db.user_sessions.create_index([("session_token", 1)], unique=True)
    await db.user_sessions.create_index([("user_id", 1)])
    await db.user_sessions.create_index([("expires_at", 1)], expireAfterSeconds=0)
    await db.password_reset_tokens.create_index([("token", 1)], unique=True)
    await db.password_reset_tokens.create_index([("expires_at", 1)], expireAfterSeconds=0)
    
    # Case features
    await db.notes.create_index([("case_id", 1), ("user_id", 1)])
    await db.timeline_events.create_index([("case_id", 1)])
    await db.deadlines.create_index([("case_id", 1), ("user_id", 1)])
    await db.checklist_items.create_index([("case_id", 1), ("user_id", 1)])
    await db.submissions_drafts.create_index([("case_id", 1), ("user_id", 1)])
    await db.activities.create_index([("case_id", 1)])
    await db.contradiction_scans.create_index([("case_id", 1)])
    
    # Payments
    await db.payments.create_index([("user_id", 1)])
    await db.payments.create_index([("case_id", 1)])
    await db.payments.create_index([("payment_id", 1)], unique=True)
    
    # Sharing
    await db.case_shares.create_index([("case_id", 1)])
    await db.share_links.create_index([("link_id", 1)], unique=True)
    
    # Analytics
    await db.visits.create_index([("timestamp", 1)])
    await db.visit_stats.create_index([("date", 1)])
    await db.contact_messages.create_index([("created_at", 1)])
    
    # Notifications
    await db.notifications.create_index([("user_id", 1), ("read", 1)])
    await db.case_messages.create_index([("case_id", 1)])
    
    # Counters
    await db.counters.create_index([("name", 1)], unique=True)

    # Minimum character targets — reports below these are incomplete
    min_recovery_targets = {
        "quick_summary": 10000,
        "full_detailed": 70000,
        "extensive_log": 120000,
    }

    five_min_ago = (datetime.now(timezone.utc) - timedelta(minutes=5)).isoformat()
    async for report in db.reports.find({"status": "generating", "generated_at": {"$lt": five_min_ago}}):
        partial = report.get("content", {}).get("analysis", "") or report.get("content", {}).get("partial_analysis", "")
        report_type = report.get("report_type", "quick_summary")
        min_target = min_recovery_targets.get(report_type, 10000)
        
        if partial and len(partial) > 5000:
            if len(partial) >= min_target:
                # Full recovery — report meets minimum target
                await db.reports.update_one(
                    {"report_id": report["report_id"]},
                    {"$set": {"status": "completed", "content.analysis": partial, "content.partial": False}}
                )
                logger.info(f"Recovered complete report {report['report_id']} ({len(partial)} chars)")
            else:
                # Partial recovery — save content but mark as failed so user can regenerate (resume)
                await db.reports.update_one(
                    {"report_id": report["report_id"]},
                    {"$set": {
                        "status": "failed",
                        "error": f"Generation interrupted after {report.get('content', {}).get('passes_completed', '?')}/8 passes ({len(partial)} chars). Click Generate to resume from where it stopped.",
                        "content.analysis": partial,
                        "content.partial": True,
                        "content.partial_analysis": partial,
                    }}
                )
                logger.info(f"Partial report {report['report_id']} below target ({len(partial)} < {min_target}), marked as failed for resume")
        else:
            # DO_NOT_UNDO — Restore from backup if available. When a user regenerates
            # a completed report and the generation fails, restore the old content
            # so they never lose their existing report.
            backup = report.get("content", {}).get("backup_analysis", "")
            if backup and len(backup) > 5000:
                await db.reports.update_one(
                    {"report_id": report["report_id"]},
                    {"$set": {
                        "status": "completed",
                        "error": None,
                        "content.analysis": backup,
                        "content.partial": False,
                    },
                    "$unset": {"content.backup_analysis": 1}}
                )
                logger.info(f"Restored report {report['report_id']} from backup ({len(backup)} chars)")
            else:
                await db.reports.update_one(
                    {"report_id": report["report_id"]},
                    {"$set": {"status": "failed", "error": "Generation interrupted by server restart. Please try again."}}
                )
                logger.info(f"Failed orphaned report {report['report_id']}")

    # ── FLAG EXISTING UNDERSIZED "COMPLETED" REPORTS (non-destructive) ──
    # Add a flag to undersized reports so the UI can show a "Regenerate for full depth" option.
    # DO NOT change status from "completed" — the user must still be able to VIEW their existing reports.
    min_completed_targets = {
        "full_detailed": 70000,
        "extensive_log": 120000,
    }
    flagged_count = 0
    for rtype, min_chars in min_completed_targets.items():
        async for report in db.reports.find({"status": "completed", "report_type": rtype}):
            analysis = (report.get("content", {}).get("analysis") or "")
            if len(analysis) < min_chars:
                await db.reports.update_one(
                    {"report_id": report["report_id"]},
                    {"$set": {"content.below_target": True, "content.actual_chars": len(analysis), "content.target_chars": min_chars}}
                )
                flagged_count += 1
    if flagged_count:
        logger.info(f"Flagged {flagged_count} undersized reports on startup")

    # Also restore any reports that were accidentally set to "failed" by a previous migration
    # DO_NOT_UNDO — Only restore if the report actually meets the minimum target.
    # Reports below the target MUST stay failed so the user can regenerate them.
    for rtype, min_chars in min_completed_targets.items():
        async for report in db.reports.find({"status": "failed", "report_type": rtype}):
            analysis = (report.get("content", {}).get("analysis") or report.get("content", {}).get("partial_analysis") or "")
            if analysis and len(analysis) >= min_chars:
                await db.reports.update_one(
                    {"report_id": report["report_id"]},
                    {"$set": {
                        "status": "completed",
                        "content.analysis": analysis,
                        "content.partial": False,
                        "content.below_target": False,
                        "content.actual_chars": len(analysis),
                        "content.target_chars": min_chars,
                        "error": None,
                    }}
                )
                logger.info(f"Restored report {report['report_id']} to completed ({len(analysis)} chars >= {min_chars} target)")


@app.on_event("startup")
async def startup_dedup_grounds():
    """DO_NOT_UNDO — Auto-cleanup duplicate grounds on every server start.
    
    Runs the fuzzy deduplication cleanup across ALL cases to merge any duplicates
    that slipped through before the dedup logic was fully applied.
    This prevents the recurring 'grounds multiplying' bug from ever persisting.
    """
    from services.ground_dedup import cleanup_duplicate_grounds, cleanup_duplicate_issues

    try:
        # Get all unique (case_id, user_id) pairs that have grounds
        pipeline = [
            {"$group": {"_id": {"case_id": "$case_id", "user_id": "$user_id"}}},
        ]
        case_pairs = await db.grounds_of_merit.aggregate(pipeline).to_list(500)

        total_removed = 0
        for pair in case_pairs:
            cid = pair["_id"]["case_id"]
            uid = pair["_id"]["user_id"]

            # Cleanup grounds
            result = await cleanup_duplicate_grounds(db, cid, uid)
            if result["removed"] > 0:
                total_removed += result["removed"]

            # Also cleanup issue_classifications
            await cleanup_duplicate_issues(db, cid, uid)

        if total_removed > 0:
            logger.info(f"Startup dedup: removed {total_removed} duplicate grounds across {len(case_pairs)} cases")
        else:
            logger.info(f"Startup dedup: no duplicates found across {len(case_pairs)} cases")
    except Exception as e:
        logger.error(f"Startup dedup failed (non-fatal): {e}")


@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()


# ── Serve frontend static files (Docker single-container deploy) ──
# Must be AFTER all API routers so /api/* routes take priority
import pathlib as _pathlib
_static_dir = _pathlib.Path(__file__).parent / "static"
if _static_dir.is_dir():
    from starlette.staticfiles import StaticFiles as _StaticFiles
    app.mount("/", _StaticFiles(directory=str(_static_dir), html=True), name="static")
