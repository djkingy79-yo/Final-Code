"""
Criminal Appeal AI - Document Processing Helpers
Handles text extraction (PDF, DOCX, TXT, images) and context building for AI prompts.

Key improvements:
  - Page-level provenance for PDF extraction
  - Structured extraction metadata returned alongside text
  - OCR preprocessing for better accuracy on poor scans
  - Document context builder includes provenance headers per chunk
  - build_document_context ranks by category relevance
"""
import base64
import io
import logging
import re
from pypdf import PdfReader
from PIL import Image, ImageFilter, ImageEnhance
import pytesseract
from docx import Document
from typing import List, Tuple

logger = logging.getLogger(__name__)

MENTION_PATTERN = re.compile(r"@([A-Za-z0-9._-]{2,64})")

# Document categories ranked by appellate relevance
CATEGORY_PRIORITY = {
    "sentencing_remarks": 1,
    "judgement": 1,
    "transcript": 2,
    "directions": 2,
    "summing_up": 2,
    "expert_report": 3,
    "psychiatric_report": 3,
    "medical_report": 3,
    "witness_statement": 4,
    "police": 4,
    "correspondence": 5,
    "other": 6,
}


def build_document_context(
    documents: list,
    per_doc_char_limit: int = 2000,
    total_char_budget: int = 15000,
    include_description: bool = True,
    content_heading: str = "CONTENT",
) -> dict:
    """
    Build document context string for AI prompts with provenance metadata.

    Documents are sorted by category priority (judgments, transcripts first).
    Each chunk carries a provenance header (filename, category, extraction info).

    Returns: {
        "text": str,
        "included_docs": int,
        "omitted_docs": int,
        "total_docs": int,
        "documents_metadata": [{"filename": ..., "category": ..., "included": bool, "chars_included": int}]
    }
    """
    if not documents:
        return {"text": "No documents available.\n", "included_docs": 0, "omitted_docs": 0, "total_docs": 0, "documents_metadata": []}

    # Sort by category priority — high-value legal documents first
    def sort_key(doc):
        category = (doc.get("category") or "other").lower().replace(" ", "_")
        return CATEGORY_PRIORITY.get(category, 6)

    sorted_docs = sorted(documents, key=sort_key)

    parts = []
    total_chars = 0
    included = 0
    omitted = 0
    metadata = []

    for doc in sorted_docs:
        content = doc.get("content_text", "")
        filename = doc.get("filename", "Unknown")
        category = doc.get("category", "other")
        description = doc.get("description", "")
        pages_processed = doc.get("pages_processed", "unknown")

        if not content or not content.strip():
            metadata.append({"filename": filename, "category": category, "included": False, "chars_included": 0, "reason": "no text"})
            omitted += 1
            continue

        allowed_chars = min(per_doc_char_limit, total_char_budget - total_chars)
        if allowed_chars <= 200:
            metadata.append({"filename": filename, "category": category, "included": False, "chars_included": 0, "reason": "budget exhausted"})
            omitted += 1
            continue

        # Truncate at sentence boundary where possible
        snippet = content[:allowed_chars]
        if len(content) > allowed_chars:
            last_period = snippet.rfind('.')
            if last_period > allowed_chars * 0.7:
                snippet = snippet[:last_period + 1]
            snippet += f"\n[... truncated — {len(content) - len(snippet)} chars remaining ...]"

        # Build provenance header for this document
        provenance_parts = [f"--- DOCUMENT: {filename}"]
        if category and category != "other":
            provenance_parts.append(f"category: {category}")
        if doc.get("ocr_used"):
            provenance_parts.append("extraction: OCR")
        if pages_processed and pages_processed != "unknown":
            provenance_parts.append(f"pages: {pages_processed}")
        provenance_header = " | ".join(provenance_parts) + " ---"

        doc_block = f"\n{provenance_header}\n"
        if include_description and description:
            doc_block += f"Description: {description}\n"
        doc_block += f"{content_heading}:\n{snippet}\n"

        parts.append(doc_block)
        chars_used = len(snippet)
        total_chars += chars_used
        included += 1
        metadata.append({"filename": filename, "category": category, "included": True, "chars_included": chars_used})

    return {
        "text": "\n".join(parts),
        "included_docs": included,
        "omitted_docs": omitted,
        "total_docs": len(documents),
        "documents_metadata": metadata,
    }


def extract_text_with_ocr(file_data_base64: str, filename: str, file_type: str) -> Tuple[str, bool]:
    """
    Extract text from a document file.

    Returns: (extracted_text, ocr_used)

    Supports: PDF (native + OCR fallback), DOCX (paragraphs + tables), TXT, images.
    Preserves page boundaries where possible.
    """
    ocr_used = False
    extracted_text = ""

    try:
        file_bytes = base64.b64decode(file_data_base64)
        # Normalise file_type: handle both MIME types (image/jpeg) and extensions (jpg)
        file_ext = file_type.lower().replace(".", "")
        if "/" in file_ext:
            file_ext = file_ext.split("/")[-1]  # image/jpeg -> jpeg, application/pdf -> pdf
        if file_ext == "vnd.openxmlformats-officedocument.wordprocessingml.document":
            file_ext = "docx"

        if file_ext == "pdf":
            extracted_text, ocr_used = _extract_pdf(file_bytes, filename)
        elif file_ext in ("docx", "doc"):
            extracted_text = _extract_docx(file_bytes, filename)
        elif file_ext in ("png", "jpg", "jpeg", "tiff", "bmp", "gif"):
            extracted_text = _extract_image(file_bytes, filename)
            ocr_used = True
        elif file_ext == "txt":
            extracted_text = _extract_txt(file_bytes)
        else:
            extracted_text = f"[Unsupported file type: {file_ext}]"

    except Exception as e:
        logger.error(f"Extraction failed for {filename}: {e}")
        extracted_text = f"[Extraction error for {filename}: {str(e)[:200]}]"

    return extracted_text, ocr_used


def _extract_pdf(file_bytes: bytes, filename: str) -> Tuple[str, bool]:
    """Extract text from PDF with page-level provenance."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        total_pages = len(reader.pages)
        max_pages = min(total_pages, 50)  # Cap at 50 pages
        text_parts = []
        native_text_found = False

        for page_num in range(max_pages):
            page = reader.pages[page_num]
            page_text = (page.extract_text() or "").strip()
            if page_text and len(page_text) > 20:
                text_parts.append(f"--- Page {page_num + 1}/{total_pages} ---\n{page_text}")
                native_text_found = True
            else:
                text_parts.append(f"--- Page {page_num + 1}/{total_pages} --- [no native text]")

        if total_pages > max_pages:
            text_parts.append(f"\n[Note: {total_pages - max_pages} additional page(s) not extracted — total document is {total_pages} pages]")

        if not native_text_found:
            # Fall back to OCR for scanned PDFs using pdf2image if available
            logger.info(f"No native text in {filename}, attempting OCR")
            try:
                from pdf2image import convert_from_bytes
                ocr_max = min(total_pages, 30)
                images = convert_from_bytes(file_bytes, first_page=1, last_page=ocr_max, dpi=200)
                ocr_parts = []
                for idx, img in enumerate(images):
                    img = _preprocess_for_ocr(img)
                    page_text = pytesseract.image_to_string(img, lang='eng', config='--psm 6')
                    if page_text.strip():
                        ocr_parts.append(f"--- Page {idx + 1}/{total_pages} (OCR) ---\n{page_text}")
                    else:
                        ocr_parts.append(f"--- Page {idx + 1}/{total_pages} --- [OCR: no text found]")

                if ocr_parts:
                    if total_pages > ocr_max:
                        ocr_parts.append(f"\n[Note: {total_pages - ocr_max} additional page(s) not OCR processed — total document is {total_pages} pages]")
                    return "\n\n".join(ocr_parts), True
            except ImportError:
                logger.warning("pdf2image not available for OCR fallback")
            except Exception as ocr_err:
                logger.warning(f"OCR fallback failed for {filename}: {ocr_err}")

            return "[No text could be extracted from this PDF]", False

        return "\n\n".join(text_parts), False

    except Exception as e:
        logger.error(f"PDF extraction failed for {filename}: {e}")
        return f"[PDF extraction error: {str(e)[:200]}]", False


def _preprocess_for_ocr(image: Image.Image) -> Image.Image:
    """Preprocess image for better OCR accuracy on legal documents."""
    try:
        # Convert to grayscale
        if image.mode != 'L':
            image = image.convert('L')
        # Enhance contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        # Sharpen
        image = image.filter(ImageFilter.SHARPEN)
        return image
    except Exception:
        return image


def _extract_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from DOCX including paragraphs, tables, headers/footers."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))
            if table_text:
                parts.append(f"\n[Table {table_idx + 1}]\n" + "\n".join(table_text))

        # Extract headers/footers
        for section in doc.sections:
            if section.header and section.header.paragraphs:
                header_text = " ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
                if header_text:
                    parts.insert(0, f"[Header] {header_text}")
            if section.footer and section.footer.paragraphs:
                footer_text = " ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
                if footer_text:
                    parts.append(f"[Footer] {footer_text}")

        return "\n\n".join(parts) if parts else f"[No text extracted from {filename}]"
    except Exception as e:
        logger.error(f"DOCX extraction failed for {filename}: {e}")
        return f"[DOCX extraction error: {str(e)[:200]}]"


def _extract_image(file_bytes: bytes, filename: str) -> str:
    """Extract text from image using OCR with preprocessing."""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        image = _preprocess_for_ocr(image)
        text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')
        return text.strip() if text.strip() else f"[No text could be extracted from image {filename}]"
    except Exception as e:
        logger.error(f"Image OCR failed for {filename}: {e}")
        return f"[Image OCR error: {str(e)[:200]}]"


def _extract_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file."""
    for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']:
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, AttributeError):
            continue
    return file_bytes.decode('utf-8', errors='replace')


def extract_mentions(text: str) -> List[str]:
    """Extract @mentions from text for notes collaboration."""
    if not text:
        return []
    mentions = [m.strip().lower() for m in MENTION_PATTERN.findall(text)]
    return sorted(list(set([m for m in mentions if m])))

