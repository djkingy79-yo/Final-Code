# ===========================================================================
# DO NOT UNDO — ENTIRE FILE PROTECTED
# Shared export footer utilities for all PDF and DOCX exports.
# All formatting standards, font sizes, footer text, and OOXML field code
# logic in this file are approved by Deb King and must be preserved.
# Do not remove imports, rename functions, or refactor without explicit instruction.
# Pay special attention to `docx.oxml.ns.qn` — it is REQUIRED for DOCX page-number
# field attributes; removing it will crash all Word exports.
# ===========================================================================
# Implements the standardised footer format:
# "Criminal Law Appeal Management / [Doc Name] — [Defendant] — [Date]" + "Page X of Y"
# Font: Times New Roman, Italic, 7pt, color #475569.

from datetime import datetime, timezone


def format_export_date(dt=None):
    """Format date as short DD/MM/YYYY Australian format (locked 24 Feb 2026 spec)."""
    if dt is None:
        dt = datetime.now(timezone.utc)
    if isinstance(dt, str):
        try:
            dt = datetime.fromisoformat(dt.replace('Z', '+00:00'))
        except Exception:
            dt = datetime.now(timezone.utc)
    return f"{dt.day:02d}/{dt.month:02d}/{dt.year}"


MISSING_CASE_NUMBER_LABEL = "No case number"


def build_footer_label(case: dict, doc_type: str, generated_at=None) -> str:
    """Canonical footer label (locked 24 Feb 2026 by owner).

    Format:
        LEFT  : 'Criminal Law Appeal Management / {doc} / {appellant} / {case_number}'
        RIGHT : '{DD/MM/YYYY} · Page X of Y'  (painted by NumberedCanvas)

    This function returns the LEFT half only. Missing case number is rendered
    as 'No case number' — never left blank.
    """
    appellant = (case.get("defendant_name") or case.get("title") or "Appellant").strip()
    case_number = (case.get("case_number") or "").strip() or MISSING_CASE_NUMBER_LABEL
    doc_type = (doc_type or "Legal Document").strip()
    return f"Criminal Law Appeal Management / {doc_type} / {appellant} / {case_number}"


# ============ ReportLab PDF "Page X of Y" Canvas ============

class NumberedCanvas:
    """Factory that produces a Canvas subclass rendering 'Page X of Y' footers
    after all pages have been built (two-pass approach)."""

    def __init__(self, footer_label):
        self._footer_label = footer_label
        
    def __call__(self, filename, **kwargs):
        from reportlab.pdfgen.canvas import Canvas as _Canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors

        outer = self

        class _Numbered(_Canvas):
            def __init__(self, *args, **kw):
                _Canvas.__init__(self, *args, **kw)
                self._saved_page_states = []

            def showPage(self):
                # Save page state and start a fresh page (don't commit yet)
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                total = len(self._saved_page_states)
                for idx, state in enumerate(self._saved_page_states):
                    self.__dict__.update(state)
                    self._draw_footer(idx + 1, total)
                    _Canvas.showPage(self)
                _Canvas.save(self)

            def _draw_footer(self, page_num, total_pages):
                self.saveState()
                footer_y = 10 * mm
                line_y = 14 * mm
                self.setStrokeColor(colors.HexColor('#1d4ed8'))
                self.setLineWidth(0.6)
                self.line(20 * mm, line_y, A4[0] - 20 * mm, line_y)
                self.setFillColor(colors.HexColor('#334155'))
                # Canonical print spec (locked 24 Feb 2026) — 8pt Times-Italic
                self.setFont('Times-Italic', 8)
                label = outer._footer_label
                # RIGHT side: "DD/MM/YYYY · Page X of Y" — date prefix computed
                # at footer draw time so every regenerated export carries the
                # actual generation date, not the request timestamp.
                from services.export_footer import format_export_date
                page_str = f"{format_export_date()}  \u00B7  Page {page_num} of {total_pages}"
                # Truncate label if too long to fit
                max_w = A4[0] - 40 * mm - self.stringWidth(page_str, 'Times-Italic', 8) - 8 * mm
                while self.stringWidth(label, 'Times-Italic', 8) > max_w and len(label) > 30:
                    label = label[:-4] + "..."
                self.drawString(20 * mm, footer_y, label)
                self.drawRightString(A4[0] - 20 * mm, footer_y, page_str)
                self.restoreState()

        return _Numbered(filename, **kwargs)


# ============ python-docx DOCX Footer ============

def apply_docx_footer(doc, footer_label):
    """Apply the standardised footer to all sections of a python-docx Document.

    Layout (locked 24 Feb 2026):
        LEFT tab   : '{footer_label}'                      — canonical 4-field legal identifier
        RIGHT tab  : 'DD/MM/YYYY · Page X of Y'             — short AU date + page counter fields

    Typography: Times New Roman, italic, 8pt, #334155.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    # DO NOT REMOVE — `qn` is used below to set OOXML attributes on the
    # page-number field runs. Removing it will crash DOCX exports.
    from docx.oxml.ns import qn  # noqa: F401 — used via qn() calls below
    from docx.oxml import OxmlElement
    from docx.shared import Inches

    def add_field(paragraph, field_code):
        run = paragraph.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = f' {field_code} '
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_end)
        return run

    def style_run(run):
        run.font.name = 'Times New Roman'
        run.font.size = Pt(8)            # Canonical footer spec (locked 24 Feb 2026)
        run.font.italic = True
        run.font.color.rgb = RGBColor(51, 65, 85)

    date_str = format_export_date()

    for section in doc.sections:
        section.footer_distance = Inches(0.4)
        footer = section.footer
        footer_para = footer.paragraphs[0]
        # Centre-tab disabled; right-tab at page width so date+page sits on the right edge.
        tab_stops = footer_para.paragraph_format.tab_stops
        tab_stops.clear_all()
        tab_stops.add_tab_stop(section.page_width - section.left_margin - section.right_margin,
                               WD_TAB_ALIGNMENT.RIGHT)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # LEFT half — the 4-field canonical identifier.
        label_run = footer_para.add_run(footer_label)
        style_run(label_run)

        # Separator (tab), then RIGHT half starts: "DD/MM/YYYY · Page X of Y".
        tab_run = footer_para.add_run('\t')
        style_run(tab_run)
        date_run = footer_para.add_run(f"{date_str}  \u00B7  Page ")
        style_run(date_run)
        page_field_run = add_field(footer_para, 'PAGE')
        style_run(page_field_run)
        of_run = footer_para.add_run(" of ")
        style_run(of_run)
        numpages_field_run = add_field(footer_para, 'NUMPAGES')
        style_run(numpages_field_run)


# ============ Report Type Labels ============

REPORT_TYPE_LABELS = {
    "quick_summary": "Case Summary Report (Free)",
    "full_detailed": "Full Detailed Report ($150.00)",
    "extensive_log": "Extensive Log Report ($200.00)",
    "barrister_view": "Barrister View Report",
}

DOC_TYPE_LABELS = {
    "timeline": "Timeline of Events",
    "grounds": "Grounds of Merit Report ($99.00)",
    "notes": "Notes",
    "documents": "Uploaded Case Documents",
    "legal_framework": "Legal Framework",
    "progress": "Progress",
    "case_export": "Case Export Pack",
    "translated": "Translated Report",
}
